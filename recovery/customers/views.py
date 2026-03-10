from __future__ import annotations


# === inventory filename helpers (ASCII only) ===
def _inv_center_slug(request: HttpRequest) -> str:
    """사업자명(센터명)을 파일명에 안전하게 넣기 위한 ASCII 슬러그"""
    try:
        from django.utils.text import slugify
    except Exception:
        slugify = None
    name = ""
    try:
        prof = getattr(request.user, "profile", None)
        name = getattr(prof, "business_name", "") or getattr(prof, "center_name", "") or ""
    except Exception:
        name = ""
    if slugify:
        s = slugify(name, allow_unicode=False) or ""
        s = s.strip("-_")
    else:
        s = re.sub(r"[^A-Za-z0-9_\-]+", "_", name).strip("_-")
    return s or "center"

def _inv_set_attachment(resp: HttpResponse, filename_ascii: str) -> HttpResponse:
    """Content-Disposition를 ASCII filename으로만 고정(브라우저 호환)"""
    safe = re.sub(r"[^A-Za-z0-9._\-]+", "_", filename_ascii)
    resp["Content-Disposition"] = f'attachment; filename="{safe}"'
    return resp

import datetime
import calendar
import json
import re
import os
import base64
import csv
import urllib.parse
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db.models import Q, Max
from django.db import transaction
from django.views.decorators.http import require_POST
from django.http import HttpRequest, HttpResponse, HttpResponseForbidden, HttpResponseBadRequest
from django.shortcuts import get_object_or_404, redirect, render
from django.utils.http import url_has_allowed_host_and_scheme
from django.http import JsonResponse
from django.utils import timezone

import io
import html
from .models import (
    Customer,
    CustomerCase,
    Consultation,
    ConsultationReservationChangeLog,
    AfterService,
    AfterServiceEvent,
    RRNAccessLog,
    RRNEditedLog,
    PaymentItem,
    PaymentTransaction,
    SalesDownloadLog,
    InventoryUnit,
    InventoryProductModel,
    InventoryStockEvent,
    _add_years_safe,
    BusinessProfile,
    CenterEvent,
    CenterEventLog,
    DocumentDownloadLog,
)
from .forms import (
    CustomerCreateForm,
    CustomerInfoInlineForm,
    RRNEditForm,
    CustomerExamForm,
    CaseProductPaymentForm,
    PaymentItemForm,
    PaymentTransactionForm,
    CaseNhisForm,
    CaseFollowupForm,
    BusinessProfileForm,
    ConsultationForm,
    AfterServiceForm,
)


# ==========================
# A/S 매출 집계(공통 헬퍼)
# ==========================
# 부가세 매출자료 다운로드에서 A/S 매출을 합산할 때 사용합니다.
# - 결제일 기준: paid_at
# - 환불: refund_after_service에서 refund_amount/refund_at 누적
def as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date) -> tuple[int, set[str]]:
    # 유상 A/S 매출은 '취소' 건은 제외합니다.
    qs = (
        AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            amount__gt=0,
            paid_at__isnull=False,
            paid_at__gte=ps,
            paid_at__lte=pe,
        )
        .exclude(status="CANCELED")
    )
    total = sum(int(a.amount or 0) for a in qs)
    taxes = set((a.tax_type or "") for a in qs)
    return total, taxes


def as_refund_sum_by_pay(ps: datetime.date, pe: datetime.date) -> tuple[int, set[str]]:
    qs = (
        AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            refund_amount__gt=0,
            refund_at__isnull=False,
            refund_at__gte=ps,
            refund_at__lte=pe,
        )
        .exclude(status="CANCELED")
    )
    total = sum(int(a.refund_amount or 0) for a in qs)
    taxes = set((a.tax_type or "") for a in qs)
    return total, taxes


def _rrn_reveal_until_key(customer_id: int) -> str:
    return f"rrn_reveal_until_customer_{customer_id}"


def _rrn_is_revealed(request: HttpRequest, customer_id: int) -> bool:
    until = request.session.get(_rrn_reveal_until_key(customer_id))
    if not until:
        return False
    try:
        until_ts = float(until)
    except (TypeError, ValueError):
        request.session.pop(_rrn_reveal_until_key(customer_id), None)
        return False

    now_ts = timezone.now().timestamp()
    if now_ts >= until_ts:
        request.session.pop(_rrn_reveal_until_key(customer_id), None)
        return False
    return True


def _ensure_base_payment_item(case: CustomerCase) -> None:
    qs = PaymentItem.objects.filter(case=case).order_by("created_at", "id")
    if not qs.exists():
        PaymentItem.objects.create(case=case, is_base=True)
        return
    if not qs.filter(is_base=True).exists():
        first = qs.first()
        first.is_base = True
        first.save(update_fields=["is_base"])



def _build_payment_groups(case: CustomerCase):
    """제품/결제 탭 '결제 내역' 표시용 그룹을 만듭니다.

    요구사항
    - 결제(양수) 1건 아래에 해당 결제의 환불(음수)을 '하위 줄'로 표시
    - 결제 목록은 최신이 위
    - 번호는 결제(양수)만 대상으로, 최신이 가장 큰 번호(상단)
    - 환불 줄에는 'n번 원거래...' 같은 메모를 쓰지 않고,
      환불 배지 + 환불일시(created_at) + 환불 사유(refund_reason)를 표시
    """

    pay_qs = PaymentTransaction.objects.filter(case=case, amount__gt=0).order_by("-paid_at", "-id")
    payments = list(pay_qs)
    total = len(payments)

    groups = []
    for idx, p in enumerate(payments):
        # 최신 결제가 가장 큰 번호
        p.display_seq = total - idx
        refunds = list(
            PaymentTransaction.objects.filter(case=case, origin_tx=p, amount__lt=0)
            .order_by("created_at", "id")
        )

        # 부분 환불을 고려한 '잔여 환불 가능 금액'
        refunded_total = 0
        for r in refunds:
            refunded_total += abs(int(r.amount or 0))
        p.refundable_remaining = max(int(p.amount or 0) - refunded_total, 0)

        groups.append({"payment": p, "refunds": refunds})

    # 과거 데이터 등으로 origin_tx가 없는 환불이 있을 수 있습니다.
    # 이 경우에도 화면에서 누락되지 않도록 '기타 환불'로 별도 그룹에 붙입니다.
    orphan_refunds = list(
        PaymentTransaction.objects.filter(case=case, amount__lt=0, origin_tx__isnull=True)
        .order_by("-created_at", "-id")
    )
    if orphan_refunds:
        groups.append({"payment": None, "refunds": orphan_refunds})

    return groups



def _get_latest_case(customer: Customer) -> CustomerCase | None:
    return (
        CustomerCase.objects.filter(customer=customer)
        .order_by("-purchase_date", "-created_at", "-id")
        .first()
    )


def _is_payment_stage_pending(case: CustomerCase | None) -> bool:
    """제품/결제 단계 진행 판정(사용자 기준 2026-02-14)

    ✅ '제품 선택 후 저장'이 되면 다음 단계(공단)로 이동해야 합니다.
    - 결제방식/결제내역 입력 여부는 '현재 단계' 판정 기준에서 제외합니다.

    최소 충족 조건(메인):
      - 제조사, 모델명, 착용일/구매일, 좌/우
    """
    if case is None:
        return True

    required = [
        (case.manufacturer or "").strip(),
        (case.model_name or "").strip(),
        case.purchase_date,
        (case.side or "").strip(),
    ]
    return not all(bool(x) for x in required)


def _is_nhis_stage_pending(case: CustomerCase | None) -> bool:
    if case is None:
        return True
    # 사용자 기준(2026-02-14): 공단 단계 완료는 '검수 + 입금일 + 입금액' 입력
    return (case.nhis_inspection_date is None) or (case.nhis_deposit_date is None) or (case.nhis_deposit_amount is None)


def _is_direct_purchase_payment_complete(case: CustomerCase | None) -> bool:
    """직접구매 완료 판정

    - 결제를 안했는데 '완료'로 표시되는 문제 방지
    - 직접구매는 제품/결제 단계가 끝이므로, 실제 수납(결제) 완료가 있어야 '완료'로 판단합니다.

    완료 조건(안전 기준):
      1) grand_total > 0
      2) 순수납(결제-환불) >= grand_total
      3) 양수 결제 내역(결제)이 1건 이상 존재
    """
    if case is None:
        return False

    try:
        grand_total = int(getattr(case, "grand_total", 0) or 0)
    except Exception:
        grand_total = 0

    if grand_total <= 0:
        return False

    qs = PaymentTransaction.objects.filter(case=case)
    try:
        net_paid = sum(int(t.amount or 0) for t in qs)
        has_positive_payment = qs.filter(amount__gt=0).exists()
    except Exception:
        net_paid = 0
        has_positive_payment = False

    return bool(has_positive_payment and net_paid >= grand_total)


def _case_receivable_total(case: 'CustomerCase | None') -> int:
    """수납/미수 기준 총액(공단 인정금액 제외)

    ✅ 가장 안전한 기준
    - 저장값(copay_amount)이 공단 포함으로 꼬여 있어도 항상 공단이 빠지도록
      grand_total(공단+본인부담+서브) - nhis_amount(공단 인정금액) 로 계산합니다.
    - 서브 포함 여부는 grand_total에 이미 반영되어 있으므로 별도 분기 불필요.
    """
    if case is None:
        return 0
    try:
        grand_total = int(getattr(case, 'grand_total', 0) or 0)
    except Exception:
        grand_total = 0
    try:
        nhis = int(getattr(case, 'nhis_amount', 0) or 0)
    except Exception:
        nhis = 0

    receivable = grand_total - nhis
    return receivable if receivable > 0 else 0


def compute_customer_stage(customer: Customer) -> str:
    """
    사용자 정의(2026-02-05):
    - [검사] 장애도 미입력 => 검사 (단, 직접구매는 검사 단계 자체 생략)
    - [제품/결제] 기본 결제방식 미입력 => 제품/결제
    - [공단] (직접구매 제외) 공단 입금일/입금액 미입력 => 공단
    - [후기적합] 공단 입금일/입금액 입력 완료 => 후기적합
    - ✅ 직접구매: 공단 단계 자체 스킵 (결제 완료 후 바로 후기적합)
    """
    # 1) 검사 단계
    # 사용자 기준(2026-02-14): '장애도' + '구분(track)' 입력 후 저장 시 다음 단계
    if customer.track != "직접구매":
        if (customer.exam_disability_level or "").strip() == "" or (customer.track or "").strip() == "":
            return "검사"

    latest_case = _get_latest_case(customer)

    # 2) 제품/결제 단계
    if _is_payment_stage_pending(latest_case):
        return "제품/결제"

    # 3) 공단 단계 (직접구매는 스킵)
    if customer.track != "직접구매":
        if _is_nhis_stage_pending(latest_case):
            return "공단"

    # 4) 후기적합 / 완료
    # ✅ 직접구매: 결제(수납) 완료가 확인되어야만 '완료'
    if customer.track == "직접구매":
        return "완료" if _is_direct_purchase_payment_complete(latest_case) else "제품/결제"

    return "후기적합"


def update_customer_stage(customer: Customer) -> None:
    new_stage = compute_customer_stage(customer)
    if customer.stage != new_stage:
        customer.stage = new_stage
        customer.save(update_fields=["stage"])


def _format_won(n: int) -> str:
    try:
        return f"{int(n):,}"
    except Exception:
        return "0"


def _chart_paths(values: list[int], labels: list[str] | None = None) -> tuple[str, str, list[dict], list[str]]:
    """값 배열 -> SVG path/points/labels 생성(대시보드용)."""
    # viewBox: 0 0 720 240, plot area x:[10..710], y:[40..180]
    # ✅ 좌/우 여백을 최소화해 그래프/하단 라벨이 카드 폭을 꽉 채우도록 고정합니다.
    x0, x1 = 10, 710
    y_top, y_bot = 40, 180
    n = max(1, len(values))
    step = (x1 - x0) / max(1, n - 1)

    vmax = max(values) if values else 0
    vmin = min(values) if values else 0
    if vmax == vmin:
        vmax = vmin + 1

    points: list[dict] = []
    for i, v in enumerate(values):
        x = x0 + step * i
        # 값이 클수록 위로
        y = y_bot - (((v - vmin) / (vmax - vmin)) * (y_bot - y_top))
        points.append({"x": round(x, 2), "y": round(y, 2)})

    if not points:
        points = [{"x": x0, "y": y_bot}]

    # line path
    d_line = "M " + " L ".join([f"{p['x']} {p['y']}" for p in points])

    # area path(하단 닫기)
    d_area = d_line + f" L {points[-1]['x']} {y_bot} L {points[0]['x']} {y_bot} Z"

    if labels is None or len(labels) != len(values):
        labels = ["" for _ in range(len(values))]

    return d_line, d_area, points, labels


@login_required
def dashboard(request: HttpRequest) -> HttpResponse:
    """홈(대시보드)"""
    real_today = timezone.localdate()
    today = real_today  # 기존 로직 호환용

    # --- 대시보드 차트 옵션
    period = (request.GET.get("period") or "week").strip().lower()   # day | week | month
    metric = (request.GET.get("metric") or "sales").strip().lower()  # sales | count
    if period not in ("day", "week", "month"):
        period = "week"
    if metric not in ("sales", "count"):
        metric = "sales"

    # 차트 기준일(네비게이션)
    # anchor=YYYY-MM-DD (없으면 오늘)
    anchor = real_today
    anchor_raw = (request.GET.get("anchor") or "").strip()
    if anchor_raw:
        try:
            anchor = datetime.date.fromisoformat(anchor_raw)
        except ValueError:
            anchor = real_today

    def _add_months(d: datetime.date, delta: int) -> datetime.date:
        y = d.year
        m = d.month + delta
        while m < 1:
            y -= 1
            m += 12
        while m > 12:
            y += 1
            m -= 12
        last_day = calendar.monthrange(y, m)[1]
        return datetime.date(y, m, min(d.day, last_day))

    def _add_years(d: datetime.date, delta: int) -> datetime.date:
        y = d.year + delta
        last_day = calendar.monthrange(y, d.month)[1]
        return datetime.date(y, d.month, min(d.day, last_day))

    def _bucket_setup(base: datetime.date):
        """
        period 기준(요청 확정):
          - day  : base 하루만
          - week : base가 속한 주(일~토) 7일
          - month: base가 속한 '연도'의 1~12월(월 단위 12개)
        """
        if period == "day":
            keys = [base]
            labels = [base.strftime("%Y.%m.%d")]
            period_label = base.strftime("%Y.%m.%d")
            return keys, labels, "일", period_label

        if period == "week":
            # 주 시작: 일요일, 주 끝: 토요일
            week_start = base - datetime.timedelta(days=(base.weekday() + 1) % 7)
            keys = [week_start + datetime.timedelta(days=i) for i in range(7)]
            labels = [k.strftime("%m/%d") for k in keys]
            period_label = f"{week_start.strftime('%Y.%m.%d')}~{(week_start + datetime.timedelta(days=6)).strftime('%Y.%m.%d')}"
            return keys, labels, "주", period_label

        # month (연도 기준 12개월)
        year = base.year
        keys = [datetime.date(year, m, 1) for m in range(1, 13)]
        # ✅ 월 표시: "1월, 2월, ..." (숫자만 표시 금지)
        labels = [f"{m}월" for m in range(1, 13)]
        period_label = f"{year}"
        return keys, labels, "월", period_label

    keys, labels, period_ko, period_label = _bucket_setup(anchor)
    bucket = {k: 0 for k in keys}

    date_min = min(keys)
    date_max = max(keys)

    cases = CustomerCase.objects.filter(customer__is_deleted=False).exclude(purchase_date__isnull=True, purchase_date_add__isnull=True)

    def _bucket_key_for_date(d: datetime.date) -> datetime.date | None:
        if period == "month":
            k = datetime.date(anchor.year, d.month, 1)
            return k if k in bucket else None
        return d if d in bucket else None

    # 집계
    for c in cases:
        if metric == "sales":
            main_total = int(c.nhis_amount or 0) + int(c.copay_amount or 0)  # 합계
            sub_total = int(c.self_pay_amount_add or 0)                      # 자부담금액(추가)
            total = main_total + sub_total

            d = c.purchase_date or c.purchase_date_add
            if d and date_min <= d <= date_max:
                k = _bucket_key_for_date(d)
                if k is not None:
                    bucket[k] += total
        else:
            if c.purchase_date and date_min <= c.purchase_date <= date_max:
                k = _bucket_key_for_date(c.purchase_date)
                if k is not None:
                    bucket[k] += 1
            if c.purchase_date_add and date_min <= c.purchase_date_add <= date_max:
                k = _bucket_key_for_date(c.purchase_date_add)
                if k is not None:
                    bucket[k] += 1

    values = [bucket[k] for k in keys]
    total_value = sum(values)

    chart_line_path, chart_area_path, chart_points, chart_labels = _chart_paths(values, labels)

    # 툴팁용 데이터(선택 지표에 따라 값 표시)
    def _fmt_value(v: int) -> str:
        if metric == "sales":
            return f"₩{v:,}"
        return f"{v}건"

    chart_point_meta: list[dict] = []
    for i, p in enumerate(chart_points[:len(values)]):
        lbl = labels[i] if i < len(labels) else ""
        v = values[i] if i < len(values) else 0
        chart_point_meta.append({
            "x": p["x"],
            "y": p["y"],
            "label": lbl,
            "value": v,
            "tip": f"{lbl} · {_fmt_value(v)}".strip(" ·"),
        })

    chart_points = chart_point_meta

    # SVG 축 라벨(tick) 좌표: 그래프 좌표와 동일한 x를 사용
    chart_ticks: list[dict] = []
    if period == "day":
        x0, x1 = 10, 710
        cx = round((x0 + x1) / 2, 2)
        chart_ticks = [{"x": cx, "label": labels[0] if labels else ""}]
    else:
        for i, p in enumerate(chart_points[:len(labels)]):
            chart_ticks.append({"x": p["x"], "label": labels[i]})

    # day(일) 전용: 네모(막대) 그래프
    chart_bars: list[dict] = []
    if period == "day":
        v = values[0] if values else 0
        vmax = max(v, 1)
        x0, x1 = 10, 710
        y_top, y_bot = 40, 180
        # 일(네모) 막대 너비를 절반 수준으로 축소
        bar_w = 120
        x = (x0 + x1) / 2 - bar_w / 2
        h = 0 if v == 0 else ((v / vmax) * (y_bot - y_top))
        y = y_bot - h
        lbl = labels[0] if labels else ""
        chart_bars = [{
            "x": round(x, 2),
            "y": round(y, 2),
            "w": round(bar_w, 2),
            "h": round(h, 2),
            "label": lbl,
            "value": v,
            "tip": f"{lbl} · {_fmt_value(v)}".strip(" ·"),
        }]

    # 기간 이동 URL(<' '오늘' '>)
    def _build_url(new_anchor: datetime.date | None) -> str:
        params = {"period": period, "metric": metric}
        if new_anchor and new_anchor != real_today:
            params["anchor"] = new_anchor.isoformat()
        return "?" + "&".join([f"{k}={v}" for k, v in params.items()])

    if period == "day":
        prev_anchor = anchor - datetime.timedelta(days=1)
        next_anchor = anchor + datetime.timedelta(days=1)
    elif period == "week":
        prev_anchor = anchor - datetime.timedelta(days=7)
        next_anchor = anchor + datetime.timedelta(days=7)
    else:
        # 월(12개월) 보기에서는 연도 이동
        prev_anchor = _add_years(anchor, -1)
        next_anchor = _add_years(anchor, 1)

    nav_prev_url = _build_url(prev_anchor)
    nav_next_url = _build_url(next_anchor)
    nav_today_url = _build_url(real_today)

    # --- 오늘 해야 할 일(사용자 확정 규칙)
    # 병원 방문: 장애도 미기입(직접구매 제외)
    # 검수 확인: 착용일/구매일 기준 30일 경과 + 검수 미입력
    # 서류 보완: 보완 내용 입력 + 보완 완료 미체크
    # 공단 접수: 검수 일자 기입 완료 + 접수일 미입력
    todo_hospital_visit = 0
    todo_inspection_confirm = 0
    todo_doc_supplement = 0
    todo_nhis_submit = 0
    todo_nhis_deposit_check = 0

    customers = Customer.objects.filter(is_deleted=False).order_by("-created_at")
    for cust in customers:
        update_customer_stage(cust)
        latest_case = _get_latest_case(cust)

        if cust.track != "직접구매" and (cust.exam_disability_level or "").strip() == "":
            todo_hospital_visit += 1

        if latest_case and latest_case.purchase_date and latest_case.nhis_inspection_date is None:
            if (real_today - latest_case.purchase_date).days >= 30:
                todo_inspection_confirm += 1

        if latest_case and (latest_case.nhis_supplement_content or "").strip() and (not latest_case.nhis_supplement_done):
            todo_doc_supplement += 1

        if latest_case and latest_case.nhis_inspection_date is not None and latest_case.nhis_submit_date is None:
            todo_nhis_submit += 1

        # 공단 입금 여부: 공단/주민센터 접수일 기준 5일 경과부터 '오늘 해야 할 일'에 노출
        # (실제 확인 기준은 7일 이후이지만, 5일 이후부터 체크 필요 건으로 먼저 노출)
        if latest_case and latest_case.nhis_submit_date is not None:
            missing_deposit = (latest_case.nhis_deposit_date is None) or (latest_case.nhis_deposit_amount in (None, ""))
            if missing_deposit:
                try:
                    if (real_today - latest_case.nhis_submit_date).days >= 5:
                        todo_nhis_deposit_check += 1
                except Exception:
                    pass

    # --- 후기적합 관리
    fu_imminent = 0
    fu_need_submit = 0
    fu_deposit_pending = 0

    for case in CustomerCase.objects.filter(customer__is_deleted=False):
        # 임박(가장 가까운 start)
        starts = []
        for n in [1, 2, 3, 4]:
            s, _e = case.followup_period(n)
            if s:
                delta = (s - today).days
                if delta >= 1:
                    starts.append(delta)
        if starts:
            nearest = min(starts)
            if 1 <= nearest <= 3:
                fu_imminent += 1

        # 차수별 상태
        for n in [1, 2, 3, 4]:
            s, e = case.followup_period(n)
            if not s or not e:
                continue
            submitted = bool(getattr(case, f"fu{n}_submitted", False))
            deposit_date = getattr(case, f"fu{n}_deposit_date")
            deposit_amount = getattr(case, f"fu{n}_deposit_amount")
            if (s <= today <= e) and (not submitted):
                fu_need_submit += 1
            if (s <= today <= e) and submitted and (deposit_date is None or deposit_amount in (None, "")):
                fu_deposit_pending += 1

    
    # --- 재신규 대상(후기 4차 end + 365일 기준)
    renew_imminent = 0  # D-5 ~ D-1
    renew_possible = 0  # D-0 이후(회차추가 전까지 계속)
    for c in Customer.objects.filter(is_deleted=False):
        latest_case = _get_latest_case(c)
        if not latest_case:
            continue
        _s4, e4 = latest_case.followup_period(4)
        if not e4:
            continue
        renew_start = e4 + datetime.timedelta(days=365)
        dday = (renew_start - today).days
        if 1 <= dday <= 5:
            renew_imminent += 1
        elif dday <= 0:
            renew_possible += 1
    renew_total = renew_imminent + renew_possible

# --- 장기 정체(기존 규칙 유지)
    stagn_exam_count = 0
    stagn_inspection_count = 0
    stagn_supplement_count = 0
    stagn_submit_count = 0
    stagn_followup_count = 0

    stagn_exam_days = 30
    stagn_inspection_days = 35
    stagn_supplement_days = 5
    stagn_submit_days = 5
    stagn_followup_days = 40

    for cust in customers:
        if cust.track != "직접구매" and (cust.exam_disability_level or "").strip() == "" and (cust.exam_hospital_name or "").strip():
            days = (today - cust.created_at.date()).days
            if days >= 30:
                stagn_exam_count += 1
                stagn_exam_days = max(stagn_exam_days, days)

        case = _get_latest_case(cust)
        if not case:
            continue

        if case.purchase_date and case.nhis_inspection_date is None:
            days = (today - case.purchase_date).days
            if days >= 35:
                stagn_inspection_count += 1
                stagn_inspection_days = max(stagn_inspection_days, days)

        if (case.nhis_supplement_content or "").strip() and (not case.nhis_supplement_done) and case.nhis_supplement_written_at:
            days = (today - case.nhis_supplement_written_at.date()).days
            if days >= 5:
                stagn_supplement_count += 1
                stagn_supplement_days = max(stagn_supplement_days, days)

        if case.nhis_inspection_date and case.nhis_submit_date is None:
            days = (today - case.nhis_inspection_date).days
            if days >= 5:
                stagn_submit_count += 1
                stagn_submit_days = max(stagn_submit_days, days)

        for n in [1, 2, 3, 4]:
            s, _e = case.followup_period(n)
            if not s:
                continue
            deposit_date = getattr(case, f"fu{n}_deposit_date")
            days = (today - s).days
            if days >= 40 and deposit_date is None:
                stagn_followup_count += 1
                stagn_followup_days = max(stagn_followup_days, days)

    # ✅ 템플릿에서 money 필터를 적용하므로 여기서는 '숫자' 그대로 전달합니다.
    # (문자열/₩/콤마가 섞이면 money 필터가 0으로 떨어져 '매출 동기화'가 깨집니다.)

    # --- A/S (대시보드)
    as_inprog_qs = AfterService.objects.select_related("customer").filter(customer__is_deleted=False, status="IN_PROGRESS")
    as_unpaid_count = as_inprog_qs.filter(is_paid=True, payment_status="UNPAID").count()
    as_followup_count = 0
    as_rows = []
    for a in as_inprog_qs.order_by("-received_at", "-id"):
        try:
            days = (today - a.received_at).days
        except Exception:
            days = 0
        followup = bool(days > 7)
        if followup:
            as_followup_count += 1
        as_rows.append({"as": a, "customer": a.customer, "days": days, "followup": followup})

    # followup 우선 정렬
    as_rows.sort(key=lambda r: (not r["followup"], -r["days"], -(r["as"].id or 0)))
    as_rows = as_rows[:12]

    ctx = {
        "sales_total": int(total_value or 0),
        "chart_line_path": chart_line_path,
        "chart_area_path": chart_area_path,
        "chart_points": chart_points,
        "chart_labels": chart_labels,
        "chart_ticks": chart_ticks,
        "chart_bars": chart_bars,
        "nav_prev_url": nav_prev_url,
        "nav_today_url": nav_today_url,
        "nav_next_url": nav_next_url,

        "period": period,
        "metric": metric,
        "period_ko": period_ko,
        "period_label": period_label,
        "anchor": anchor,

        "todo_hospital_visit": todo_hospital_visit,
        "todo_inspection_confirm": todo_inspection_confirm,
        "todo_doc_supplement": todo_doc_supplement,
        "todo_nhis_submit": todo_nhis_submit,
        "todo_nhis_deposit_check": todo_nhis_deposit_check,

        "as_followup_count": as_followup_count,
        "as_unpaid_count": as_unpaid_count,
        "as_rows": as_rows,

        "fu_imminent": fu_imminent,
        "fu_need_submit": fu_need_submit,
        "fu_deposit_pending": fu_deposit_pending,

        "renew_total": renew_total,
        "renew_imminent": renew_imminent,
        "renew_possible": renew_possible,

        "stagn_exam_count": stagn_exam_count,
        "stagn_inspection_count": stagn_inspection_count,
        "stagn_supplement_count": stagn_supplement_count,
        "stagn_submit_count": stagn_submit_count,
        "stagn_followup_count": stagn_followup_count,

        "stagn_exam_days": stagn_exam_days,
        "stagn_inspection_days": stagn_inspection_days,
        "stagn_supplement_days": stagn_supplement_days,
        "stagn_submit_days": stagn_submit_days,
        "stagn_followup_days": stagn_followup_days,

        "as_followup_count": as_followup_count,
        "as_unpaid_count": as_unpaid_count,
        "as_rows": as_rows,
    }
    return render(request, "customers/dashboard.html", ctx)

@login_required
def dashboard_filter(request: HttpRequest, kind: str) -> HttpResponse:
    """대시보드에서 클릭한 카드/행을 조건별로 필터링한 목록."""
    today = timezone.localdate()
    q = (request.GET.get("q") or "").strip()

    base = Customer.objects.filter(is_deleted=False).order_by("-created_at")
    items: list[Customer] = []

    def _add_if_match(c: Customer, latest_case: CustomerCase | None, match: bool):
        if not match:
            return
        items.append(c)

    for c in base:
        update_customer_stage(c)
        latest_case = _get_latest_case(c)

        if kind == "todo_hospital":
            _add_if_match(c, latest_case, c.track != "직접구매" and (c.exam_disability_level or "").strip() == "")
        elif kind == "todo_inspection":
            ok = False
            if latest_case and latest_case.purchase_date and latest_case.nhis_inspection_date is None:
                ok = (today - latest_case.purchase_date).days >= 30
            _add_if_match(c, latest_case, ok)
        elif kind == "todo_supplement":
            ok = bool(latest_case and (latest_case.nhis_supplement_content or "").strip() and (not latest_case.nhis_supplement_done))
            _add_if_match(c, latest_case, ok)
        elif kind == "todo_submit":
            ok = bool(latest_case and latest_case.nhis_inspection_date is not None and latest_case.nhis_submit_date is None)
            _add_if_match(c, latest_case, ok)
        elif kind == "todo_deposit":
            ok = False
            if latest_case and latest_case.nhis_submit_date is not None:
                missing_deposit = (latest_case.nhis_deposit_date is None) or (latest_case.nhis_deposit_amount in (None, ""))
                if missing_deposit:
                    try:
                        ok = (today - latest_case.nhis_submit_date).days >= 5
                    except Exception:
                        ok = False
            _add_if_match(c, latest_case, ok)

        elif kind == "fu_imminent":
            ok = False
            if latest_case:
                starts = []
                for n in [1, 2, 3, 4]:
                    s, _e = latest_case.followup_period(n)
                    if s:
                        delta = (s - today).days
                        if delta >= 1:
                            starts.append(delta)
                if starts:
                    nearest = min(starts)
                    ok = (1 <= nearest <= 3)
            _add_if_match(c, latest_case, ok)

        elif kind == "fu_need_submit":
            ok = False
            if latest_case:
                for n in [1, 2, 3, 4]:
                    s, e = latest_case.followup_period(n)
                    if not s or not e:
                        continue
                    submitted = bool(getattr(latest_case, f"fu{n}_submitted", False))
                    if (s <= today <= e) and (not submitted):
                        ok = True
                        break
            _add_if_match(c, latest_case, ok)

        elif kind == "fu_deposit_pending":
            ok = False
            if latest_case:
                for n in [1, 2, 3, 4]:
                    s, e = latest_case.followup_period(n)
                    if not s or not e:
                        continue
                    if not (s <= today <= e):
                        continue

                    submitted = bool(getattr(latest_case, f"fu{n}_submitted", False))
                    deposit_date = getattr(latest_case, f"fu{n}_deposit_date")
                    deposit_amount = getattr(latest_case, f"fu{n}_deposit_amount")

                    if submitted and (deposit_date is None or deposit_amount in (None, "")):
                        ok = True
                        break
            _add_if_match(c, latest_case, ok)

        
        elif kind == "renew":
            ok = False
            if latest_case:
                _s4, e4 = latest_case.followup_period(4)
                if e4:
                    renew_start = e4 + datetime.timedelta(days=365)
                    dday = (renew_start - today).days
                    ok = (dday <= 5)  # 임박(D-5~D-1) + 가능(D-0 이후)
            _add_if_match(c, latest_case, ok)

        elif kind == "stagn_exam":
            ok = (c.track != "직접구매" and (c.exam_disability_level or "").strip() == "" and (c.exam_hospital_name or "").strip() and (today - c.created_at.date()).days >= 30)
            _add_if_match(c, latest_case, ok)

        elif kind == "stagn_inspection":
            ok = bool(latest_case and latest_case.purchase_date and latest_case.nhis_inspection_date is None and (today - latest_case.purchase_date).days >= 35)
            _add_if_match(c, latest_case, ok)

        elif kind == "stagn_supplement":
            ok = bool(latest_case and (latest_case.nhis_supplement_content or "").strip() and (not latest_case.nhis_supplement_done) and latest_case.nhis_supplement_written_at and (today - latest_case.nhis_supplement_written_at.date()).days >= 5)
            _add_if_match(c, latest_case, ok)

        elif kind == "stagn_submit":
            ok = bool(latest_case and latest_case.nhis_inspection_date and latest_case.nhis_submit_date is None and (today - latest_case.nhis_inspection_date).days >= 5)
            _add_if_match(c, latest_case, ok)

        elif kind == "stagn_followup":
            ok = False
            if latest_case:
                for n in [1, 2, 3, 4]:
                    s, _e = latest_case.followup_period(n)
                    if not s:
                        continue
                    deposit_date = getattr(latest_case, f"fu{n}_deposit_date")
                    if (today - s).days >= 40 and deposit_date is None:
                        ok = True
                        break
            _add_if_match(c, latest_case, ok)

    if q:
        ql = q.lower()
        items = [c for c in items if (ql in (c.name or "").lower()) or (ql in (c.phone or "").lower())]

    title_map = {
        "todo_hospital": ("오늘 해야 할 일 · 병원 방문", "장애도 미기입 고객(직접구매 제외)"),
        "todo_inspection": ("오늘 해야 할 일 · 검수 확인", "착용일/구매일 30일 경과 + 검수 미입력"),
        "todo_supplement": ("오늘 해야 할 일 · 서류 보완", "보완 내용 입력 + 보완 완료 미체크"),
        "todo_submit": ("오늘 해야 할 일 · 공단 접수", "검수 일자 기입 완료 + 접수일 미입력"),
        "todo_deposit": ("오늘 해야 할 일 · 공단 입금 여부", "접수일 5일+ (7일 이후 확인 권장)"),
        "fu_imminent": ("후기적합 관리 · 임박", "D-1~D-3"),
        "fu_need_submit": ("후기적합 관리 · 접수 필요", "기간 내 제출 필요"),
        "fu_deposit_pending": ("후기적합 관리 · 입금 예정", "제출완료 후 입금일/입금액 미입력"),
        "renew": ("후기적합 관리 · 재신규 대상", "후기 4차 종료 + 365일 기준"),
        "stagn_exam": ("장기 정체 · 검사 대기", "30일 이상 정체"),
        "stagn_inspection": ("장기 정체 · 검수 대기", "35일 이상 정체"),
        "stagn_supplement": ("장기 정체 · 서류 보완", "5일 이상 정체"),
        "stagn_submit": ("장기 정체 · 공단 접수", "5일 이상 정체"),
        "stagn_followup": ("장기 정체 · 후기 적합 진행", "40일 이상 정체"),
    }
    page_title, page_desc = title_map.get(kind, ("필터 결과", ""))

    # 대시보드 필터 화면: 검색 좌측 드롭다운(요청사항)
    filter_options = None
    if kind.startswith("todo_"):
        filter_options = [
            {"value": "todo_hospital", "label": "병원 방문"},
            {"value": "todo_inspection", "label": "검수 확인"},
            {"value": "todo_supplement", "label": "서류 보완"},
            {"value": "todo_submit", "label": "공단 접수"},
            {"value": "todo_deposit", "label": "공단 입금 여부"},
        ]
    elif kind.startswith("fu_") or kind == "renew":
        filter_options = [
            {"value": "fu_imminent", "label": "후기적합 임박"},
            {"value": "fu_need_submit", "label": "후기 접수 필요"},
            {"value": "fu_deposit_pending", "label": "입금 예정"},
            {"value": "renew", "label": "재신규 대상"},
        ]

    return render(
        request,
        "customers/stage_list.html",
        {
            "customers": items,
            "q": q,
            "page_title": page_title,
            "page_desc": page_desc,
            "table_title": page_title,
            "target_tab": "",
            "filter_options": filter_options,
            "selected_filter": kind,
        },
    )



@login_required
def customer_list(request: HttpRequest) -> HttpResponse:
    q = (request.GET.get("q") or "").strip()
    per_page = request.GET.get("per_page") or "30"
    if per_page not in ("30", "50", "100", "200"):
        per_page = "30"
    per_page_n = int(per_page)

    # 기본 고객 쿼리 (표시 대상 기준)
    base_customers_qs = Customer.objects.filter(is_deleted=False).order_by("-created_at")

    # 검색(q)은 필터 옵션 생성 기준에도 반영되어야 함(예: 검색 결과에 따라 표시되는 회차만 보이도록)
    if q:
        base_customers_for_options = base_customers_qs.filter(name__icontains=q) | base_customers_qs.filter(phone__icontains=q) | base_customers_qs.filter(
            rrn_full__icontains=q
        ) | base_customers_qs.filter(address_summary__icontains=q)
    else:
        base_customers_for_options = base_customers_qs

    # 회차 필터 옵션: 현재 표시 대상(검색 결과 기반)의 실제 케이스(회차) 기준으로 distinct
    from .models import CustomerCase, Consultation
    from django.db.models import OuterRef, Subquery

    cycle_range = (
        CustomerCase.objects.filter(customer__in=base_customers_for_options)
        .values_list("cycle_no", flat=True)
        .distinct()
        .order_by("cycle_no")
    )

    # 실제로 화면에 표시할 customers_qs: base_customers_qs에 단계(stage) / 회차 / 검색(q) / 상태 필터를 적용
    customers_qs = base_customers_qs

    filter_stage = (request.GET.get("filter_stage") or "").strip()
    if filter_stage:
        customers_qs = customers_qs.filter(stage=filter_stage)

    filter_cycle = (request.GET.get("filter_cycle") or "").strip()
    if filter_cycle and filter_cycle.isdigit():
        customers_qs = customers_qs.filter(current_cycle=int(filter_cycle))

    # 상담 상태(진행/보류/취소)는 최신 Consultation.outcome 기준으로 필터
    filter_status = (request.GET.get("filter_status") or "").strip()
    if filter_status:
        latest_outcome_sq = Subquery(
            Consultation.objects.filter(customer=OuterRef('pk')).order_by('-created_at').values('outcome')[:1]
        )
        customers_qs = customers_qs.annotate(latest_outcome=latest_outcome_sq).filter(latest_outcome=filter_status)

    if q:
        customers_qs = customers_qs.filter(name__icontains=q) | customers_qs.filter(phone__icontains=q) | customers_qs.filter(
            rrn_full__icontains=q
        ) | customers_qs.filter(address_summary__icontains=q)

    paginator = Paginator(customers_qs, per_page_n)
    page_number = request.GET.get("page") or "1"
    page_obj = paginator.get_page(page_number)

    # UI에 실제로 표시되는 페이지 항목 기준으로 우선적으로 회차 옵션을 제한합니다.
    # (요청: 현재 상황에서는 보여지는 항목이 2회차뿐이면 옵션도 2회차만 보여야 함)
    visible_cycle_nums = sorted({int(c.current_cycle) for c in page_obj.object_list if getattr(c, 'current_cycle', None) is not None})
    if visible_cycle_nums:
        cycle_range = visible_cycle_nums

    return render(
        request,
        "customers/customer_list.html",
        {"customers": page_obj.object_list, "page_obj": page_obj, "paginator": paginator, "q": q, "per_page": per_page, "filter_cycle": filter_cycle, "cycle_range": cycle_range},
    )



@login_required
def stage_exam_list(request: HttpRequest) -> HttpResponse:
    q = (request.GET.get("q") or "").strip()

    # ✅ "최종 검사 단계"만 목록에 노출: stage == '검사'
    base_qs = Customer.objects.filter(is_deleted=False).order_by("-created_at")
    pending = []
    for c in base_qs:
        update_customer_stage(c)
        if c.stage == "검사":
            pending.append(c)

    if q:
        ql = q.lower()
        pending = [c for c in pending if (ql in (c.name or "").lower()) or (ql in (c.phone or "").lower())]

    return render(
        request,
        "customers/stage_list.html",
        {
            "customers": pending,
            "q": q,
            "page_title": "검사",
            "page_desc": "장애도 미입력 고객 (직접구매 제외)",
            "table_title": "검사 단계 대상",
            "target_tab": "검사",
        },
    )


@login_required
def stage_payment_list(request: HttpRequest) -> HttpResponse:
    q = (request.GET.get("q") or "").strip()

    customers = Customer.objects.filter(is_deleted=False).order_by("-created_at")

    pending = []
    for c in customers:
        update_customer_stage(c)
        if c.stage == "제품/결제":
            pending.append(c)

    if q:
        pending = [c for c in pending if (q.lower() in (c.name or "").lower()) or (q.lower() in (c.phone or "").lower())]

    return render(
        request,
        "customers/stage_list.html",
        {
            "customers": pending,
            "q": q,
            "page_title": "제품/결제",
            "page_desc": "기본 결제방식 미입력 고객",
            "table_title": "제품/결제 단계 대상",
            "target_tab": "제품/결제",
        },
    )


@login_required
def stage_nhis_list(request: HttpRequest) -> HttpResponse:
    q = (request.GET.get("q") or "").strip()

    customers = Customer.objects.filter(is_deleted=False).order_by("-created_at")

    pending = []
    for c in customers:
        update_customer_stage(c)
        if c.stage == "공단":
            pending.append(c)

    if q:
        pending = [c for c in pending if (q.lower() in (c.name or "").lower()) or (q.lower() in (c.phone or "").lower())]

    return render(
        request,
        "customers/stage_list.html",
        {
            "customers": pending,
            "q": q,
            "page_title": "공단",
            "page_desc": "입금일/입금액 미입력 고객",
            "table_title": "공단 단계 대상",
            "target_tab": "공단",
        },
    )


@login_required
def stage_followup_list(request: HttpRequest) -> HttpResponse:
    q = (request.GET.get("q") or "").strip()

    qs = (
        CustomerCase.objects.select_related("customer")
        .filter(customer__is_deleted=False, nhis_deposit_date__isnull=False, nhis_deposit_amount__isnull=False)
        .order_by("-nhis_deposit_date", "-id")
    )

    rows = []
    for case in qs:
        customer = case.customer
        update_customer_stage(customer)
        rows.append({"customer": customer, "case": case})

    if q:
        ql = q.lower()
        rows = [r for r in rows if (ql in (r["customer"].name or "").lower()) or (ql in (r["customer"].phone or "").lower())]

    return render(request, "customers/stage_followup_list.html", {"rows": rows, "q": q})


@login_required
def after_service_list(request: HttpRequest) -> HttpResponse:
    """사이드바 A/S 탭(전체 모아보기)"""

    status = (request.GET.get("status") or "").strip() or "IN_PROGRESS"
    q = (request.GET.get("q") or "").strip()
    # 기존 체크박스 파라미터(호환)
    followup_only = (request.GET.get("followup") or "") == "1"
    unpaid_only = (request.GET.get("unpaid") or "") == "1"
    paid_only = (request.GET.get("paid") or "") == "1"

    # 신규 드롭다운(단일 선택)
    flag_filter = (request.GET.get("flag") or "").strip().upper()  # FOLLOWUP / UNPAID / PAID / ""
    if flag_filter in ("FOLLOWUP", "UNPAID", "PAID"):
        followup_only = flag_filter == "FOLLOWUP"
        unpaid_only = flag_filter == "UNPAID"
        paid_only = flag_filter == "PAID"
    else:
        # 구버전 파라미터로 들어온 경우 드롭다운 선택값을 유추
        if followup_only:
            flag_filter = "FOLLOWUP"
        elif unpaid_only:
            flag_filter = "UNPAID"
        elif paid_only:
            flag_filter = "PAID"
        else:
            flag_filter = ""
    is_paid_filter = (request.GET.get("is_paid") or "").strip()  # "1" or "0" or ""
    owner_filter = (request.GET.get("owner") or "").strip()

    qs = AfterService.objects.select_related("customer").filter(customer__is_deleted=False)
    if status in ("IN_PROGRESS", "COMPLETED", "CANCELED"):
        qs = qs.filter(status=status)
    elif status == "ALL":
        pass

    if q:
        ql = q.lower()
        qs = qs.filter(Q(customer__name__icontains=q) | Q(customer__phone__icontains=q) | Q(reason_text__icontains=q))

    if is_paid_filter in ("1", "0"):
        qs = qs.filter(is_paid=(is_paid_filter == "1"))

    if owner_filter:
        qs = qs.filter(owner__icontains=owner_filter)

    if unpaid_only:
        qs = qs.filter(is_paid=True, payment_status="UNPAID")
    if paid_only:
        qs = qs.filter(is_paid=True, payment_status="PAID")

    today = timezone.localdate()
    rows = []
    for a in qs.order_by("-received_at", "-id"):
        days = 0
        try:
            days = (today - a.received_at).days
        except Exception:
            days = 0
        followup_needed = (a.status == "IN_PROGRESS" and days > 7)
        if followup_only and not followup_needed:
            continue
        rows.append({
            "as": a,
            "customer": a.customer,
            "days": days,
            "followup_needed": followup_needed,
        })

    # KPI
    all_qs = AfterService.objects.select_related("customer").filter(customer__is_deleted=False)
    inprog_count = all_qs.filter(status="IN_PROGRESS").count()
    followup_count = 0
    unpaid_count = all_qs.filter(is_paid=True, payment_status="UNPAID").count()
    for a in all_qs.filter(status="IN_PROGRESS"):
        try:
            if (today - a.received_at).days > 7:
                followup_count += 1
        except Exception:
            continue

    # 고객 검색(신규 접수: 고객 먼저 선택)
    pick_q = (request.GET.get("pick_q") or "").strip()
    picked_customers = []
    if pick_q:
        picked_customers = list(
            Customer.objects.filter(is_deleted=False).filter(
                Q(name__icontains=pick_q) | Q(phone__icontains=pick_q) | Q(rrn_full__icontains=pick_q)
            ).order_by("-created_at")[:20]
        )

    return render(
        request,
        "customers/after_service_list.html",
        {
            "rows": rows,
            "status": status,
            "q": q,
            "followup_only": followup_only,
            "unpaid_only": unpaid_only,
            "paid_only": paid_only,
            "flag_filter": flag_filter,
            "is_paid_filter": is_paid_filter,
            "owner_filter": owner_filter,
            "kpi_inprog": inprog_count,
            "kpi_followup": followup_count,
            "kpi_unpaid": unpaid_count,
            "pick_q": pick_q,
            "picked_customers": picked_customers,
        },
    )


@login_required
def customer_create(request: HttpRequest) -> HttpResponse:
    if request.method == "POST":
        form = CustomerCreateForm(request.POST)
        if form.is_valid():
            customer = form.save()
            # Create default CustomerCase with safe defaults for non-null boolean fields
            _case_kwargs = {"customer": customer}
            _field_names = {f.name for f in CustomerCase._meta.fields}
            if "nhis_supplement_completed" in _field_names:
                _case_kwargs["nhis_supplement_completed"] = False
            if "nhis_supplement_done" in _field_names:
                _case_kwargs["nhis_supplement_done"] = False
                _case_kwargs = {"customer": customer}
            # DB/모델 상태가 꼬여있어도 NOT NULL 컬럼에 기본값을 넣어 안전하게 생성
            _case_field_names = {f.name for f in CustomerCase._meta.get_fields() if hasattr(f, "name")}
            if "nhis_supplement_completed" in _case_field_names:
                _case_kwargs["nhis_supplement_completed"] = False
            if "nhis_supplement_done" in _case_field_names:
                _case_kwargs["nhis_supplement_done"] = False

            case = CustomerCase.objects.create(**_case_kwargs)
            PaymentItem.objects.create(case=case, is_base=True)
            update_customer_stage(customer)
            return redirect(f"/customers/{customer.id}/")
    else:
        form = CustomerCreateForm()
    return render(request, "customers/customer_form.html", {"form": form, "mode": "create"})


@login_required
def customer_detail(request: HttpRequest, pk: int) -> HttpResponse:
    customer = get_object_or_404(Customer, pk=pk, is_deleted=False)
    update_customer_stage(customer)

    tab = request.GET.get("tab", "고객정보")

    # ✅ 단계 순서(확정): 고객정보 -> 상담 -> 검사 -> 제품/결제 -> 공단 -> 후기적합
    # ✅ 직접구매면 '검사'/'공단'/'후기적합' 탭은 숨김(단, 상담은 유지)
    if customer.track == "직접구매":
        allowed_tabs = ["고객정보", "상담", "제품/결제", "A/S"]
        if tab in ("검사", "공단", "후기적합"):
            tab = "제품/결제"
    else:
        allowed_tabs = ["고객정보", "상담", "검사", "제품/결제", "공단", "후기적합", "A/S"]

    if tab not in allowed_tabs:
        tab = allowed_tabs[0]


    # ✅ redirect 시 scroll/hl 같은 쿼리스트링을 보존해야 자동 스크롤이 정상 동작합니다.
    def _redirect_url(tab_name: str, case_obj=None, extra: dict | None = None) -> str:
        params: dict[str, str] = {"tab": tab_name}
        if case_obj is not None:
            params["case"] = str(getattr(case_obj, "id", case_obj))
        # 상담 탭에서 넘어올 때 scroll 파라미터가 redirect 과정에서 사라지지 않도록 보존
        for k in ("scroll", "hl"):
            v = request.GET.get(k)
            if v:
                params[k] = v
        if extra:
            for k, v in extra.items():
                if v is None:
                    continue
                params[k] = str(v)
        return f"/customers/{customer.id}/?{urllib.parse.urlencode(params, doseq=True)}"

    case_tabs_enabled = tab in ["제품/결제", "공단", "후기적합"]
    if customer.track == "직접구매":
        case_tabs_enabled = tab in ["제품/결제"]

    # 최초 1회차 케이스 보장
    if not CustomerCase.objects.filter(customer=customer).exists():
        case0 = CustomerCase.objects.create(customer=customer, cycle_no=1)
        PaymentItem.objects.create(case=case0, is_base=True)
        if customer.current_cycle != 1:
            customer.current_cycle = 1
            customer.save(update_fields=["current_cycle"])

    cases = list(CustomerCase.objects.filter(customer=customer).order_by("cycle_no", "created_at", "id"))

    # 회차 PILL: 재구매 표시 (base PaymentItem.repurchase_yn)
    try:
        base_items = {it.case_id: it for it in PaymentItem.objects.filter(case__in=cases, is_base=True)}
        for c in cases:
            setattr(c, "is_repurchase", bool(getattr(base_items.get(c.id), "repurchase_yn", False)))
    except Exception:
        for c in cases:
            setattr(c, "is_repurchase", False)


    # ✅ 상담 탭 '고객정보 요약'은 항상 '가장 최근 회차' 기준으로 표시합니다.
    # - 착용 제품 / 결제 요약 / 공단 요약 모두 최신 회차(가장 큰 cycle_no) 기준
    latest_case = cases[-1] if cases else None

    # ✅ selected_case는 "제품/결제/공단/후기적합" 탭에서만 필요한 것이 아니라,
    # 상담 탭의 "고객정보 요약"(착용제품/결제/공단 요약)에서도 항상 필요합니다.
    # 따라서 기본값(현재 회차 우선)을 항상 잡아두고,
    # 케이스 관련 탭에서는 ?case= 로 선택을 허용합니다.
    selected_case = None
    selected_case_id = request.GET.get("case")

    if cases:
        # 기본: 현재 회차 우선
        selected_case = next((c for c in cases if c.cycle_no == customer.current_cycle), None) or cases[0]

        # 케이스 탭에서만 URL 파라미터 선택 허용
        if case_tabs_enabled and selected_case_id and selected_case_id.isdigit():
            selected_case = next((c for c in cases if c.id == int(selected_case_id)), selected_case)

    # ✅ 1회차가 "비어있는 상태"에서 2회차 이상이 존재할 수 있습니다.
    # - 이 경우 1회차로 이동했을 때 안내 모달을 띄우기 위한 플래그를 제공합니다.
    case_update_needed = False
    try:
        if selected_case and cases:
            max_cycle = max([c.cycle_no for c in cases])
            if max_cycle >= 2 and int(selected_case.cycle_no) < int(max_cycle):
                c = selected_case
                important_vals = [
                    c.manufacturer, c.model_name, c.serial_number, c.standard_code,
                    c.manufacture_date, c.receiver, c.pre_fit_date, c.purchase_date,
                    c.side, c.earmold_made_date, c.nhis_amount, c.copay_amount,
                    c.manufacturer_add, c.model_name_add, c.serial_number_add, c.standard_code_add,
                    c.manufacture_date_add, c.receiver_add, c.pre_fit_date_add, c.purchase_date_add,
                    c.side_add, c.earmold_made_date_add, c.self_pay_amount_add,
                    c.nhis_inspection_date, c.nhis_center_name, c.nhis_submit_date, c.nhis_submit_method,
                    c.nhis_supplement_content, c.nhis_supplement_done, c.nhis_deposit_date, c.nhis_deposit_amount,
                    c.fu1_start_override, c.fu1_end_override, c.fu1_submitted, c.fu1_deposit_date, c.fu1_deposit_amount, c.fu1_note,
                    c.fu2_start_override, c.fu2_end_override, c.fu2_submitted, c.fu2_deposit_date, c.fu2_deposit_amount, c.fu2_note,
                    c.fu3_start_override, c.fu3_end_override, c.fu3_submitted, c.fu3_deposit_date, c.fu3_deposit_amount, c.fu3_note,
                    c.fu4_start_override, c.fu4_end_override, c.fu4_submitted, c.fu4_deposit_date, c.fu4_deposit_amount, c.fu4_note,
                ]
                has_any = any([(v not in (None, "", 0, False)) for v in important_vals])
                has_tx = PaymentTransaction.objects.filter(case=c).exists()
                case_update_needed = (not has_any and not has_tx)
    except Exception:
        case_update_needed = False

    # 요약(상담 탭)에서 사용할 '최신 회차' 케이스 (선택 케이스와 별개)
    summary_case = latest_case

    show_full_rrn = request.user.is_staff and _rrn_is_revealed(request, customer.id)
    rrn_edit_mode = request.user.is_staff and (request.GET.get("rrn_edit") == "1")

    customer_info_form = CustomerInfoInlineForm(instance=customer) if tab == "고객정보" else None
    rrn_edit_form = RRNEditForm(instance=customer) if (tab == "고객정보" and rrn_edit_mode) else None
    exam_form = CustomerExamForm(instance=customer) if tab == "검사" else None

    # 상담 탭
    consultations = list(Consultation.objects.filter(customer=customer).order_by("-created_at", "-id"))
    latest_consultation = consultations[0] if consultations else None

    # 방문예약 변경 이력(캘린더/상담 공통) — 타임라인에 합쳐서 노출
    reservation_logs = list(
        ConsultationReservationChangeLog.objects.filter(customer=customer).select_related("changed_by").order_by("-created_at", "-id")
    )

    # 상담 타임라인: 상담 기록 + 방문예약 변경 로그를 한 리스트로 합칩니다.
    timeline_items: list[dict] = []
    for c in consultations:
        timeline_items.append({"type": "consult", "obj": c, "dt": c.created_at})
    for lg in reservation_logs:
        timeline_items.append({"type": "reserve_change", "obj": lg, "dt": lg.created_at})
    timeline_items.sort(key=lambda x: (x["dt"], getattr(x["obj"], "id", 0)), reverse=True)

    consultation_form = ConsultationForm() if tab == "상담" else None
    consultation_edit_form = None
        # ✅ 트랙별 공단 인정금액 자동 기입(표시/저장 일관성 유지)
    if selected_case is not None:
        track = (getattr(customer, "track", "") or "").strip()
        fixed = None
        if track == "일반":
            fixed = 999_000
        elif track in ("의료", "차상위"):
            fixed = 1_110_000
        elif track == "직접구매":
            fixed = 0
        if fixed is not None and int(getattr(selected_case, "nhis_amount", -1) or -1) != int(fixed):
            try:
                selected_case.nhis_amount = fixed
                selected_case.save(update_fields=["nhis_amount"])
            except Exception:
                selected_case.nhis_amount = fixed

    pp_form = CaseProductPaymentForm(instance=selected_case, customer=customer) if (tab == "제품/결제" and selected_case) else None
    nhis_form = CaseNhisForm(instance=selected_case) if (tab == "공단" and selected_case) else None
    followup_form = None
    if tab == "후기적합" and selected_case and customer.track != "직접구매":
        followup_form = CaseFollowupForm(instance=selected_case)

    # A/S 탭
    # A/S 좌측 카드 정렬 규칙
    # - '신규 생성된 건'이 최상단에 고정되도록 생성일시(created_at) 기준 내림차순 정렬
    # - 사용자가 접수일(received_at)을 과거로 입력하더라도 신규 카드가 위로 오게 하기 위함
    after_services = list(
        AfterService.objects.filter(customer=customer).order_by("-created_at", "-id")
    )
    # A/S 카드 잔액(결제-환불) 계산: 템플릿에서 연산하지 않도록 여기서 주입
    for a in after_services:
        a.remaining_amount = (getattr(a, "amount", 0) or 0) - (getattr(a, "refund_amount", 0) or 0)
        if a.remaining_amount < 0:
            a.remaining_amount = 0

    # A/S 선택 규칙
    # - 사용자가 왼쪽에서 특정 A/S건을 클릭했을 때만 상세를 보여줍니다.
    # - 최초 A/S 탭 진입 시(파라미터 없음)에는 오른쪽을 비워두기 위해 자동 선택을 하지 않습니다.
    selected_as = None
    selected_as_id = request.GET.get("as_id")
    if selected_as_id and str(selected_as_id).isdigit():
        selected_as = next((a for a in after_services if a.id == int(selected_as_id)), None)

    as_new = (request.GET.get("as_new") == "1")
    as_form = None
    as_events = []
    as_payment_groups = []
    if tab == "A/S":
        # ✅ A/S '대상' 선택지는 제품/결제 탭에 등록된 좌/우를 기준으로 제한
        # - "제일 최근" 제품/결제(회차) 기준으로 제한
        #   * 메인만 존재: 메인에 등록된 좌/우만
        #   * 메인+서브 존재: 좌/우/양이
        allowed_target_sides = None
        try:
            # ✅ 안정성 강화
            # - 사용자 지시: A/S '대상'은 "가장 최근 메인/서브" 기준
            # - 실제 서비스에서는 '현재 회차(current_cycle)'가 가장 최신 회차이며,
            #   최신 회차의 제품/결제 좌우 값이 곧 A/S 대상 제한의 기준이 됩니다.
            # - 일부 데이터(구버전/초기값)에서 purchase_date 정렬(latest_case)이
            #   현재 회차와 어긋날 수 있어, 우선 current_cycle 회차를 기준으로 잡습니다.
            base = (
                CustomerCase.objects.filter(customer=customer, cycle_no=customer.current_cycle)
                .order_by("-id")
                .first()
            ) or latest_case
            main_side_raw = ((getattr(base, "side", "") or "").strip() if base else "")
            sub_side_raw = ((getattr(base, "side_add", "") or "").strip() if base else "")
            has_sub_flag = bool(getattr(base, "has_sub", False)) if base else False

            def _norm_side(v: str) -> str:
                v = (v or "").strip()
                # 메인/서브 좌우는 케이스 모델에서 '좌'/'우'로 저장되는 경우가 많습니다.
                if v in ["좌", "L", "LEFT", "Left", "left"]:
                    return "LEFT"
                if v in ["우", "R", "RIGHT", "Right", "right"]:
                    return "RIGHT"
                if v in ["양", "양이", "BOTH", "both", "Both"]:
                    return "BOTH"
                # 혹시 내부코드가 이미 LEFT/RIGHT 인 경우 대비
                u = v.upper()
                if u in ["LEFT", "RIGHT", "BOTH"]:
                    return u
                return ""

            main_side = _norm_side(main_side_raw)
            sub_side = _norm_side(sub_side_raw)

            # ✅ '서브 존재' 판정(안정성 강화)
            # - has_sub 플래그가 켜져 있어도, 수정 과정/구버전 데이터에서 서브 값이 일부 남아
            #   A/S '대상'이 좌/우 전체로 풀리는 문제를 방지합니다.
            # - side_add(좌우)가 실제 입력되어 있고,
            #   서브 제품 관련 필드 중 하나라도 값이 있을 때만 '서브 존재'로 취급합니다.
            sub_has_any_value = False
            if base and has_sub_flag:
                sub_fields = [
                    "manufacturer_add",
                    "model_name_add",
                    "serial_number_add",
                    "standard_code_add",
                    "receiver_add",
                    "pre_fit_date_add",
                    "purchase_date_add",
                    "side_add",
                    "self_pay_amount_add",
                ]
                for fn in sub_fields:
                    val = getattr(base, fn, None)
                    if isinstance(val, str):
                        if val.strip():
                            sub_has_any_value = True
                            break
                    elif val not in (None, "", 0):
                        sub_has_any_value = True
                        break

            sub_exists = bool(has_sub_flag) and bool(sub_side) and bool(sub_has_any_value)

            if sub_exists:
                # 메인+서브가 있으면 좌/우/양이(=BOTH) 노출
                allowed_target_sides = ["LEFT", "RIGHT", "BOTH"]
            else:
                # 메인만 있으면 메인에 입력된 좌/우만 노출
                if main_side == "LEFT":
                    allowed_target_sides = ["LEFT"]
                elif main_side == "RIGHT":
                    allowed_target_sides = ["RIGHT"]
                elif main_side == "BOTH":
                    allowed_target_sides = ["LEFT", "RIGHT", "BOTH"]
                else:
                    allowed_target_sides = None

            # ✅ 최후 안전장치: main_side가 인식되면 최소한 그 값만이라도 제한 적용
            if (not allowed_target_sides) and (main_side in ("LEFT", "RIGHT")):
                allowed_target_sides = [main_side]
        except Exception:
            allowed_target_sides = None

        # 오른쪽 상세 영역은
        # - as_new=1 이거나
        # - as_id가 지정된 경우에만 렌더합니다.
        if as_new:
            as_form = AfterServiceForm(initial={
                "received_at": timezone.localdate(),
                "status": "IN_PROGRESS",
                "owner": (customer.담당자 or "").strip(),
                "target_side": (allowed_target_sides[0] if allowed_target_sides else "LEFT"),
            }, allowed_target_sides=allowed_target_sides)
        elif selected_as:
            # ✅ 기존 A/S 건(완료/취소/진행 포함)은 '그 건의 target_side'만 노출
            # - 메인 제품의 좌/우가 이후에 변경되더라도 과거 A/S 기록의 대상은 고정되어야 함
            fixed_side = (getattr(selected_as, "target_side", None) or "").strip()
            fixed_allowed = [fixed_side] if fixed_side else allowed_target_sides
            as_form = AfterServiceForm(instance=selected_as, allowed_target_sides=fixed_allowed)

        if selected_as:
            as_events = list(AfterServiceEvent.objects.filter(after_service=selected_as).order_by("-created_at", "-id"))
            # A/S 결제/환불 내역(제품/결제 '결제내역'과 동일한 표 형태로 노출하기 위한 데이터)
            # - 결제: A/S 유상 금액(1건)
            # - 환불: REFUND 이벤트(여러 건) → 음수 금액으로 표시
            try:
                if selected_as and int(selected_as.amount or 0) > 0 and (bool(selected_as.is_paid) or int(getattr(selected_as, "refund_amount", 0) or 0) > 0):
                    pay_row = {
                        "created_at": selected_as.created_at,
                        "amount": int(selected_as.amount or 0),
                        "method": (selected_as.payment_method or ""),
                        "tax_type": (selected_as.tax_type or ""),
                        "reason": selected_as.get_reason_code_display(),
                        "reason_detail": (selected_as.reason_text or "").strip(),
                        "refundable_remaining": max(int(selected_as.amount or 0) - int(selected_as.refund_amount or 0), 0),
                        "as_id": selected_as.id,
                    }
                else:
                    pay_row = None

                refund_rows = []
                for ev in as_events:
                    if getattr(ev, "event_type", "") != "REFUND":
                        continue
                    msg = (getattr(ev, "message", "") or "").strip()
                    m = re.search(r"환불\s+([0-9,]+)원\s*·\s*(.+)$", msg)
                    if not m:
                        continue
                    try:
                        amt = int(m.group(1).replace(",", ""))
                    except Exception:
                        amt = 0
                    reason = (m.group(2) or "").strip()
                    if amt <= 0:
                        continue
                    refund_rows.append({
                        "created_at": ev.created_at,
                        "amount": -amt,
                        "refund_reason": reason or msg,
                    })

                if pay_row or refund_rows:
                    as_payment_groups = [{"payment": pay_row, "refunds": refund_rows}]
            except Exception:
                as_payment_groups = []


    # A/S 요약(상담 탭 상단 고객정보 요약 카드에서 사용)
    today_local = timezone.localdate()
    in_prog = [a for a in after_services if a.status == "IN_PROGRESS"]
    followup_cutoff = today_local - datetime.timedelta(days=7)
    followup_needed = [a for a in in_prog if a.received_at and a.received_at <= followup_cutoff]
    unpaid_paid = [a for a in after_services if (a.status != "CANCELED" and a.is_paid and a.payment_status == "UNPAID" and int(a.amount or 0) > 0)]
    oldest_in_progress = None
    if in_prog:
        try:
            oldest_in_progress = min([a.received_at for a in in_prog if a.received_at])
        except Exception:
            oldest_in_progress = None
    as_summary = {
        "in_progress_count": len(in_prog),
        "followup_needed_count": len(followup_needed),
        "unpaid_count": len(unpaid_paid),
        "oldest_received_at": oldest_in_progress,
    }

    payment_status = None
    payment_overdue = False
    paid_total = 0
    outstanding = 0
    base_payment_item = None
    payment_groups = []
    payment_tx_form = None
    payment_target_total = 0
    overdue_days = 0
    tx_edit = None

    payment_forms = []

    # --- 결제 요약 계산
    # 1) selected_case: 제품/결제 탭 계산/표시(케이스 선택 가능)
    # 2) summary_case: 상담 탭 상단 '고객정보 요약'은 항상 최신 회차 기준
    summary_paid_total = 0
    summary_outstanding = 0

    if selected_case:
        try:
            paid_total = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=selected_case))
            grand_total = int(getattr(selected_case, "grand_total", 0) or 0)
            # 기본(outstanding)은 기존 로직 유지(총계=공단+본인부담+서브).
            # 제품/결제 탭의 수납/미수는 아래에서 별도로 '공단 제외 총액' 기준으로 재계산합니다.
            outstanding = grand_total - paid_total
        except Exception:
            paid_total = 0
            outstanding = 0

    if summary_case:
        try:
            summary_paid_total = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=summary_case))
            summary_grand_total = int(getattr(summary_case, "grand_total", 0) or 0)
            summary_outstanding = summary_grand_total - summary_paid_total
        except Exception:
            summary_paid_total = 0
            summary_outstanding = 0

    if tab == "제품/결제" and selected_case:
        _ensure_base_payment_item(selected_case)
        items = list(PaymentItem.objects.filter(case=selected_case).order_by("created_at", "id"))
        payment_forms = [(it, PaymentItemForm(instance=it)) for it in items]

        # 수납/미수(제품/결제 탭 전용)
        base_payment_item = next((it for it in items if it.is_base), None)
        payment_groups = _build_payment_groups(selected_case)

        # ✅ 수납/미수(제품/결제 탭): 공단 인정금액을 제외한 '수납 대상 총액' 기준
        payment_target_total = _case_receivable_total(selected_case)
        # 결제금액(누적)이 수납 대상 총액을 초과(과납)하더라도
        # 수납/미수에서는 잔여금액이 음수로 내려가거나(또는 표시가 꼬이거나)
        # 미수 섹션이 노출되지 않도록 0으로 클램프합니다.
        outstanding = payment_target_total - paid_total
        if outstanding < 0:
            outstanding = 0

        # 수납/미수는 결제내역(결제/환불)을 그대로 반영한 '순결제(=결제-환불)' 기준으로 계산합니다.

        today = timezone.localdate()
        unpaid_due_date = getattr(base_payment_item, "unpaid_due_date", None) if base_payment_item else None

        if outstanding <= 0:
            payment_status = "완납"
            payment_overdue = False
        else:
            if unpaid_due_date and today > unpaid_due_date:
                payment_status = "연체"
                payment_overdue = True
            else:
                payment_status = "미수"
                payment_overdue = False

        overdue_days = 0
        if outstanding > 0 and unpaid_due_date and today > unpaid_due_date:
            overdue_days = (today - unpaid_due_date).days

        tx_edit = None
        tx_edit_id = request.GET.get("tx_edit")
        if tx_edit_id and str(tx_edit_id).isdigit():
            tx_edit = PaymentTransaction.objects.filter(id=int(tx_edit_id), case=selected_case).first()

        payment_tx_form = PaymentTransactionForm(instance=tx_edit) if tx_edit else PaymentTransactionForm()


    followups = []
    if tab == "후기적합" and selected_case:
        today = timezone.localdate()
        for n in [1, 2, 3, 4]:
            start, end = selected_case.followup_period(n)
            deposit_date = getattr(selected_case, f"fu{n}_deposit_date")
            overdue = False
            if start and (today > (start + datetime.timedelta(days=30))) and not deposit_date:
                overdue = True
            followups.append({"n": n, "start": start, "end": end, "overdue": overdue, "deposit_date": deposit_date})

    # --- 상담/단계 제어용 계산값
    is_admin = bool(request.user.is_superuser or request.user.is_staff)
    latest_outcome = (latest_consultation.outcome if latest_consultation else "")
    is_consult_hold_or_cancel = latest_outcome in ("보류", "취소")
    is_consult_cancelled = latest_outcome == "취소"

    # 리마인드(보류)
    reminder_days = 30
    try:
        profile = getattr(request.user, "business_profile", None)
        if profile and getattr(profile, "consultation_reminder_days", None):
            reminder_days = int(profile.consultation_reminder_days)
    except Exception:
        reminder_days = 30
    days_since_last_consult = None
    remind_needed = False
    if latest_consultation:
        try:
            last_day = timezone.localtime(latest_consultation.created_at).date()
        except Exception:
            last_day = timezone.localdate()
        days_since_last_consult = (timezone.localdate() - last_day).days
        if latest_outcome == "보류" and days_since_last_consult is not None and days_since_last_consult >= reminder_days:
            remind_needed = True

    # 단계 진행 여부(기록 있음/완료 판단)
    customer_info_done = all(
        (
            (customer.name or "").strip(),
            (customer.phone or "").strip(),
            (customer.rrn_full or "").strip(),
            (customer.address_summary or "").strip(),
            (customer.담당자 or "").strip(),
        )
    )
    consult_done = bool(consultations)
    # 사용자 기준(2026-02-14): 검사 단계 완료는 '장애도 + 구분(track)' 입력
    exam_done = ((customer.exam_disability_level or "").strip() != "") and ((customer.track or "").strip() != "")

    def _product_payment_done(case: CustomerCase | None) -> bool:
        if case is None:
            return False
        # 기본 결제 아이템
        base = (
            PaymentItem.objects.filter(case=case, is_base=True)
            .order_by("created_at", "id")
            .first()
        )
        base_method = (base.payment_method or "").strip() if base else ""

        required = [
            (case.manufacturer or "").strip(),
            (case.model_name or "").strip(),
            (case.serial_number or "").strip(),
            (case.standard_code or "").strip(),
            case.manufacture_date,
            (case.receiver or "").strip(),
            case.purchase_date,
            (case.side or "").strip(),
            case.copay_amount,
            base_method,
        ]
        return all(bool(x) for x in required)

    product_done = _product_payment_done(selected_case)
    nhis_done = False
    if selected_case is not None:
        nhis_done = bool(
            selected_case.nhis_inspection_date
            and selected_case.nhis_deposit_date
            and (selected_case.nhis_deposit_amount is not None)
        )

    # 탭 비활성화(진입 차단) 계산
    tab_disabled: dict[str, bool] = {}
    for t in allowed_tabs:
        tab_disabled[t] = False
    if is_consult_hold_or_cancel:
        # 상담 이후 단계 중 '미진행'인 탭만 비활성화
        if "검사" in tab_disabled and (not exam_done):
            tab_disabled["검사"] = True
        if "제품/결제" in tab_disabled and (not product_done):
            tab_disabled["제품/결제"] = True
        if "공단" in tab_disabled and (not nhis_done):
            tab_disabled["공단"] = True
        if "후기적합" in tab_disabled:
            # 후기적합은 v1에서는 진행 판정 보류, 단 미진행이면 비활성 취급
            tab_disabled["후기적합"] = True

    # 탭 UI에 사용할 href/disabled 묶음
    tab_items = []
    for t in allowed_tabs:
        href = f"?tab={t}"
        if selected_case:
            href = href + f"&case={selected_case.id}"
        tab_items.append({"name": t, "href": href, "disabled": bool(tab_disabled.get(t, False))})

    # '상담' 탭에서 다음 버튼 비활성화 여부
    next_from_consult_disabled = bool(tab_disabled.get("검사", False))
    next_from_consult_disabled_direct = bool(tab_disabled.get("제품/결제", False))

    # ✅ 필수항목 누락 안내용 팝업(저장 차단)
    required_popup = None

    def _set_required_popup(title: str, items: list[str]):
        nonlocal required_popup
        # 중복 제거 + 순서 유지
        seen = set()
        unique = []
        for it in items:
            if not it:
                continue
            if it in seen:
                continue
            seen.add(it)
            unique.append(it)
        required_popup = {"title": title, "items": unique}

    if request.method == "POST":
        action = request.POST.get("action")

        # --- 상담(상담 탭)
        if action in ("add_consultation", "edit_consultation"):
            if tab != "상담":
                return HttpResponseForbidden("상담 탭에서만 가능합니다.")

            # 취소 상태에서는 관리자만 수정/추가 가능
            if is_consult_cancelled and (not is_admin):
                return HttpResponseForbidden("취소 상태에서는 직원 계정에서 수정할 수 없습니다.")

            if action == "add_consultation":
                form = ConsultationForm(request.POST)
                if form.is_valid():
                    obj: Consultation = form.save(commit=False)
                    obj.customer = customer
                    obj.created_by = request.user
                    obj.updated_by = request.user
                    obj.save()
                    # 상담 리스트 갱신 후 단계 업데이트
                    update_customer_stage(customer)
                    return redirect(f"/customers/{customer.id}/?tab=상담")
                consultation_form = form

            if action == "edit_consultation":
                cid = request.POST.get("consultation_id")
                cobj = None
                if cid and str(cid).isdigit():
                    cobj = Consultation.objects.filter(id=int(cid), customer=customer).first()
                if cobj is None:
                    return redirect(f"/customers/{customer.id}/?tab=상담")

                form = ConsultationForm(request.POST, instance=cobj)
                if form.is_valid():
                    obj: Consultation = form.save(commit=False)
                    obj.updated_by = request.user
                    obj.save()
                    update_customer_stage(customer)
                    return redirect(f"/customers/{customer.id}/?tab=상담")
                consultation_edit_form = form

            # 폼 에러 시 아래 렌더로 진행

        # --- A/S (A/S 탭)
        if action in ("save_after_service", "complete_after_service", "cancel_after_service"):
            if tab != "A/S":
                return HttpResponseForbidden("A/S 탭에서만 가능합니다.")

            as_id = request.POST.get("as_id")
            obj = None
            if as_id and str(as_id).isdigit():
                obj = AfterService.objects.filter(id=int(as_id), customer=customer).first()

            # 생성/수정
            form = AfterServiceForm(request.POST, instance=obj)
            if form.is_valid():
                prev = None
                if obj is not None:
                    # 변경 감지용(타임라인/이벤트)
                    prev = {
                        "status": obj.status,
                        "is_paid": bool(obj.is_paid),
                        "amount": int(obj.amount or 0),
                        "payment_method": (obj.payment_method or ""),
                        "tax_type": (obj.tax_type or ""),
                        "paid_at": obj.paid_at,
                        "refund_amount": int(getattr(obj, "refund_amount", 0) or 0),
                    }
                creating = obj is None
                as_obj: AfterService = form.save(commit=False)
                as_obj.customer = customer
                if not (as_obj.owner or "").strip():
                    as_obj.owner = (customer.담당자 or "").strip()
                # 상태 버튼 처리(완료/취소는 날짜 자동 세팅)
                today = timezone.localdate()
                if action == "complete_after_service":
                    as_obj.status = "COMPLETED"
                    as_obj.completed_at = today
                    as_obj.canceled_at = None
                elif action == "cancel_after_service":
                    as_obj.status = "CANCELED"
                    as_obj.canceled_at = today
                    as_obj.completed_at = None
                else:
                    # 일반 저장: 상태에 따라 날짜 정리
                    if as_obj.status == "IN_PROGRESS":
                        as_obj.completed_at = None
                        as_obj.canceled_at = None
                    elif as_obj.status == "COMPLETED" and not as_obj.completed_at:
                        as_obj.completed_at = today
                        as_obj.canceled_at = None
                    elif as_obj.status == "CANCELED" and not as_obj.canceled_at:
                        as_obj.canceled_at = today
                        as_obj.completed_at = None

                # 기타 사유면 상세 필수(추가 방어)
                if (as_obj.reason_code == "ETC") and not (as_obj.reason_text or "").strip():
                    form.add_error("reason_text", "기타 사유를 입력해 주세요.")
                else:
                    as_obj.save()
                    # 이벤트 기록(간단 메시지)
                    if creating:
                        AfterServiceEvent.objects.create(after_service=as_obj, event_type="CREATED", message="A/S 접수")

                    # ✅ 유/무상 전환 이벤트(무상→유상): 사유 기록
                    try:
                        if (prev is not None) and (not bool(prev.get("is_paid"))) and bool(as_obj.is_paid):
                            reason = (request.POST.get("paid_toggle_reason") or "").strip()
                            if reason:
                                AfterServiceEvent.objects.create(
                                    after_service=as_obj,
                                    event_type="MEMO",
                                    message=f"유/무상 전환(무상→유상) · {reason}",
                                )
                    except Exception:
                        pass

                    # ✅ 유/무상 전환 이벤트(유상→무상): 전액 환불 상태에서만(폼에서 보장)
                    try:
                        if (prev is not None) and bool(prev.get("is_paid")) and (not bool(as_obj.is_paid)):
                            prev_amt = int(prev.get("amount", 0) or 0)
                            prev_ref = int(prev.get("refund_amount", 0) or 0)
                            if prev_amt > 0 and prev_ref >= prev_amt:
                                AfterServiceEvent.objects.create(
                                    after_service=as_obj,
                                    event_type="MEMO",
                                    message="유/무상 전환(유상→무상)",
                                )
                    except Exception:
                        pass

                    # ✅ 결제 이벤트: 유상 저장 시(신규/변경) 1회 기록
                    try:
                        if bool(as_obj.is_paid) and int(as_obj.amount or 0) > 0:
                            changed = False
                            if creating:
                                changed = True
                            elif prev is not None:
                                if (not prev.get("is_paid")) and bool(as_obj.is_paid):
                                    changed = True
                                if int(prev.get("amount", 0)) != int(as_obj.amount or 0):
                                    changed = True
                                if (prev.get("payment_method", "") or "") != (as_obj.payment_method or ""):
                                    changed = True
                                if (prev.get("tax_type", "") or "") != (as_obj.tax_type or ""):
                                    changed = True
                                if prev.get("paid_at") != as_obj.paid_at:
                                    changed = True

                            if changed:
                                amt = int(as_obj.amount or 0)
                                method = (as_obj.payment_method or "").strip()
                                tax = (as_obj.tax_type or "").strip()
                                extra = []
                                if method:
                                    extra.append(method)
                                if tax:
                                    extra.append(tax)
                                suffix = (" · " + " · ".join(extra)) if extra else ""
                                AfterServiceEvent.objects.create(
                                    after_service=as_obj,
                                    event_type="PAYMENT",
                                    message=f"결제 {amt:,}원{suffix}",
                                )
                    except Exception:
                        pass

                    if action == "complete_after_service":
                        AfterServiceEvent.objects.create(after_service=as_obj, event_type="STATUS", message="완료 처리")
                    elif action == "cancel_after_service":
                        AfterServiceEvent.objects.create(after_service=as_obj, event_type="STATUS", message="취소 처리")
                    else:
                        # 일반 저장은 타임라인 이벤트를 남기지 않습니다.
                        # (요청사항: 타임라인에는 접수/취소/결제/환불만 노출)
                        pass

                    return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={as_obj.id}")



            # 폼 에러 시 아래 렌더로 진행
            as_form = form

        # --- A/S 환불 (A/S 탭 상세의 '환불' 버튼)
        if action == "refund_after_service":
            if tab != "A/S":
                return HttpResponseForbidden("A/S 탭에서만 가능합니다.")
            as_id = request.POST.get("as_id")
            if not (as_id and str(as_id).isdigit()):
                return redirect(f"/customers/{customer.id}/?tab=A/S")
            as_obj = AfterService.objects.filter(id=int(as_id), customer=customer).first()
            if as_obj is None:
                return redirect(f"/customers/{customer.id}/?tab=A/S")

            if not as_obj.is_paid or int(as_obj.amount or 0) <= 0:
                return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={as_obj.id}")

            # 입력값
            def _to_int(v):
                try:
                    import re
                    n = int(re.sub(r"[^0-9]", "", str(v or "")) or "0")
                    return n
                except Exception:
                    return 0

            req_amt = _to_int(request.POST.get("refund_amount"))
            reason = (request.POST.get("refund_reason") or "").strip()
            # 환불일은 제품/결제 환불 UX처럼 '오늘'로 자동 처리합니다.
            if req_amt <= 0 or not reason:
                # 폼 리렌더 대신 간단히 되돌림(모달에서 필수 체크하므로 서버는 방어만)
                return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={as_obj.id}")

            # 잔여 환불 가능 금액
            current_refund = int(as_obj.refund_amount or 0)
            remain = max(int(as_obj.amount or 0) - current_refund, 0)
            if remain <= 0:
                return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={as_obj.id}")
            if req_amt > remain:
                req_amt = remain

            refund_at = timezone.localdate()

            # 누적 환불
            as_obj.refund_amount = current_refund + req_amt
            as_obj.refund_at = refund_at
            as_obj.save(update_fields=["refund_amount", "refund_at", "updated_at"])

            AfterServiceEvent.objects.create(
                after_service=as_obj,
                event_type="REFUND",
                message=f"환불 {req_amt:,}원 · {reason}",
            )

            # ✅ 환불 완료 시 A/S 건은 '취소' 처리(전액 환불 기준)
            try:
                now_refund = int(as_obj.refund_amount or 0)
                total_amt = int(as_obj.amount or 0)
                if total_amt > 0 and now_refund >= total_amt and as_obj.status != "CANCELED":
                    as_obj.status = "CANCELED"
                    as_obj.canceled_at = refund_at
                    as_obj.completed_at = None
                    as_obj.save(update_fields=["status", "canceled_at", "completed_at", "updated_at"])
                    AfterServiceEvent.objects.create(after_service=as_obj, event_type="STATUS", message="취소 처리(환불)")
            except Exception:
                pass

            # ✅ 유상 → 무상 자동 전환(전액 환불 완료 후)
            try:
                auto_unpaid = (request.POST.get("auto_unpaid") or "").strip() == "1"
                if auto_unpaid:
                    now_refund = int(as_obj.refund_amount or 0)
                    total_amt = int(as_obj.amount or 0)
                    if total_amt > 0 and now_refund >= total_amt:
                        if bool(as_obj.is_paid):
                            as_obj.is_paid = False
                            as_obj.save(update_fields=["is_paid", "updated_at"])
                        AfterServiceEvent.objects.create(
                            after_service=as_obj,
                            event_type="MEMO",
                            message="유/무상 전환(유상→무상 · 전액환불 완료)",
                        )
            except Exception:
                pass

            return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={as_obj.id}")

        # ✅ 2차 적용: 과거 회차 수정 권한
        case_edit_actions = {
            "save_product_payment",
            "toggle_sub",
            "swap_main_sub",
            "save_payment_item",
            "add_payment_item",
            "delete_payment_item",
            "save_nhis",
            "save_followups",
        }
        can_edit_past_cycles = bool(request.user.is_superuser or request.user.is_staff)
        is_past_cycle = bool(selected_case and (selected_case.cycle_no != customer.current_cycle))
        if is_past_cycle and (action in case_edit_actions) and (not can_edit_past_cycles):
            return HttpResponseForbidden("과거 회차는 직원 계정에서 수정할 수 없습니다.")

        # ✅ 직접구매 → 공단 전환 (새 회차 생성)
        if action == "convert_to_nhis":
            if not (request.user.is_superuser or request.user.is_staff):
                return HttpResponseForbidden("권한이 없습니다.")
            track_to = (request.POST.get("track_to") or "").strip()
            if track_to not in ("일반", "의료", "차상위"):
                return redirect(f"/customers/{customer.id}/?tab=고객정보")

            # ✅ 직접구매에서만 "새 회차 생성"을 동반합니다.
            if customer.track == "직접구매":
                max_cycle = (
                    CustomerCase.objects.filter(customer=customer)
                    .order_by("-cycle_no")
                    .values_list("cycle_no", flat=True)
                    .first()
                ) or 0
                new_cycle_no = int(max_cycle) + 1
                new_case = CustomerCase.objects.create(customer=customer, cycle_no=new_cycle_no)
                PaymentItem.objects.create(case=new_case, is_base=True)
                customer.track = track_to
                customer.current_cycle = new_cycle_no
                customer.save(update_fields=["track", "current_cycle"])
                update_customer_stage(customer)
                return redirect(f"/customers/{customer.id}/?tab=검사&case={new_case.id}")

            # ✅ 일반/의료/차상위 간 전환은 "구분만 변경" (기존 회차/데이터 유지)
            if customer.track != track_to:
                customer.track = track_to
                customer.save(update_fields=["track"])
                update_customer_stage(customer)
            return redirect(f"/customers/{customer.id}/?tab=고객정보")


        if action == "save_customer_info":
            form = CustomerInfoInlineForm(request.POST, instance=customer)
            if form.is_valid():
                rrn_existing = bool((customer.rrn_full or "").strip())
                if not rrn_existing:
                    new_rrn = (request.POST.get("rrn_full") or "").strip()
                    if not new_rrn:
                        form.add_error(None, "필수 항목이 누락되어 저장할 수 없습니다.")
                    else:
                        form.save()
                        old = (customer.rrn_full or "")
                        customer.rrn_full = new_rrn
                        customer.save(update_fields=["rrn_full"])
                        RRNEditedLog.objects.create(
                            user=request.user, customer=customer, old_rrn=old, new_rrn=new_rrn
                        )
                        request.session.pop(_rrn_reveal_until_key(customer.id), None)

                        update_customer_stage(customer)

                        # 구분(track) 변경 시: 현재 회차 공단 인정금액 자동 반영
                        try:
                            cur_case = CustomerCase.objects.filter(customer=customer, cycle_no=customer.current_cycle).first()
                            if cur_case is not None:
                                if customer.track == "일반":
                                    cur_case.nhis_amount = 999_000
                                elif customer.track in ("의료", "차상위"):
                                    cur_case.nhis_amount = 1_110_000
                                elif customer.track == "직접구매":
                                    cur_case.nhis_amount = 0
                                cur_case.save(update_fields=["nhis_amount"])
                        except Exception:
                            pass

                        next_tab = tab
                        if customer.track == "직접구매" and next_tab in ("검사", "공단"):
                            next_tab = "제품/결제"
                        return redirect(f"/customers/{customer.id}/?tab={next_tab}")
                else:
                    form.save()
                    update_customer_stage(customer)

                    # 구분(track) 변경 시: 현재 회차 공단 인정금액 자동 반영
                    try:
                        cur_case = CustomerCase.objects.filter(customer=customer, cycle_no=customer.current_cycle).first()
                        if cur_case is not None:
                            if customer.track == "일반":
                                cur_case.nhis_amount = 999_000
                            elif customer.track in ("의료", "차상위"):
                                cur_case.nhis_amount = 1_110_000
                            elif customer.track == "직접구매":
                                cur_case.nhis_amount = 0
                            cur_case.save(update_fields=["nhis_amount"])
                    except Exception:
                        pass

                    next_tab = tab
                    if customer.track == "직접구매" and next_tab in ("검사", "공단"):
                        next_tab = "제품/결제"
                    return redirect(f"/customers/{customer.id}/?tab={next_tab}")

            customer_info_form = form

        elif action == "save_rrn":
            if not request.user.is_staff:
                return HttpResponseForbidden("권한이 없습니다.")
            form = RRNEditForm(request.POST, instance=customer)
            if form.is_valid():
                old = (customer.rrn_full or "")
                obj = form.save()

                if customer.track == "일반":
                    obj.nhis_amount = 999_000
                elif customer.track in ("의료", "차상위"):
                    obj.nhis_amount = 1_110_000
                elif customer.track == "직접구매":
                    obj.nhis_amount = 0

                new = (obj.rrn_full or "")
                RRNEditedLog.objects.create(user=request.user, customer=customer, old_rrn=old, new_rrn=new)
                request.session.pop(_rrn_reveal_until_key(customer.id), None)
                return redirect(f"/customers/{customer.id}/?tab=고객정보")
            rrn_edit_form = form

        elif action == "reveal_rrn":
            if not request.user.is_staff:
                return HttpResponseForbidden("권한이 없습니다.")
            until = timezone.now() + timezone.timedelta(seconds=60)
            request.session[_rrn_reveal_until_key(customer.id)] = str(until.timestamp())
            RRNAccessLog.objects.create(user=request.user, customer=customer)
            return redirect(f"/customers/{customer.id}/?tab=고객정보")

        elif action == "hide_rrn":
            if not request.user.is_staff:
                return HttpResponseForbidden("권한이 없습니다.")
            request.session.pop(_rrn_reveal_until_key(customer.id), None)
            return redirect(f"/customers/{customer.id}/?tab=고객정보")

        elif action == "save_exam":
            if tab != "검사":
                return HttpResponseForbidden("검사 탭에서만 가능합니다.")
            form = CustomerExamForm(request.POST, instance=customer)
            if form.is_valid():
                obj = form.save(commit=False)
                # 결제방식 선택일(미수 기준일): 구버전 필드(payment_method)가 있을 때만 최초 선택 시 자동 기록
                if hasattr(obj, 'payment_method') and hasattr(obj, 'payment_method_selected_date'):
                    if (getattr(obj, 'payment_method') or '').strip() and getattr(obj, 'payment_method_selected_date') is None:
                        obj.payment_method_selected_date = timezone.localdate()
                obj.save()
                update_customer_stage(customer)
                return redirect(f"/customers/{customer.id}/?tab=검사")
            # 필수(장애도) 누락 시 팝업
            if form.errors and "exam_disability_level" in form.errors:
                _set_required_popup("필수 항목 누락", ["검사: 장애도"])
            exam_form = form

        elif action == "toggle_sub":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            turning_on = not selected_case.has_sub
            selected_case.has_sub = not selected_case.has_sub

            if turning_on:
                # 켜는 경우: 기본값(메인 날짜) 복사
                selected_case.pre_fit_date_add = selected_case.pre_fit_date
                selected_case.purchase_date_add = selected_case.purchase_date
            else:
                # ✅ 보청기 추가 해제 시: 서브(추가) 관련 기록을 모두 제거
                # (A/S 대상 드롭다운이 과거 서브 side_add 값을 끌고 오지 않도록)
                selected_case.manufacturer_add = ""
                selected_case.model_name_add = ""
                selected_case.serial_number_add = ""
                selected_case.standard_code_add = ""
                selected_case.manufacture_date_add = None
                selected_case.receiver_add = ""
                selected_case.receiver_manufacturer_add = ""
                selected_case.receiver_serial_number_add = ""
                selected_case.receiver_standard_code_add = ""
                selected_case.receiver_manufacture_date_add = None
                selected_case.pre_fit_date_add = None
                selected_case.purchase_date_add = None
                selected_case.side_add = ""
                selected_case.self_pay_amount_add = None

            selected_case.save()
            # ✅ 보청기 추가/해제 클릭 시 화면이 아래로 튀지 않도록 상단 고정
            return redirect(_redirect_url("제품/결제", selected_case, extra={}))


        elif action == "swap_main_sub":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            if not getattr(selected_case, "has_sub", False):
                return redirect(_redirect_url("제품/결제", selected_case))

            # 메인<->서브 스위칭
            swap_pairs = [
                ("manufacturer", "manufacturer_add"),
                ("model_name", "model_name_add"),
                ("standard_code", "standard_code_add"),
                ("serial_number", "serial_number_add"),
                ("manufacture_date", "manufacture_date_add"),
                ("receiver", "receiver_add"),
                ("receiver_manufacturer", "receiver_manufacturer_add"),
                ("receiver_serial_number", "receiver_serial_number_add"),
                ("receiver_standard_code", "receiver_standard_code_add"),
                ("receiver_manufacture_date", "receiver_manufacture_date_add"),
                ("pre_fit_date", "pre_fit_date_add"),
                ("purchase_date", "purchase_date_add"),
                ("side", "side_add"),
                ("earmold_made_date", "earmold_made_date_add"),
            ]

            # ✅ 사용자 UX: 스위칭 시 금액도 함께 스왑(초기화/공백처럼 보이는 현상 방지)
            swap_pairs.append(("copay_amount", "self_pay_amount_add"))
            for a,b in swap_pairs:
                av = getattr(selected_case, a, None)
                bv = getattr(selected_case, b, None)
                setattr(selected_case, a, bv)
                setattr(selected_case, b, av)

            selected_case.save()
            return redirect(_redirect_url("제품/결제", selected_case, extra={}))

        elif action == "add_repurchase_cycle":
            # 재구매: 새 회차 생성 + 빨간 PILL 표기(repurchase_yn)
            if customer.track == "직접구매":
                return redirect(f"/customers/{customer.id}/?tab=제품/결제")

            max_cycle = (
                CustomerCase.objects.filter(customer=customer)
                .order_by("-cycle_no")
                .values_list("cycle_no", flat=True)
                .first()
            ) or 0
            new_cycle_no = int(max_cycle) + 1
            new_case = CustomerCase.objects.create(customer=customer, cycle_no=new_cycle_no)
            base_item = PaymentItem.objects.create(case=new_case, is_base=True)
            try:
                base_item.repurchase_yn = True
                base_item.save(update_fields=["repurchase_yn"])
            except Exception:
                pass

            customer.current_cycle = new_cycle_no
            customer.save(update_fields=["current_cycle"])
            update_customer_stage(customer)
            return redirect(f"/customers/{customer.id}/?tab=제품/결제&case={new_case.id}")

        elif action == "save_product_payment":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")

            if (request.POST.get("do") or "").strip() == "clear":
                selected_case.manufacturer = ""
                selected_case.model_name = ""
                selected_case.serial_number = ""
                selected_case.standard_code = ""
                selected_case.manufacture_date = None
                selected_case.receiver = ""
                selected_case.receiver_manufacturer = ""
                selected_case.receiver_serial_number = ""
                selected_case.receiver_standard_code = ""
                selected_case.receiver_manufacture_date = None
                selected_case.pre_fit_date = None
                selected_case.purchase_date = None
                selected_case.side = ""
                selected_case.earmold_made_date = None

                if customer.track == "일반":
                    selected_case.nhis_amount = 999_000
                elif customer.track in ("의료", "차상위"):
                    selected_case.nhis_amount = 1_110_000
                elif customer.track == "직접구매":
                    selected_case.nhis_amount = 0
                else:
                    selected_case.nhis_amount = None

                selected_case.copay_amount = None

                selected_case.manufacturer_add = ""
                selected_case.model_name_add = ""
                selected_case.serial_number_add = ""
                selected_case.standard_code_add = ""
                selected_case.manufacture_date_add = None
                selected_case.receiver_add = ""
                selected_case.receiver_manufacturer_add = ""
                selected_case.receiver_serial_number_add = ""
                selected_case.receiver_standard_code_add = ""
                selected_case.receiver_manufacture_date_add = None
                selected_case.pre_fit_date_add = None
                selected_case.purchase_date_add = None
                selected_case.side_add = ""
                selected_case.self_pay_amount_add = None

                selected_case.has_sub = False
                selected_case.save()
                update_customer_stage(customer)
                return redirect(_redirect_url("제품/결제", selected_case))

            # NOTE: 일부 환경에서 date input 값이 POST 에 누락되는 현상이 있어
            #       템플릿에서 mirror(hidden) 필드를 함께 전송하고, 서버에서 우선 보정합니다.
            post = request.POST.copy()
            if not (post.get("purchase_date") or "").strip():
                v = (post.get("purchase_date_mirror") or "").strip()
                if v:
                    post["purchase_date"] = v
            if not (post.get("purchase_date_add") or "").strip():
                v = (post.get("purchase_date_add_mirror") or "").strip()
                if v:
                    post["purchase_date_add"] = v

            form = CaseProductPaymentForm(post, instance=selected_case, customer=customer)
            if form.is_valid():
                # ✅ 제품/결제(제품 정보)는 결제방식과 무관하게 저장 가능 (결제방식은 별도 저장)
                obj = form.save()
                obj.save(update_fields=["nhis_amount"])
                if obj.has_sub:
                    if obj.pre_fit_date and not obj.pre_fit_date_add:
                        obj.pre_fit_date_add = obj.pre_fit_date
                    if obj.purchase_date and not obj.purchase_date_add:
                        obj.purchase_date_add = obj.purchase_date
                    obj.save()
                return redirect(_redirect_url("제품/결제", selected_case))

            # 필수 누락 팝업(메인/서브)
            if form.errors:
                missing = []
                m = {
                    "manufacturer": "제조사",
                    "model_name": "모델명",
                    "serial_number": "제조번호",
                    "standard_code": "표준코드",
                    "manufacture_date": "제조일",
                    "receiver": "리시버",
                    "purchase_date": "착용일/구매일",
                    "side": "좌우",
                    "copay_amount": "본인부담액",
                    "manufacturer_add": "제조사",
                    "model_name_add": "모델명",
                    "serial_number_add": "제조번호",
                    "standard_code_add": "표준코드",
                    "manufacture_date_add": "제조일",
                    "receiver_add": "리시버",
                    "purchase_date_add": "착용일/구매일",
                    "side_add": "좌우",
                    "self_pay_amount_add": "자부담금액",
                }
                for key in form.errors.keys():
                    if key == "__all__":
                        continue
                    label = m.get(key)
                    if not label:
                        continue
                    group = "제품/결제(서브)" if key.endswith("_add") or key == "self_pay_amount_add" else "제품/결제(메인)"
                    missing.append(f"{group}: {label}")
                if missing:
                    _set_required_popup("필수 항목 누락", missing)
            pp_form = form

        elif action == "add_payment_item":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            _ensure_base_payment_item(selected_case)
            PaymentItem.objects.create(case=selected_case, is_base=False)
            return redirect(_redirect_url("제품/결제", selected_case))

        elif action == "save_payment_item":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            item_id = request.POST.get("item_id")
            item = get_object_or_404(PaymentItem, pk=item_id, case=selected_case)
            form = PaymentItemForm(request.POST, instance=item)
            if form.is_valid():
                obj = form.save(commit=False)
                # 결제방식 선택일(미수 기준일): 구버전 필드(payment_method)가 있을 때만 최초 선택 시 자동 기록
                if hasattr(obj, 'payment_method') and hasattr(obj, 'payment_method_selected_date'):
                    if (getattr(obj, 'payment_method') or '').strip() and getattr(obj, 'payment_method_selected_date') is None:
                        obj.payment_method_selected_date = timezone.localdate()
                obj.save()
                update_customer_stage(customer)
                return redirect(_redirect_url("제품/결제", selected_case))

            # 필수(결제방식) 누락 시 팝업
            if form.errors and "payment_method" in form.errors:
                _set_required_popup("필수 항목 누락", ["제품/결제(메인): 결제방식"])

            items = list(PaymentItem.objects.filter(case=selected_case).order_by("created_at", "id"))
            payment_forms = []
            for it in items:
                if it.id == item.id:
                    payment_forms.append((it, form))
                else:
                    payment_forms.append((it, PaymentItemForm(instance=it)))


        elif action == "save_unpaid_info":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            # 기본 결제 아이템에만 저장
            base = PaymentItem.objects.filter(case=selected_case, is_base=True).order_by("created_at", "id").first()
            if base is None:
                _ensure_base_payment_item(selected_case)
                base = PaymentItem.objects.filter(case=selected_case, is_base=True).order_by("created_at", "id").first()

            due_str = (request.POST.get("unpaid_due_date") or "").strip()
            note = (request.POST.get("unpaid_note") or "").strip()

            old_due = base.unpaid_due_date
            new_due = None
            if due_str:
                try:
                    new_due = datetime.date.fromisoformat(due_str)
                except Exception:
                    new_due = None

            base.unpaid_due_date = new_due
            base.unpaid_note = note

            # 변경 로그는 메모에 자동 1줄 추가(선택적)
            today = timezone.localdate().isoformat()
            if old_due != new_due:
                old_txt = old_due.isoformat() if old_due else "없음"
                new_txt = new_due.isoformat() if new_due else "없음"
                line = f"{today} 납부 예정일 변경 ({old_txt} → {new_txt})"
                if base.unpaid_note:
                    base.unpaid_note = base.unpaid_note + "\n" + line
                else:
                    base.unpaid_note = line

            base.save()
            return redirect(_redirect_url("제품/결제", selected_case))
        elif action == "add_payment_tx":
            # ✅ 결제 모달은 제품/결제 탭에서만 노출되지만,
            #    일부 환경에서 form submit 시 querystring(tab/case)이 누락될 수 있습니다.
            #    이 경우에도 결제 저장이 되도록 selected_case 를 최대한 복구합니다.
            if selected_case is None:
                try:
                    selected_case = summary_case
                except Exception:
                    selected_case = None
                if selected_case is None:
                    selected_case = (
                        CustomerCase.objects.filter(customer=customer)
                        .order_by("-round", "-id")
                        .first()
                    )
            if selected_case is None:
                return HttpResponseForbidden("제품/결제 케이스가 없습니다.")
            form = PaymentTransactionForm(request.POST)
            if form.is_valid():
                amount = int(form.cleaned_data.get("amount") or 0)
                payment_target_total = _case_receivable_total(selected_case)
                # 순결제(결제-환불) 기준으로 과납을 방지합니다.
                current_paid_total = sum(
                    int(t.amount or 0)
                    for t in PaymentTransaction.objects.filter(case=selected_case)
                )

                # 과납/0원 방지 (과납은 사용자 확인 후 허용)
                if amount <= 0:
                    form.add_error("amount", "1원 이상의 금액을 입력해주세요.")
                elif (current_paid_total + amount) > payment_target_total:
                    allow_overpay = (request.POST.get("allow_overpay") or "").strip()
                    if allow_overpay != "1":
                        form.add_error("amount", "결제 금액이 총 결제금액(공단 제외)을 초과합니다. 계속 진행하려면 확인이 필요합니다.")

                if not form.errors:
                    tx = form.save(commit=False)
                    tx.case = selected_case
                    # ✅ 과세 구분은 UI 미노출 / 과세로 고정
                    try:
                        if not getattr(tx, "tax_type", ""):
                            tx.tax_type = "과세"
                        else:
                            tx.tax_type = "과세"
                    except Exception:
                        pass
                    tx.save()

                    # ✅ 재고 차감: 결제 이력 생성(=결제 완료로 간주) 시 즉시 차감
                    # - 보청기(시리얼형): 제조번호 IN_STOCK -> SHIPPED
                    # - 리시버(수량형): qty_current - 1
                    # (A+C) 중복 차감 방지: 결제TX 마커 기반

                    tx_marker = f"PAYMENT_TX:{tx.id}"

                    # 1) 보청기(시리얼형)
                    serials = []

                    def _add_sn(v):
                        v = (v or "").strip()
                        if v:
                            serials.append(v)

                    _add_sn(getattr(selected_case, "serial_number", ""))
                    if getattr(selected_case, "has_sub", False):
                        _add_sn(getattr(selected_case, "serial_number_add", ""))

                    if serials:
                        try:
                            with transaction.atomic():
                                for sn in serials:
                                    u = (
                                        InventoryUnit.objects.select_for_update()
                                        .filter(serial_no=sn)
                                        .first()
                                    )
                                    if not u:
                                        continue
                                    if (u.status or "").strip() != "IN_STOCK":
                                        continue
                                    if InventoryStockEvent.objects.filter(unit=u, event_type="SHIP", reason__icontains=tx_marker).exists():
                                        continue
                                    before = _inv_unit_snapshot(u)
                                    u.status = "SHIPPED"
                                    u.save(update_fields=["status"])
                                    try:
                                        InventoryStockEvent.objects.create(
                                            unit=u,
                                            product_model=None,
                                            qty_delta=0,
                                            event_type="SHIP",
                                            progress_status="DONE",
                                            reason=f"결제(제품/결제) · {tx_marker}",
                                            before_json=json.dumps(before, ensure_ascii=False),
                                            after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                            created_by=request.user,
                                        )
                                    except Exception:
                                        pass
                        except Exception:
                            pass

                    # 2) 리시버(수량형)
                    def _decrement_receiver(model_name: str, manufacturer: str):
                        model_name = (model_name or "").strip()
                        manufacturer = (manufacturer or "").strip()
                        if not model_name or not manufacturer:
                            return
                        pm = (
                            InventoryProductModel.objects.select_related("manufacturer")
                            .filter(
                                item_type="QTY",
                                product_kind="RECEIVER",
                                manufacturer__name=manufacturer,
                                model_name=model_name,
                                is_deleted=False,
                            )
                            .first()
                        )
                        if not pm:
                            return
                        if InventoryStockEvent.objects.filter(product_model=pm, event_type="SHIP", reason__icontains=tx_marker).exists():
                            return
                        before = _inv_qty_snapshot(pm)
                        pm.qty_current = int(pm.qty_current or 0) - 1
                        pm.save(update_fields=["qty_current"])
                        _inv_check_and_notify(pm)
                        try:
                            InventoryStockEvent.objects.create(
                                unit=None,
                                product_model=pm,
                                qty_delta=-1,
                                event_type="SHIP",
                                progress_status="DONE",
                                reason=f"결제(제품/결제) · {tx_marker}",
                                before_json=json.dumps(before, ensure_ascii=False),
                                after_json=json.dumps(_inv_qty_snapshot(pm), ensure_ascii=False),
                                created_by=request.user,
                            )
                        except Exception:
                            pass

                    _decrement_receiver(getattr(selected_case, "receiver", ""), getattr(selected_case, "receiver_manufacturer", ""))
                    if getattr(selected_case, "has_sub", False):
                        _decrement_receiver(getattr(selected_case, "receiver_add", ""), getattr(selected_case, "receiver_manufacturer_add", ""))
                        return redirect(_redirect_url("제품/결제", selected_case, extra={}))

            # 폼 에러가 있으면 다시 렌더링(아래 컨텍스트를 위해 재구성)
            _ensure_base_payment_item(selected_case)
            items = list(PaymentItem.objects.filter(case=selected_case).order_by("created_at", "id"))
            payment_forms = [(it, PaymentItemForm(instance=it)) for it in items]

            base_payment_item = next((it for it in items if it.is_base), None)
            payment_groups = _build_payment_groups(selected_case)
            paid_total = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=selected_case))
            # ✅ 수납/미수(제품/결제 탭): 공단 인정금액 제외 '수납 대상 총액' 기준으로 통일
            payment_target_total = _case_receivable_total(selected_case)
            outstanding = payment_target_total - paid_total
            if outstanding < 0:
                outstanding = 0
            today = timezone.localdate()
            unpaid_due_date = getattr(base_payment_item, "unpaid_due_date", None) if base_payment_item else None

            if outstanding <= 0:
                payment_status = "완납"
                payment_overdue = False
            else:
                if unpaid_due_date and today > unpaid_due_date:
                    payment_status = "연체"
                    payment_overdue = True
                else:
                    payment_status = "미수"
                    payment_overdue = False

            overdue_days = 0
            if outstanding > 0 and unpaid_due_date and today > unpaid_due_date:
                overdue_days = (today - unpaid_due_date).days

            tx_edit = None
            payment_tx_form = form

        
        elif action == "edit_payment_tx":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            tx_id = request.POST.get("tx_id")
            tx = get_object_or_404(PaymentTransaction, pk=tx_id, case=selected_case)

            form = PaymentTransactionForm(request.POST, instance=tx)
            if form.is_valid():
                new_amount = int(form.cleaned_data.get("amount") or 0)
                payment_target_total = _case_receivable_total(selected_case)
                current_paid_total = sum(
                    int(t.amount or 0)
                    for t in PaymentTransaction.objects.filter(case=selected_case)
                )
                adjusted_paid_total = current_paid_total - int(tx.amount or 0) + new_amount

                # 환불(음수)은 음수 유지 / 일반 수납(양수)은 양수 유지
                if int(tx.amount or 0) < 0:
                    if new_amount >= 0:
                        form.add_error("amount", "환불 이력은 음수 금액으로만 수정할 수 있습니다.")
                else:
                    if new_amount <= 0:
                        form.add_error("amount", "1원 이상의 금액을 입력해주세요.")
                    elif adjusted_paid_total > payment_target_total:
                        form.add_error("amount", "결제 금액이 총계를 초과했습니다. (과납 방지)")

                if not form.errors:
                    form.save()
                    return redirect(_redirect_url("제품/결제", selected_case))

            # 폼 에러가 있으면 다시 렌더링(아래 컨텍스트를 위해 재구성)
            _ensure_base_payment_item(selected_case)
            items = list(PaymentItem.objects.filter(case=selected_case).order_by("created_at", "id"))
            payment_forms = [(it, PaymentItemForm(instance=it)) for it in items]

            base_payment_item = next((it for it in items if it.is_base), None)
            payment_groups = _build_payment_groups(selected_case)
            paid_total = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=selected_case))
            # ✅ 수납/미수(제품/결제 탭): 공단 인정금액 제외 '수납 대상 총액' 기준으로 통일
            payment_target_total = _case_receivable_total(selected_case)
            outstanding = payment_target_total - paid_total
            if outstanding < 0:
                outstanding = 0
            today = timezone.localdate()
            unpaid_due_date = getattr(base_payment_item, "unpaid_due_date", None) if base_payment_item else None

            if outstanding <= 0:
                payment_status = "완납"
                payment_overdue = False
            else:
                if unpaid_due_date and today > unpaid_due_date:
                    payment_status = "연체"
                    payment_overdue = True
                else:
                    payment_status = "미수"
                    payment_overdue = False

            overdue_days = 0
            if outstanding > 0 and unpaid_due_date and today > unpaid_due_date:
                overdue_days = (today - unpaid_due_date).days

            tx_edit = tx
            payment_tx_form = form

        elif action == "delete_payment_tx":
            return HttpResponseForbidden("수납 이력은 삭제할 수 없습니다.")

        
        elif action == "refund_after_service":
            # A/S 탭 환불 (제품/결제 탭의 환불 모달/버튼과 동일 UX)
            if tab != "A/S":
                return HttpResponseForbidden("A/S 탭에서만 가능합니다.")
            as_id = (request.POST.get("as_id") or "").strip()
            if not as_id.isdigit():
                return HttpResponseForbidden("A/S ID가 올바르지 않습니다.")
            a = get_object_or_404(AfterService, pk=int(as_id), customer=customer)

            if (not a.is_paid) or int(a.amount or 0) <= 0:
                return HttpResponseForbidden("환불 대상이 아닙니다.")
            if a.status == "CANCELED":
                return HttpResponseForbidden("취소된 A/S는 환불할 수 없습니다.")

            reason = (request.POST.get("refund_reason") or "").strip()
            if not reason:
                return HttpResponseForbidden("환불 사유가 필요합니다.")

            raw_amt = (request.POST.get("refund_amount") or "").strip()
            try:
                refund_amt = int(raw_amt)
            except Exception:
                refund_amt = 0
            if refund_amt <= 0:
                return HttpResponseForbidden("환불 금액이 올바르지 않습니다.")

            refundable_remaining = max(int(a.amount or 0) - int(a.refund_amount or 0), 0)
            if refundable_remaining <= 0:
                return HttpResponseForbidden("이미 전액 환불된 건입니다.")
            if refund_amt > refundable_remaining:
                return HttpResponseForbidden("환불 금액은 잔여 환불 가능 금액을 초과할 수 없습니다.")

            a.refund_amount = int(a.refund_amount or 0) + refund_amt
            a.refund_at = timezone.localdate()
            a.save(update_fields=["refund_amount", "refund_at", "updated_at"])

            AfterServiceEvent.objects.create(
                after_service=a,
                event_type="REFUND",
                message=f"환불 {refund_amt:,}원 · {reason}",
            )

            # A/S 화면으로 복귀
            return redirect(f"/customers/{customer.id}/?tab=A/S&as_id={a.id}")

        elif action == "refund_payment_tx":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            tx_id = request.POST.get("tx_id")
            src = get_object_or_404(PaymentTransaction, pk=tx_id, case=selected_case)
            # 원거래가 이미 환불(음수)인 경우 환불 불가
            if int(src.amount or 0) <= 0:
                return HttpResponseForbidden("환불 대상이 아닙니다.")

            reason = (request.POST.get("refund_reason") or "").strip()
            if not reason:
                return HttpResponseForbidden("환불 사유가 필요합니다.")

            raw_amt = (request.POST.get("refund_amount") or "").strip()
            try:
                refund_amt = int(raw_amt)
            except Exception:
                refund_amt = 0

            if refund_amt <= 0:
                return HttpResponseForbidden("환불 금액이 올바르지 않습니다.")

            # 기존 환불을 고려한 '잔여 환불 가능' 금액 계산
            already_refunded = 0
            for r in PaymentTransaction.objects.filter(case=selected_case, origin_tx=src, amount__lt=0):
                already_refunded += abs(int(r.amount or 0))
            refundable_remaining = int(src.amount or 0) - already_refunded

            if refundable_remaining <= 0:
                return HttpResponseForbidden("이미 전액 환불된 결제건입니다.")
            if refund_amt > refundable_remaining:
                return HttpResponseForbidden("환불 금액은 잔여 환불 가능 금액을 초과할 수 없습니다.")

            # 원거래 번호(1,2,3..): 양수 거래만 기준, 오래된 순 (내부 저장용)
            origin_seq = 0
            pos_txs = list(PaymentTransaction.objects.filter(case=selected_case, amount__gt=0).order_by("paid_at", "id"))
            for i, t in enumerate(pos_txs, start=1):
                if t.id == src.id:
                    origin_seq = i
                    break

            today = timezone.localdate()

            PaymentTransaction.objects.create(
                case=selected_case,
                paid_at=today,
                amount=-abs(refund_amt),
                method=src.method,
                tax_type=src.tax_type,
                memo="",
                origin_tx=src,
                origin_seq=origin_seq,
                refund_reason=reason,
            )

            # ✅ 재고 원복: 전액 환불(순결제=0)일 때만 원복
            net_paid_total = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=selected_case))
            full_refund_marker = f"FULL_REFUND_CASE:{selected_case.id}"

            if net_paid_total == 0 and not InventoryStockEvent.objects.filter(reason__icontains=full_refund_marker).exists():
                # 1) 보청기(시리얼형) 재입고
                serials = []
                main_sn = (getattr(selected_case, "serial_number", "") or "").strip()
                if main_sn:
                    serials.append(main_sn)
                if getattr(selected_case, "has_sub", False):
                    sub_sn = (getattr(selected_case, "serial_number_add", "") or "").strip()
                    if sub_sn:
                        serials.append(sub_sn)

                for sn in serials:
                    u = InventoryUnit.objects.filter(serial_no=sn).first()
                    if not u:
                        continue
                    if (u.status or "").strip() != "SHIPPED":
                        continue
                    before = _inv_unit_snapshot(u)
                    u.status = "IN_STOCK"
                    u.save(update_fields=["status"])
                    try:
                        InventoryStockEvent.objects.create(
                            unit=u,
                            product_model=None,
                            qty_delta=0,
                            event_type="RECEIVE",
                            progress_status="DONE",
                            reason=f"전액환불 재고원복 · {full_refund_marker}",
                            before_json=json.dumps(before, ensure_ascii=False),
                            after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                            created_by=request.user,
                        )
                    except Exception:
                        pass

                # 2) 리시버(수량형) +1
                def _restore_receiver(model_name: str, manufacturer: str):
                    model_name = (model_name or "").strip()
                    manufacturer = (manufacturer or "").strip()
                    if not model_name or not manufacturer:
                        return
                    pm = (
                        InventoryProductModel.objects.select_related("manufacturer")
                        .filter(
                            item_type="QTY",
                            product_kind="RECEIVER",
                            manufacturer__name=manufacturer,
                            model_name=model_name,
                            is_deleted=False,
                        )
                        .first()
                    )
                    if not pm:
                        return
                    before = _inv_qty_snapshot(pm)
                    pm.qty_current = int(pm.qty_current or 0) + 1
                    pm.save(update_fields=["qty_current"])
                    _inv_check_and_notify(pm)
                    try:
                        InventoryStockEvent.objects.create(
                            unit=None,
                            product_model=pm,
                            qty_delta=+1,
                            event_type="ADJUST",
                            progress_status="DONE",
                            adjust_kind="DATA_CORRECTION",
                            reason=f"전액환불 재고원복 · {full_refund_marker}",
                            before_json=json.dumps(before, ensure_ascii=False),
                            after_json=json.dumps(_inv_qty_snapshot(pm), ensure_ascii=False),
                            created_by=request.user,
                        )
                    except Exception:
                        pass

                _restore_receiver(getattr(selected_case, "receiver", ""), getattr(selected_case, "receiver_manufacturer", ""))
                if getattr(selected_case, "has_sub", False):
                    _restore_receiver(getattr(selected_case, "receiver_add", ""), getattr(selected_case, "receiver_manufacturer_add", ""))

            return redirect(_redirect_url("제품/결제", selected_case))

        elif action == "delete_payment_item":
            if tab != "제품/결제" or selected_case is None:
                return HttpResponseForbidden("제품/결제 탭에서만 가능합니다.")
            item_id = request.POST.get("item_id")
            item = get_object_or_404(PaymentItem, pk=item_id, case=selected_case)
            if item.is_base:
                return HttpResponseForbidden("기본 결제카드는 삭제할 수 없습니다.")
            item.delete()
            _ensure_base_payment_item(selected_case)
            return redirect(_redirect_url("제품/결제", selected_case))

        elif action == "save_nhis":
            if tab != "공단" or selected_case is None:
                return HttpResponseForbidden("공단 탭에서만 가능합니다.")
            form = CaseNhisForm(request.POST, instance=selected_case)
            if form.is_valid():
                obj = form.save(commit=False)

                content = (obj.nhis_supplement_content or "").strip()
                if content and obj.nhis_supplement_written_at is None:
                    obj.nhis_supplement_written_at = timezone.now()
                if not content:
                    obj.nhis_supplement_written_at = None
                    # 보완 내용이 없더라도 '보완 완료' 체크/완료일시는 유지될 수 있어야 합니다.
                    # (사용자가 단순 완료 처리만 하는 케이스)

                if obj.nhis_supplement_done:
                    if obj.nhis_supplement_done_at is None:
                        obj.nhis_supplement_done_at = timezone.now()
                else:
                    obj.nhis_supplement_done_at = None

                obj.save()
                update_customer_stage(customer)
                return redirect(f"/customers/{customer.id}/?tab=공단&case={selected_case.id}")

            # 필수(검수) 누락 시 팝업
            if form.errors and "nhis_inspection_date" in form.errors:
                _set_required_popup("필수 항목 누락", ["공단: 검수"])
            nhis_form = form

        elif action == "save_followups":
            if tab != "후기적합" or selected_case is None:
                return HttpResponseForbidden("후기적합 탭에서만 가능합니다.")
            form = CaseFollowupForm(request.POST, instance=selected_case)
            if form.is_valid():
                obj = form.save(commit=False)

                for n in [1, 2, 3, 4]:
                    submitted = bool(getattr(obj, f"fu{n}_submitted"))
                    at = getattr(obj, f"fu{n}_submitted_at")
                    if submitted:
                        if at is None:
                            setattr(obj, f"fu{n}_submitted_at", timezone.now())
                    else:
                        setattr(obj, f"fu{n}_submitted_at", None)

                obj.save()
                update_customer_stage(customer)
                return redirect(f"/customers/{customer.id}/?tab=후기적합&case={selected_case.id}")
            followup_form = form

    # A/S 유상 체크(UI 노출 여부)는 "폼 값"을 최우선으로 사용합니다.
    # - POST 검증 에러로 리렌더될 때 DB값(selected_as.is_paid)을 쓰면 사용자가 체크를 해제/변경해도
    #   유상 영역이 계속 노출되는 문제가 발생할 수 있습니다.
    as_paid_checked = False
    if as_form is not None:
        try:
            if getattr(as_form, "is_bound", False):
                as_paid_checked = bool(as_form.data.get("is_paid"))
            else:
                init_v = None
                try:
                    init_v = as_form.initial.get("is_paid")
                except Exception:
                    init_v = None
                if init_v is None and selected_as is not None:
                    init_v = bool(getattr(selected_as, "is_paid", False))
                as_paid_checked = bool(init_v)
        except Exception:
            as_paid_checked = bool(selected_as and getattr(selected_as, "is_paid", False))

    return render(
        request,
        "customers/customer_detail.html",
        {
            "customer": customer,
            "tab": tab,
            "allowed_tabs": allowed_tabs,
            "tab_items": tab_items,
            "tab_disabled": tab_disabled,
            "case_tabs_enabled": case_tabs_enabled,
            "cases": cases,
            "selected_case": selected_case,
            "case_update_needed": case_update_needed,
            # 상담 탭 상단 '고객정보 요약'은 항상 최신 회차 기준
            "summary_case": summary_case,
            "summary_payment_outstanding": summary_outstanding,
            "consultations": consultations,
            "timeline_items": timeline_items,
        "latest_consultation_id": (latest_consultation.id if latest_consultation else None),
            "latest_outcome": latest_outcome,
            "days_since_last_consult": days_since_last_consult,
            "after_services": after_services,
            "selected_as": selected_as,
            "as_is_paid": bool(selected_as and selected_as.is_paid),
            "as_paid_checked": as_paid_checked,
            "as_refund_remaining": (max(int(getattr(selected_as, "amount", 0) or 0) - int(getattr(selected_as, "refund_amount", 0) or 0), 0) if selected_as else 0),
            "as_refund_remaining": (
                max(int(getattr(selected_as, "amount", 0) or 0) - int(getattr(selected_as, "refund_amount", 0) or 0), 0)
                if selected_as else 0
            ),
            "as_form": as_form,
            "as_events": as_events,
            "as_payment_groups": as_payment_groups,
            "as_new": as_new,
            "as_summary": as_summary,
            "today": timezone.localdate(),
            "reminder_days": reminder_days,
            "remind_needed": remind_needed,
            "consultation_form": consultation_form,
            "consultation_edit_form": consultation_edit_form,
            "is_consult_cancelled": is_consult_cancelled,
            "is_superuser": bool(request.user.is_superuser),
            "is_admin": is_admin,
            "next_from_consult_disabled": next_from_consult_disabled,
            "next_from_consult_disabled_direct": next_from_consult_disabled_direct,
            "customer_info_form": customer_info_form,
            "exam_form": exam_form,
            "pp_form": pp_form,
            "nhis_form": nhis_form,
            "followup_form": followup_form,
            "followups": followups,
            "show_full_rrn": show_full_rrn,
            "payment_forms": payment_forms,
            "payment_groups": payment_groups,
            "payment_tx_form": payment_tx_form,
            "payment_paid_total": paid_total,
            "payment_target_total": payment_target_total,
            "payment_outstanding": outstanding,
            "payment_status": payment_status,
            "payment_overdue": payment_overdue,
            "payment_overdue_days": overdue_days,
            "payment_tx_edit": tx_edit,

            "base_payment_item": base_payment_item,
            "rrn_edit_mode": rrn_edit_mode,
            "rrn_edit_form": rrn_edit_form,
            "user_role": ("운영자" if request.user.is_superuser else ("관리자" if request.user.is_staff else "직원")),
            "can_edit_past_cycles": bool(request.user.is_superuser or request.user.is_staff),
            "is_past_cycle": bool(selected_case and (selected_case.cycle_no != customer.current_cycle)),
            "required_popup": required_popup,
            "rrn_exists": bool((customer.rrn_full or "").strip()),
        },
    )

@login_required
def customers_bulk_trash(request: HttpRequest) -> HttpResponse:
    if request.method != "POST":
        return redirect("/customers/")

    ids = request.POST.getlist("ids")
    if not ids:
        return redirect("/customers/")

    now = timezone.now()
    qs = Customer.objects.filter(id__in=ids, is_deleted=False)
    for c in qs:
        c.is_deleted = True
        c.deleted_at = now
        c.deleted_by = request.user
        c.save(update_fields=["is_deleted", "deleted_at", "deleted_by"])
        from .models import CustomerTrashLog
        CustomerTrashLog.objects.create(customer=c, user=request.user, action="trash")

    return redirect("/customers/")


@login_required
def settings_home(request: HttpRequest) -> HttpResponse:
    return render(request, "customers/settings.html", {})


@login_required
def settings_upload_management(request: HttpRequest) -> HttpResponse:
    if not (request.user.is_superuser or request.user.is_staff):
        return HttpResponseForbidden("Forbidden")
    return render(request, "customers/upload_management.html", {})


@login_required
def settings_upload_management_template(request: HttpRequest, kind: str) -> HttpResponse:
    if not (request.user.is_superuser or request.user.is_staff):
        return HttpResponseForbidden("Forbidden")

    BASIC_HEADERS = ["구분", "고객명", "고객ID", "회차", "연락처", "주소", "보호자_연락처", "보호자_연락처2", "메모", "담당자", "병원명", "방문병원_1차", "방문병원_2차", "방문병원_3차", "제출일", "재검일자", "재검병원", "장애결정일자", "장애도", "제조사", "모델명", "제조번호", "표준코드", "제조일", "리시버", "선착용일", "착용일_구매일", "좌우", "이어몰드_제작일", "공단_인정_금액", "제품_결제일", "본인부담액", "제조사_추가", "모델명_추가", "제조번호_추가", "표준코드_추가", "제조일_추가", "리시버_추가", "선착용일_추가", "착용일_구매일_추가", "좌우_추가", "이어몰드제작일_추가", "자부담금액_추가", "결제일", "결제금액", "결제방식", "메모", "총 결제금액", "검수", "공단_주민센터명", "공단_주민센터_접수일", "제출방법", "보완 내용", "보완 완료", "입금일", "입금액", "후기1차_기간", "제출완료", "입금일", "입금액", "비고", "후기2차_기간", "제출완료", "입금일", "입금액", "비고", "후기3차_기간", "제출완료", "입금일", "입금액", "비고", "후기4차_기간", "제출완료", "입금일", "입금액", "비고"]
    RRN_HEADERS = ["구분", "고객명", "고객ID", "회차", "연락처", "주민등록번호", "주소", "보호자_연락처", "보호자_연락처2", "메모", "담당자", "병원명", "방문병원_1차", "방문병원_2차", "방문병원_3차", "제출일", "재검일자", "재검병원", "장애결정일자", "장애도", "제조사", "모델명", "제조번호", "표준코드", "제조일", "리시버", "선착용일", "착용일_구매일", "좌우", "이어몰드_제작일", "공단_인정_금액", "제품_결제일", "본인부담액", "제조사_추가", "모델명_추가", "제조번호_추가", "표준코드_추가", "제조일_추가", "리시버_추가", "선착용일_추가", "착용일_구매일_추가", "좌우_추가", "이어몰드제작일_추가", "자부담금액_추가", "결제일", "결제금액", "결제방식", "메모", "총 결제금액", "검수", "공단_주민센터명", "공단_주민센터_접수일", "제출방법", "보완 내용", "보완 완료", "입금일", "입금액", "후기1차_기간", "제출완료", "입금일", "입금액", "비고", "후기2차_기간", "제출완료", "입금일", "입금액", "비고", "후기3차_기간", "제출완료", "입금일", "입금액", "비고", "후기4차_기간", "제출완료", "입금일", "입금액", "비고"]

    headers = BASIC_HEADERS if kind == "basic" else (RRN_HEADERS if kind == "rrn" else None)
    if headers is None:
        return HttpResponseBadRequest("Invalid template kind")

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(headers)
    content = "\ufeff" + buf.getvalue()  # UTF-8 BOM
    fn = "customer_upload_basic.csv" if kind == "basic" else "customer_upload_with_rrn.csv"
    resp = HttpResponse(content, content_type="text/csv; charset=utf-8")
    resp["Content-Disposition"] = f'attachment; filename="{fn}"'
    return resp


@require_POST
@login_required
def api_customer_csv_upload(request: HttpRequest) -> JsonResponse:
    if not (request.user.is_superuser or request.user.is_staff):
        return JsonResponse({"ok": False, "error": "Forbidden"}, status=403)

    BASIC_HEADERS = ["구분", "고객명", "고객ID", "회차", "연락처", "주소", "보호자_연락처", "보호자_연락처2", "메모", "담당자", "병원명", "방문병원_1차", "방문병원_2차", "방문병원_3차", "제출일", "재검일자", "재검병원", "장애결정일자", "장애도", "제조사", "모델명", "제조번호", "표준코드", "제조일", "리시버", "선착용일", "착용일_구매일", "좌우", "이어몰드_제작일", "공단_인정_금액", "제품_결제일", "본인부담액", "제조사_추가", "모델명_추가", "제조번호_추가", "표준코드_추가", "제조일_추가", "리시버_추가", "선착용일_추가", "착용일_구매일_추가", "좌우_추가", "이어몰드제작일_추가", "자부담금액_추가", "결제일", "결제금액", "결제방식", "메모", "총 결제금액", "검수", "공단_주민센터명", "공단_주민센터_접수일", "제출방법", "보완 내용", "보완 완료", "입금일", "입금액", "후기1차_기간", "제출완료", "입금일", "입금액", "비고", "후기2차_기간", "제출완료", "입금일", "입금액", "비고", "후기3차_기간", "제출완료", "입금일", "입금액", "비고", "후기4차_기간", "제출완료", "입금일", "입금액", "비고"]
    RRN_HEADERS = ["구분", "고객명", "고객ID", "회차", "연락처", "주민등록번호", "주소", "보호자_연락처", "보호자_연락처2", "메모", "담당자", "병원명", "방문병원_1차", "방문병원_2차", "방문병원_3차", "제출일", "재검일자", "재검병원", "장애결정일자", "장애도", "제조사", "모델명", "제조번호", "표준코드", "제조일", "리시버", "선착용일", "착용일_구매일", "좌우", "이어몰드_제작일", "공단_인정_금액", "제품_결제일", "본인부담액", "제조사_추가", "모델명_추가", "제조번호_추가", "표준코드_추가", "제조일_추가", "리시버_추가", "선착용일_추가", "착용일_구매일_추가", "좌우_추가", "이어몰드제작일_추가", "자부담금액_추가", "결제일", "결제금액", "결제방식", "메모", "총 결제금액", "검수", "공단_주민센터명", "공단_주민센터_접수일", "제출방법", "보완 내용", "보완 완료", "입금일", "입금액", "후기1차_기간", "제출완료", "입금일", "입금액", "비고", "후기2차_기간", "제출완료", "입금일", "입금액", "비고", "후기3차_기간", "제출완료", "입금일", "입금액", "비고", "후기4차_기간", "제출완료", "입금일", "입금액", "비고"]

    def _is_english_filename(fn: str) -> bool:
        return bool(re.match(r"^[A-Za-z0-9._-]+$", fn or ""))

    def _norm(s: str | None) -> str:
        return (s or "").strip()

    def _parse_int(s: str | None) -> int | None:
        t = _norm(s)
        if not t:
            return None
        t = t.replace(",", "")
        try:
            return int(float(t))
        except Exception:
            return None

    def _parse_customer_id(s: str | None) -> int | None:
        t = _norm(s)
        if not t:
            return None
        t = t.lstrip("#")
        try:
            return int(t)
        except Exception:
            # 엑셀/CSV에서 숫자형으로 저장되면 "1234.0" 또는 "1.234E+03" 형태가 될 수 있음
            # -> 기존 고객ID 중복 판정이 누락되지 않도록 float 파싱을 허용
            try:
                return int(float(t.replace(",", "")))
            except Exception:
                return None

    def _parse_date(s: str | None) -> datetime.date | None:
        t = _norm(s)
        if not t:
            return None
        t = t.replace(".", "-").replace("/", "-")
        try:
            return datetime.date.fromisoformat(t)
        except Exception:
            return None

    def _parse_bool(s: str | None) -> bool | None:
        t = _norm(s)
        if t == "":
            return None
        tt = t.upper()
        if tt in ("Y", "YES", "TRUE", "1", "완료"):
            return True
        if tt in ("N", "NO", "FALSE", "0", "미완료"):
            return False
        return None

    def _parse_period(s: str | None) -> tuple[datetime.date | None, datetime.date | None, str | None]:
        t = _norm(s)
        if not t:
            return (None, None, None)
        t = t.replace(" ", "")
        if "~" in t:
            a, b = t.split("~", 1)
        elif "-" in t:
            parts = t.split("-")
            if len(parts) >= 6:
                a = "-".join(parts[:3])
                b = "-".join(parts[3:6])
            else:
                return (None, None, "기간 포맷 오류")
        else:
            return (None, None, "기간 포맷 오류")
        sd = _parse_date(a)
        ed = _parse_date(b)
        if not sd or not ed:
            return (None, None, "기간 날짜 파싱 실패")
        if sd > ed:
            sd, ed = ed, sd
        return (sd, ed, None)

    confirm_dup = _norm(request.POST.get("confirm_duplicate")) == "1"
    token = _norm(request.POST.get("token"))

    file_bytes: bytes | None = None
    filename = ""

    if confirm_dup and token:
        sess = request.session.get("msh_customer_upload_pending") or {}
        if sess.get("token") != token or not sess.get("b64"):
            return JsonResponse({"ok": False, "error": "업로드 토큰이 만료되었습니다. 다시 업로드해 주세요."}, status=400)
        filename = sess.get("filename") or "customer_upload.csv"
        try:
            file_bytes = base64.b64decode(sess.get("b64"))
        except Exception:
            return JsonResponse({"ok": False, "error": "업로드 데이터 복원 실패"}, status=400)
    else:
        f = request.FILES.get("csv_file")
        if not f:
            return JsonResponse({"ok": False, "error": "CSV 파일이 없습니다."}, status=400)
        filename = getattr(f, "name", "") or ""
        if not filename.lower().endswith(".csv"):
            return JsonResponse({"ok": False, "error": "CSV 파일만 업로드 가능합니다."}, status=400)
        if not _is_english_filename(filename):
            return JsonResponse({"ok": False, "error": "CSV 파일명은 영문/숫자/._- 만 가능합니다."}, status=400)
        file_bytes = f.read()

    assert file_bytes is not None

    try:
        text = file_bytes.decode("utf-8-sig")
    except Exception:
        return JsonResponse({"ok": False, "error": "CSV 인코딩은 UTF-8(권장: UTF-8 BOM)만 지원합니다."}, status=400)

    # delimiter sniff (comma "CSV" or tab-delimited)
    sample = text.splitlines()[0] if text else ""
    try:
        if "	" in sample and "," not in sample:
            dialect = csv.excel_tab
        else:
            dialect = csv.Sniffer().sniff(sample, delimiters=",	")
    except Exception:
        dialect = csv.excel

    reader = csv.reader(io.StringIO(text), dialect)
    rows = list(reader)
    if not rows:
        return JsonResponse({"ok": False, "error": "빈 CSV 입니다."}, status=400)

    header = [ (c or "").strip() for c in (rows[0] or []) ]
    expected = RRN_HEADERS if (len(header) == len(RRN_HEADERS) and header == RRN_HEADERS) else (BASIC_HEADERS if (len(header) == len(BASIC_HEADERS) and header == BASIC_HEADERS) else None)
    if expected is None:
        return JsonResponse({"ok": False, "error": "CSV 헤더(컬럼명/순서)가 양식과 일치하지 않습니다."}, status=400)
    has_rrn = (expected is RRN_HEADERS)

    def _row_to_values(r: list[str]) -> list[str]:
        if len(r) < len(expected):
            r = r + [""] * (len(expected) - len(r))
        return r[: len(expected)]

    # data_rows는 "원본 CSV 상의 실제 행 번호(엑셀 행 번호 기준)"를 보존해야 합니다.
    # (중간 빈 줄을 건너뛰면서 행 번호가 1씩 밀리는 혼동을 방지)
    data_rows: list[tuple[int, list[str]]] = []
    for orig_row_no, r in enumerate(rows[1:], start=2):
        rr = _row_to_values(r)
        if all(_norm(x) == "" for x in rr):
            continue
        data_rows.append((orig_row_no, rr))

    if not data_rows:
        return JsonResponse({"ok": False, "error": "업로드 데이터가 없습니다."}, status=400)

    # ⚠️ 중복(이미 등록된 고객)은 "성공"이 아니라 "실패"로 집계합니다.
    # - 사용자 요구사항: 이미 등록된 고객은 반영하지 않으며(현상 유지), 실패 총 건수에 포함하고
    #   실패 사유에 "이미 등록 된 고객 N명"으로 표시합니다.
    #
    # 중복 판정 기준(안전/최소):
    # 1) 고객ID가 존재하고 DB에 이미 있으면 중복
    # 2) (주민번호 양식의 경우) 주민등록번호가 존재하고 DB에 이미 있으면 중복
    # 3) (고객ID/주민번호가 비어있는 경우 대비) 고객명+연락처 조합이 DB에 이미 있으면 중복
    existing_cids: set[int] = set()
    existing_rrns: set[str] = set()
    existing_name_phones: set[tuple[str, str]] = set()
    existing_names: set[str] = set()

    idx_cid = expected.index("고객ID")
    idx_name = expected.index("고객명")
    idx_phone = expected.index("연락처")
    idx_rrn = expected.index("주민등록번호") if has_rrn else -1
    idx_addr = expected.index("주소") if "주소" in expected else -1

    cids_in_file: set[int] = set()
    rrns_in_file: set[str] = set()
    name_phones_in_file: set[tuple[str, str]] = set()
    name_addrs_in_file: set[tuple[str, str]] = set()

    def _norm_rrn(s: str | None) -> str:
        t = _norm(s)
        if not t:
            return ""
        # 숫자만 추출하여 비교(하이픈/공백 혼재 방지)
        digits = "".join([c for c in t if c.isdigit()])
        return digits

    def _norm_phone(s: str | None) -> str:
        """연락처 정규화

        - 하이픈/공백 제거
        - 엑셀에서 숫자/지수표기로 들어오는 경우(예: 1.023456789E10, 10234567890.0)도 정규화
        """
        t = _norm(s)
        if not t:
            return ""
        tt = t.replace(",", "")
        if re.match(r"^\d+(?:\.\d+)?(?:[eE][+-]?\d+)?$", tt):
            try:
                from decimal import Decimal

                d = Decimal(tt)
                if d.is_finite():
                    tt = str(int(d))
            except Exception:
                pass
        digits = "".join([c for c in tt if c.isdigit()])
        return digits

    def _norm_addr(s: str | None) -> str:
        t = _norm(s)
        if not t:
            return ""
        t = re.sub(r"\s+", " ", t).strip().lower()
        return t

    for _orig_row_no, rr in data_rows:
        cid = _parse_customer_id(rr[idx_cid])
        if cid:
            cids_in_file.add(cid)
        if has_rrn and idx_rrn >= 0:
            rrn_digits = _norm_rrn(rr[idx_rrn])
            if rrn_digits:
                rrns_in_file.add(rrn_digits)
        nm = _norm(rr[idx_name])
        ph = _norm(rr[idx_phone])
        ph_digits = _norm_phone(ph)
        if nm and ph_digits:
            name_phones_in_file.add((nm, ph_digits))
        if nm and (not ph_digits) and idx_addr >= 0:
            ad = _norm_addr(rr[idx_addr])
            if ad:
                name_addrs_in_file.add((nm, ad))

    if cids_in_file:
        existing_cids = set(Customer.objects.filter(id__in=list(cids_in_file)).values_list("id", flat=True))
    if rrns_in_file:
        # DB에는 다양한 포맷(하이픈 포함)이 있을 수 있어 digits 기준으로 후처리 비교
        qs = Customer.objects.exclude(rrn_full="").values_list("rrn_full", flat=True)
        for r in qs:
            d = _norm_rrn(r)
            if d and d in rrns_in_file:
                existing_rrns.add(d)
    if name_phones_in_file:
        names = sorted({n for n, _p in name_phones_in_file})
        phone_digits_set = {p for _n, p in name_phones_in_file}
        # DB에는 하이픈/공백 등 다양한 포맷이 섞여 있을 수 있으므로,
        # name으로 1차 필터 후 phone을 digits로 정규화하여 비교합니다.
        for n, p in Customer.objects.filter(name__in=names).values_list("name", "phone"):
            pd = _norm_phone(p)
            if pd and pd in phone_digits_set:
                existing_name_phones.add((_norm(n), pd))

    existing_name_addrs: set[tuple[str, str]] = set()
    if name_addrs_in_file:
        names = sorted({n for n, _a in name_addrs_in_file})
        addr_set = {a for _n, a in name_addrs_in_file}
        for n, a in Customer.objects.filter(name__in=names).values_list("name", "address_summary"):
            ad = _norm_addr(a)
            if ad and ad in addr_set:
                existing_name_addrs.add((_norm(n), ad))

    # 연락처가 비어있는 행(또는 연락처 파싱 불가)은 고객명만으로 중복 방지
    # - 사용자 요구사항: 동일 CSV를 재업로드할 때 신규 생성으로 잡히는 케이스 방지
    existing_names = set(Customer.objects.exclude(name="").values_list("name", flat=True))

    failures: list[dict] = []
    success = 0
    duplicate_count = 0
    # 실패행 다운로드(CSV)에는 "진짜 실패(유효성/파싱 등 에러)"만 포함합니다.
    # 중복(이미 등록된 고객)은 실패 건수에는 포함하되, 다운로드 대상에서는 제외합니다.
    failed_rows: list[list[str]] = []

    def _set_if(obj, field: str, val):
        if val is None:
            return
        if isinstance(val, str) and val.strip() == "":
            return
        setattr(obj, field, val)

    fu_starts = {n: expected.index(f"후기{n}차_기간") for n in (1, 2, 3, 4)}

    for row_no, rr in data_rows:
        def v(col: str) -> str:
            return rr[expected.index(col)]

        # ✅ 실패 기준: 고객명 누락 ONLY
        # - 사용자 요구사항: 나머지 누락/포맷 문제는 실패로 잡지 않습니다.
        name = _norm(v("고객명"))
        if not name:
            failures.append({"row": row_no, "reason": "고객명 누락"})
            failed_rows.append(rr)
            continue

        try:
            track = _norm(v("구분"))
            cid = _parse_customer_id(v("고객ID"))
            cycle_no = _parse_int(v("회차")) or 1
            if cycle_no < 1:
                cycle_no = 1

            # 이미 등록된 고객은 반영하지 않고 '실패'로 집계
            # - 고객ID 기반
            # - 주민번호(숫자만) 기반
            # - 고객명+연락처 기반
            is_dup = False
            if cid and cid in existing_cids:
                is_dup = True
            if (not is_dup) and has_rrn:
                rrn_digits = _norm_rrn(v("주민등록번호"))
                if rrn_digits and rrn_digits in existing_rrns:
                    is_dup = True
            if (not is_dup):
                ph_digits = _norm_phone(v("연락처"))
                if name and ph_digits:
                    if (name, ph_digits) in existing_name_phones:
                        is_dup = True
                else:
                    if idx_addr >= 0:
                        ad = _norm_addr(v("주소"))
                        if ad and (name, ad) in existing_name_addrs:
                            is_dup = True
                    if (not is_dup) and (name in existing_names):
                        is_dup = True

            if is_dup:
                duplicate_count += 1
                # 행별 실패 목록에는 쌓지 않습니다(혼동 방지). 요약 사유로만 제공.
                continue

            # 고객ID가 주어졌지만 DB에 없으면: 신규로 생성(서버 자동 ID)
            customer = None
            if cid:
                customer = Customer.objects.filter(id=cid).first()
            if not customer:
                customer = Customer(name=name)

            # 구분 값이 유효하지 않으면 무시(실패 처리하지 않음)
            if track and (track in dict(Customer.TRACK_CHOICES)):
                customer.track = track

            _set_if(customer, "name", name)
            _set_if(customer, "phone", _norm(v("연락처")))
            _set_if(customer, "address_summary", _norm(v("주소")))
            _set_if(customer, "guardian_phone", _norm(v("보호자_연락처")))
            _set_if(customer, "guardian_phone_2", _norm(v("보호자_연락처2")))
            _set_if(customer, "memo", _norm(v("메모")))
            _set_if(customer, "담당자", _norm(v("담당자")))
            if has_rrn:
                _set_if(customer, "rrn_full", _norm(v("주민등록번호")))

            _set_if(customer, "exam_hospital_name", _norm(v("병원명")))
            _set_if(customer, "exam_visit_1_date", _parse_date(v("방문병원_1차")))
            _set_if(customer, "exam_visit_2_date", _parse_date(v("방문병원_2차")))
            _set_if(customer, "exam_visit_3_date", _parse_date(v("방문병원_3차")))
            _set_if(customer, "exam_submit_date", _parse_date(v("제출일")))
            _set_if(customer, "exam_retest_date", _parse_date(v("재검일자")))
            _set_if(customer, "exam_retest_hospital", _norm(v("재검병원")))
            _set_if(customer, "exam_disability_decision_date", _parse_date(v("장애결정일자")))
            dl = _norm(v("장애도"))
            if dl and dl in dict(Customer.DISABILITY_LEVEL_CHOICES):
                customer.exam_disability_level = dl

            customer.save()

            # 신규 생성된 고객도 즉시 중복 집계 기준에 포함(동일 파일 반복 업로드 시 신규 생성 방지)
            existing_names.add(_norm(customer.name))
            if has_rrn:
                rrn_digits = _norm_rrn(getattr(customer, "rrn_full", ""))
                if rrn_digits:
                    existing_rrns.add(rrn_digits)
            phd = _norm_phone(getattr(customer, "phone", ""))
            if phd:
                existing_name_phones.add((_norm(customer.name), phd))
            if idx_addr >= 0:
                ad = _norm_addr(getattr(customer, "address_summary", ""))
                if ad:
                    existing_name_addrs.add((_norm(customer.name), ad))

            case, _ = CustomerCase.objects.get_or_create(customer=customer, cycle_no=int(cycle_no))

            _set_if(case, "manufacturer", _norm(v("제조사")))
            _set_if(case, "model_name", _norm(v("모델명")))
            _set_if(case, "serial_number", _norm(v("제조번호")))
            _set_if(case, "standard_code", _norm(v("표준코드")))
            _set_if(case, "manufacture_date", _parse_date(v("제조일")))
            _set_if(case, "receiver", _norm(v("리시버")))
            _set_if(case, "pre_fit_date", _parse_date(v("선착용일")))
            _set_if(case, "purchase_date", _parse_date(v("착용일_구매일")))
            # 제품_결제일이 있으면 최우선으로 착용일_구매일을 덮어씀
            _ppd = _parse_date(v("제품_결제일"))
            if _ppd:
                case.purchase_date = _ppd
            side = _norm(v("좌우"))
            if side in ("좌", "우", ""):
                _set_if(case, "side", side)
            _set_if(case, "earmold_made_date", _parse_date(v("이어몰드_제작일")))
            _set_if(case, "nhis_amount", _parse_int(v("공단_인정_금액")))
            _set_if(case, "copay_amount", _parse_int(v("본인부담액")))

            add_any = False
            add_map = [
                ("제조사_추가", "manufacturer_add", "str"),
                ("모델명_추가", "model_name_add", "str"),
                ("제조번호_추가", "serial_number_add", "str"),
                ("표준코드_추가", "standard_code_add", "str"),
                ("제조일_추가", "manufacture_date_add", "date"),
                ("리시버_추가", "receiver_add", "str"),
                ("선착용일_추가", "pre_fit_date_add", "date"),
                ("착용일_구매일_추가", "purchase_date_add", "date"),
                ("이어몰드제작일_추가", "earmold_made_date_add", "date"),
                ("자부담금액_추가", "self_pay_amount_add", "int"),
            ]
            for col, field, typ in add_map:
                raw = v(col)
                if _norm(raw) != "":
                    add_any = True
                if typ == "date":
                    _set_if(case, field, _parse_date(raw))
                elif typ == "int":
                    _set_if(case, field, _parse_int(raw))
                else:
                    _set_if(case, field, _norm(raw))
            side_add = _norm(v("좌우_추가"))
            if side_add in ("좌", "우", ""):
                if side_add != "":
                    add_any = True
                _set_if(case, "side_add", side_add)
            case.has_sub = bool(case.has_sub or add_any)

            _set_if(case, "nhis_inspection_date", _parse_date(v("검수")))
            _set_if(case, "nhis_center_name", _norm(v("공단_주민센터명")))
            _set_if(case, "nhis_submit_date", _parse_date(v("공단_주민센터_접수일")))
            sm = _norm(v("제출방법"))
            if sm in ("FAX", "방문제출", ""):
                _set_if(case, "nhis_submit_method", sm)
            _set_if(case, "nhis_supplement_content", _norm(v("보완 내용")))
            bd = _parse_bool(v("보완 완료"))
            if bd is not None:
                case.nhis_supplement_done = bd
                if bd and not case.nhis_supplement_done_at:
                    case.nhis_supplement_done_at = timezone.now()
            _set_if(case, "nhis_deposit_date", _parse_date(v("입금일")))
            _set_if(case, "nhis_deposit_amount", _parse_int(v("입금액")))

            for n in (1, 2, 3, 4):
                start_idx = fu_starts[n]
                period = rr[start_idx]
                subm = rr[start_idx + 1]
                depd = rr[start_idx + 2]
                depa = rr[start_idx + 3]
                note = rr[start_idx + 4]

                sd, ed, perr = _parse_period(period)
                # 기간 포맷 오류는 무시(실패 처리하지 않음)
                if perr:
                    sd, ed = None, None
                if sd or ed:
                    _set_if(case, f"fu{n}_start_override", sd)
                    _set_if(case, f"fu{n}_end_override", ed)
                sb = _parse_bool(subm)
                if sb is not None:
                    setattr(case, f"fu{n}_submitted", sb)
                    if sb and not getattr(case, f"fu{n}_submitted_at"):
                        setattr(case, f"fu{n}_submitted_at", timezone.now())
                _set_if(case, f"fu{n}_deposit_date", _parse_date(depd))
                _set_if(case, f"fu{n}_deposit_amount", _parse_int(depa))
                _set_if(case, f"fu{n}_note", _norm(note))

            case.save()

            # ✅ 결제일 기준 매출상세용 자동 TX 생성 (제품_결제일 > 착용일_구매일)
            _auto_paid_at = _parse_date(v("제품_결제일")) or _parse_date(v("착용일_구매일"))
            _auto_amt = (_parse_int(v("본인부담액")) or 0) + (_parse_int(v("자부담금액_추가")) or 0)
            if _auto_paid_at and _auto_amt > 0:
                if not PaymentTransaction.objects.filter(case=case, paid_at=_auto_paid_at, amount=_auto_amt, memo="upload:auto:product_payment").exists():
                    PaymentTransaction.objects.create(case=case, paid_at=_auto_paid_at, amount=_auto_amt, method="", memo="upload:auto:product_payment")

            paid_at = _parse_date(v("결제일"))
            amt = _parse_int(v("결제금액"))
            method = _norm(v("결제방식"))
            memo = _norm(v("메모"))
            if paid_at and amt is not None:
                if not PaymentTransaction.objects.filter(case=case, paid_at=paid_at, amount=amt, method=method, memo=memo).exists():
                    PaymentTransaction.objects.create(case=case, paid_at=paid_at, amount=amt, method=method, memo=memo)

            success += 1
        except Exception:
            # 사용자 요구사항: 고객명 누락 외 실패로 잡지 않음
            # (예상치 못한 포맷/파싱 문제는 해당 행을 조용히 스킵)
            continue

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(expected)
    for r in failed_rows:
        w.writerow(r)
    failed_csv = "\ufeff" + buf.getvalue()

    if request.session.get("msh_customer_upload_pending"):
        request.session.pop("msh_customer_upload_pending", None)
        request.session.modified = True

    return JsonResponse({
        "ok": True,
        "status": "done",
        "success_count": success,
        # ✅ 실패 건수는 '고객명 누락'만
        "fail_count": int(len(failures)),
        # ✅ 중복은 별도 표기(반영 안 됨)
        "duplicate_count": int(duplicate_count),
        "failures": failures[:500],
        "failed_csv": failed_csv,
    })


@login_required
def sales_analysis(request: HttpRequest) -> HttpResponse:
    """설정 > 매출분석: 결제일 기준 / 입금일 기준 집계"""
    # 매출분석/다운로드는 관리자만
    if not (request.user.is_staff or request.user.is_superuser):
        return HttpResponseForbidden("Forbidden")
    today = timezone.localdate()

    def _parse_date(key: str, default: datetime.date) -> datetime.date:
        s = (request.GET.get(key) or "").strip()
        if not s:
            return default
        try:
            return datetime.date.fromisoformat(s)
        except Exception:
            return default

    end = _parse_date("end", today)
    start = _parse_date("start", end - datetime.timedelta(days=30))
    if start > end:
        start, end = end, start

    gran = (request.GET.get("gran") or "day").strip()
    if gran not in ("day", "week", "month"):
        gran = "day"

    view_mode = (request.GET.get("view") or "data").strip()
    if view_mode not in ("data", "detail"):
        view_mode = "data"

    # ✅ 매출상세는 일(day) 기준만 허용(요구사항)
    if view_mode == "detail":
        gran = "day"

    # ✅ 조회 버튼을 누르기 전에는 데이터 로딩을 하지 않습니다.
    searched = (request.GET.get("searched") or "").strip() == "1"

    # 매출상세(view=detail)은 매출자료와 동일한 폼/집계 단위를 사용합니다.

    start_month = (request.GET.get("start_month") or "").strip()
    end_month = (request.GET.get("end_month") or "").strip()

    # 집계 단위가 '월'이면 월(YYYY-MM) 기간(시작월~종료월)만 선택하도록: 해당 월 범위로 기간 고정
    if view_mode != "detail" and gran == "month":
        def _ym_to_first(ym: str) -> datetime.date | None:
            try:
                y_str, m_str = ym.split("-")
                y = int(y_str)
                m = int(m_str)
                if m < 1 or m > 12:
                    return None
                return datetime.date(y, m, 1)
            except Exception:
                return None

        def _ym_to_last(ym: str) -> datetime.date | None:
            try:
                y_str, m_str = ym.split("-")
                y = int(y_str)
                m = int(m_str)
                if m < 1 or m > 12:
                    return None
                last_day = calendar.monthrange(y, m)[1]
                return datetime.date(y, m, last_day)
            except Exception:
                return None

        # 기본값: 종료일 기준 월로 시작월/종료월 세팅
        if not end_month:
            end_month = f"{end.year:04d}-{end.month:02d}"
        if not start_month:
            start_month = end_month

        s_first = _ym_to_first(start_month)
        e_last = _ym_to_last(end_month)

        if not s_first or not e_last:
            # 값이 없거나 잘못된 경우: 종료일 기준 월로 강제
            start_month = f"{end.year:04d}-{end.month:02d}"
            end_month = start_month
            s_first = _ym_to_first(start_month)
            e_last = _ym_to_last(end_month)

        # 역순이면 swap
        if s_first and e_last and s_first > e_last:
            start_month, end_month = end_month, start_month
            s_first = _ym_to_first(start_month)
            e_last = _ym_to_last(end_month)

        if s_first and e_last:
            start, end = s_first, e_last
            start_month = f"{start.year:04d}-{start.month:02d}"
            end_month = f"{end.year:04d}-{end.month:02d}"

    
    # 집계 단위가 '주'이면: 선택한 날짜 범위를 해당 주(일~토) 단위로 자동 정렬(서버에서도 강제)
    if view_mode != "detail" and gran == "week":
        # 일~토: 일요일이 주 시작(0), 토요일이 주 끝(6)
        def _to_week_start(d: datetime.date) -> datetime.date:
            # Python weekday: Mon=0..Sun=6 -> Sunday start offset
            return d - datetime.timedelta(days=(d.weekday() + 1) % 7)

        def _to_week_end(d: datetime.date) -> datetime.date:
            sunday_index = (d.weekday() + 1) % 7  # Sun=0..Sat=6
            return d + datetime.timedelta(days=(6 - sunday_index))

        start = _to_week_start(start)
        end = _to_week_end(end)
        if start > end:
            start, end = end, start


    def _periods(start_d: datetime.date, end_d: datetime.date, g: str):
        periods: list[tuple[str, datetime.date, datetime.date]] = []
        if g == "day":
            cur = start_d
            while cur <= end_d:
                periods.append((cur.strftime("%Y-%m-%d"), cur, cur))
                cur = cur + datetime.timedelta(days=1)
            return periods
        if g == "week":
            # 일~토: 일요일이 주 시작
            cur = start_d - datetime.timedelta(days=(start_d.weekday() + 1) % 7)
            while cur <= end_d:
                w_start = cur
                w_end = cur + datetime.timedelta(days=6)
                ps = max(w_start, start_d)
                pe = min(w_end, end_d)
                label = f"{ps.strftime('%Y-%m-%d')}~{pe.strftime('%Y-%m-%d')}"
                periods.append((label, ps, pe))
                cur = cur + datetime.timedelta(days=7)
            return [p for p in periods if p[1] <= p[2]]

        # month
        cur = datetime.date(start_d.year, start_d.month, 1)
        while cur <= end_d:
            last_day = calendar.monthrange(cur.year, cur.month)[1]
            m_start = cur
            m_end = datetime.date(cur.year, cur.month, last_day)
            label = cur.strftime("%Y.%m")
            periods.append((label, max(m_start, start_d), min(m_end, end_d)))
            if cur.month == 12:
                cur = datetime.date(cur.year + 1, 1, 1)
            else:
                cur = datetime.date(cur.year, cur.month + 1, 1)
        return [p for p in periods if p[1] <= p[2]]

    periods = _periods(start, end, gran)

    # 조회 전에는 무거운 ORM 집계를 아예 수행하지 않습니다.
    if not searched:
        return render(
            request,
            "customers/sales_analysis.html",
            {
                "start": start,
                "end": end,
                "gran": gran,
                "start_month": start_month,
                "end_month": end_month,
                "view_mode": view_mode,
                "searched": searched,
                "pay_rows": [],
                "dep_rows": [],
                "detail_pay_rows": [],
                "detail_dep_rows": [],
            },
        )

    cases = list(CustomerCase.objects.select_related("customer").all())

    FOLLOWUP_AMOUNT = 50000  # 후기적합 차수별 고정 금액

    def _followup_sum_by_submitted(ps: datetime.date, pe: datetime.date) -> int:
        total = 0
        for c in cases:
            for n in (1, 2, 3, 4):
                dt = getattr(c, f"fu{n}_submitted_at", None)
                if not dt:
                    continue
                try:
                    d = dt.date()
                except Exception:
                    continue
                if ps <= d <= pe:
                    total += FOLLOWUP_AMOUNT
        return total

    def _followup_sum_by_deposit(ps: datetime.date, pe: datetime.date) -> int:
        total = 0
        for c in cases:
            for n in (1, 2, 3, 4):
                d = getattr(c, f"fu{n}_deposit_date", None)
                if not d:
                    continue
                if ps <= d <= pe:
                    total += FOLLOWUP_AMOUNT
        return total

    def _case_amounts(c: CustomerCase):
        nhis = int(getattr(c, "nhis_amount", 0) or 0)
        copay = int(getattr(c, "copay_amount", 0) or 0)
        add = int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0
        total = nhis + copay + add
        return nhis, copay, add, total

    def _case_qty(c: CustomerCase) -> int:
        return 2 if getattr(c, "has_sub", False) else 1

    def _tax_label(taxes: set[str]) -> str:
        taxes = {t for t in taxes if t}
        if not taxes:
            return "-"
        if len(taxes) >= 2:
            return "혼합"
        return next(iter(taxes))

    # 케이스별 첫 결제일(없으면 구매일) + 첫 트랜잭션(과세구분 판단)
    case_first_date: dict[int, datetime.date | None] = {}
    case_first_tx: dict[int, PaymentTransaction | None] = {}
    for c in cases:
        txs = list(PaymentTransaction.objects.filter(case=c).order_by("paid_at", "id"))
        if txs:
            case_first_date[c.id] = txs[0].paid_at
            case_first_tx[c.id] = txs[0]
        elif c.purchase_date:
            case_first_date[c.id] = c.purchase_date
            case_first_tx[c.id] = None
        else:
            case_first_date[c.id] = None
            case_first_tx[c.id] = None

    # 환불(수납 이력에서 음수 금액)
    def _refund_sum(ps: datetime.date, pe: datetime.date) -> int:
        refund_txs = PaymentTransaction.objects.filter(paid_at__gte=ps, paid_at__lte=pe, amount__lt=0)
        return sum(abs(int(t.amount or 0)) for t in refund_txs)

    # ✅ A/S 매출/환불 합계
    # 요구사항:
    # - 유상(is_paid=True) & 결제완료(결제일=paid_at 존재)만 매출로 포함
    # - 취소(CANCELED) 제외
    # - 입금일 기준도 "입금일"을 보지 않고 결제완료 즉시 포함 => paid_at 기준으로 동일 집계
    from django.db.models import Q

    def _as_paid_net_sum_by_pay(ps: datetime.date, pe: datetime.date):
        qs = (
            AfterService.objects.filter(
                customer__is_deleted=False,
                is_paid=True,
                amount__gt=0,
                paid_at__isnull=False,
                paid_at__gte=ps,
                paid_at__lte=pe,
            )
            .exclude(status="CANCELED")
        )
        # net = amount - refund_amount (0 미만은 0 처리)
        total = 0
        taxes: set[str] = set()
        for amount, refund_amount, tax_type in qs.values_list("amount", "refund_amount", "tax_type"):
            amt = int(amount or 0)
            ref = int(refund_amount or 0)
            net = amt - ref
            if net < 0:
                net = 0
            total += net
            if tax_type:
                taxes.add(tax_type)
        return total, taxes

    def _as_paid_net_sum_by_dep(ps: datetime.date, pe: datetime.date):
        # 입금일 기준도 결제완료 즉시 포함해야 하므로 paid_at 기준으로 동일 집계
        return _as_paid_net_sum_by_pay(ps, pe)

    def _as_refund_sum(ps: datetime.date, pe: datetime.date):
        qs = (
            AfterService.objects.filter(
                customer__is_deleted=False,
                refund_amount__gt=0,
                refund_at__isnull=False,
                refund_at__gte=ps,
                refund_at__lte=pe,
            )
            .exclude(status="CANCELED")
        )
        total = sum(int(a.refund_amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes


    # ✅ A/S 매출/환불 합계 (부가세 매출자료 다운로드용)
    # - 결제일 기준: paid_at
    # - 입금일 기준: deposited_at이 있으면 deposited_at, 없으면 paid_at
    def _as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            paid_at__gte=ps,
            paid_at__lte=pe,
        )
        total = sum(int(a.amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    def _as_paid_sum_by_dep(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="PAID")
        total = 0
        taxes: set[str] = set()
        for a in qs:
            d = a.deposited_at or a.paid_at
            if not d or d < ps or d > pe:
                continue
            total += int(a.amount or 0)
            taxes.add(a.tax_type or "")
        return total, taxes

    def _as_refund_sum(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, refund_at__gte=ps, refund_at__lte=pe)
        total = sum(int(a.refund_amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    # ✅ A/S 매출/환불 합계 (부가세 매출자료 다운로드용)
    # - 결제일 기준: paid_at
    # - 입금일 기준: deposited_at이 있으면 deposited_at, 없으면 paid_at
    def _as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            paid_at__gte=ps,
            paid_at__lte=pe,
        )
        total = sum(int(a.amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    def _as_paid_sum_by_dep(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="PAID")
        total = 0
        taxes = set()
        for a in qs:
            d = a.deposited_at or a.paid_at
            if not d or d < ps or d > pe:
                continue
            total += int(a.amount or 0)
            taxes.add(a.tax_type or "")
        return total, taxes

    def _as_refund_sum(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, refund_at__gte=ps, refund_at__lte=pe)
        total = sum(int(a.refund_amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    def _as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="PAID", paid_at__gte=ps, paid_at__lte=pe)
        total = sum(int(a.amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    def _as_paid_sum_by_dep(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="PAID")
        total = 0
        taxes = set()
        for a in qs:
            d = a.deposited_at or a.paid_at
            if not d or d < ps or d > pe:
                continue
            total += int(a.amount or 0)
            taxes.add(a.tax_type or "")
        return total, taxes

    def _as_refund_sum(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, refund_at__gte=ps, refund_at__lte=pe)
        total = sum(int(a.refund_amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    # A/S 매출/환불 (결제일/입금일 기준)
    def _as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            paid_at__gte=ps,
            paid_at__lte=pe,
        )
        total = sum(int(a.amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    def _as_paid_sum_by_dep(ps: datetime.date, pe: datetime.date):
        # 입금일이 있으면 입금일, 없으면 결제일을 사용
        qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="PAID")
        items = []
        taxes = set()
        total = 0
        for a in qs:
            d = a.deposited_at or a.paid_at
            if not d or d < ps or d > pe:
                continue
            total += int(a.amount or 0)
            taxes.add(a.tax_type or "")
            items.append(a)
        return total, taxes

    def _as_refund_sum(ps: datetime.date, pe: datetime.date):
        qs = AfterService.objects.filter(customer__is_deleted=False, refund_at__gte=ps, refund_at__lte=pe)
        total = sum(int(a.refund_amount or 0) for a in qs)
        taxes = set((a.tax_type or "") for a in qs)
        return total, taxes

    # 결제일 기준(케이스 1회 집계)
    pay_rows: list[dict] = []
    pay_total = {"qty": 0, "total": 0, "real": 0, "nhis": 0, "copay": 0, "self_add": 0, "followup": 0, "as_cost": 0, "refund": 0, "tax_label": "-"}
    pay_tax_flags: set[str] = set()

    for label, ps, pe in periods:
        taxes: set[str] = set()
        row = {"period": label, "qty": 0, "total": 0, "real": 0, "nhis": 0, "copay": 0, "self_add": 0, "followup": 0, "as_cost": 0, "refund": 0, "tax_label": "-"}

        for c in cases:
            d = case_first_date.get(c.id)
            if not d or d < ps or d > pe:
                continue
            nhis, copay, add, total = _case_amounts(c)
            qty = _case_qty(c)
            row["qty"] += qty
            row["total"] += total
            row["nhis"] += nhis
            row["copay"] += copay
            row["self_add"] += add
            tx = case_first_tx.get(c.id)
            tax = (getattr(tx, "tax_type", "") or "") if tx else ""
            if tax:
                taxes.add(tax)
                pay_tax_flags.add(tax)

        # 후기적합(제출완료일 기준): 차수별 50,000원
        fu_sum = _followup_sum_by_submitted(ps, pe)
        row["followup"] += fu_sum
        row["total"] += fu_sum

        # A/S (결제일 기준): 유상 & 수납완료만 매출로 포함
        as_paid, as_tax = _as_paid_net_sum_by_pay(ps, pe)
        as_refund, as_ref_tax = _as_refund_sum(ps, pe)
        row["total"] += as_paid
        row["as_cost"] += as_paid
        row["refund"] += as_refund
        for t in (as_tax | as_ref_tax):
            if t:
                taxes.add(t)
                pay_tax_flags.add(t)

        row["refund"] += _refund_sum(ps, pe)
        row["real"] = row["total"] - row["refund"]
        row["tax_label"] = _tax_label(taxes)
        pay_rows.append(row)

        # totals
        pay_total["qty"] += row["qty"]
        pay_total["total"] += row["total"]
        pay_total["nhis"] += row["nhis"]
        pay_total["copay"] += row["copay"]
        pay_total["self_add"] += row["self_add"]
        pay_total["followup"] += row.get("followup", 0)
        pay_total["as_cost"] += row["as_cost"]
        pay_total["refund"] += row["refund"]

    pay_total["tax_label"] = _tax_label(pay_tax_flags)

    # 입금일 기준(수납 + 공단입금)
    dep_rows: list[dict] = []
    dep_total = {"qty": 0, "total": 0, "real": 0, "nhis_deposit": 0, "copay": 0, "self_add": 0, "followup": 0, "unpaid": 0, "unpaid_cum": 0, "as_cost": 0, "refund": 0, "tax_label": "-"}
    dep_tax_flags: set[str] = set()

    for label, ps, pe in periods:
        taxes: set[str] = set()
        row = {
            "period": label,
            "qty": 0,
            "total": 0,  # 총 매출(입금/수납 유입 총액, 환불 제외)
            "nhis_deposit": 0,
            "copay": 0,
            "self_add": 0,
            "followup": 0,
            "unpaid": 0,
            "unpaid_cum": 0,
            "as_cost": 0,
            "refund": 0,
            "tax_label": "-",
        }

        # 수납 트랜잭션(양수=수납, 음수=환불)
        txs = list(PaymentTransaction.objects.filter(paid_at__gte=ps, paid_at__lte=pe))
        pos = [t for t in txs if int(t.amount or 0) > 0]
        neg = [t for t in txs if int(t.amount or 0) < 0]
        row["total"] += sum(int(t.amount or 0) for t in pos)
        row["refund"] += sum(abs(int(t.amount or 0)) for t in neg)

        for t in txs:
            tax = (t.tax_type or "")
            if tax:
                taxes.add(tax)
                dep_tax_flags.add(tax)

        # 공단 입금(면세로 라벨 판단)
        dep_cases = list(CustomerCase.objects.filter(nhis_deposit_date__gte=ps, nhis_deposit_date__lte=pe))
        dep_sum = sum(int(getattr(c, "nhis_deposit_amount", 0) or 0) for c in dep_cases)
        if dep_sum:
            taxes.add("면세")
            dep_tax_flags.add("면세")
        row["nhis_deposit"] += dep_sum
        row["total"] += dep_sum

        # 후기적합(입금일 기준): 차수별 50,000원
        fu_sum = _followup_sum_by_deposit(ps, pe)
        row["followup"] += fu_sum
        row["total"] += fu_sum

        # A/S (입금일 기준): 유상 & 수납완료만 매출로 포함
        as_paid, as_tax = _as_paid_net_sum_by_dep(ps, pe)
        as_refund, as_ref_tax = _as_refund_sum(ps, pe)
        row["total"] += as_paid
        row["as_cost"] += as_paid
        row["refund"] += as_refund
        for t in (as_tax | as_ref_tax):
            if t:
                taxes.add(t)
                dep_tax_flags.add(t)

        # 보청기 수량: 기간 내 수납/환불/공단입금이 있었던 케이스 기준(1회)
        touched_case_ids = set([t.case_id for t in txs]) | set([c.id for c in dep_cases])
        for c in cases:
            if c.id not in touched_case_ids:
                continue
            row["qty"] += _case_qty(c)
            _, copay, add, _ = _case_amounts(c)
            row["copay"] += copay
            row["self_add"] += add

        # 미수/누적미수(기간 종료일 기준)
        for c in cases:
            if not c.purchase_date:
                continue
            nhis, copay, add, total = _case_amounts(c)
            paid_upto = sum(int(t.amount or 0) for t in PaymentTransaction.objects.filter(case=c, paid_at__lte=pe))
            nhis_upto = 0
            if getattr(c, "nhis_deposit_date", None) and c.nhis_deposit_date <= pe:
                nhis_upto = int(getattr(c, "nhis_deposit_amount", 0) or 0)
            outstanding = total - paid_upto - nhis_upto
            if outstanding <= 0:
                continue
            if c.purchase_date <= pe:
                row["unpaid_cum"] += outstanding
            if ps <= c.purchase_date <= pe:
                row["unpaid"] += outstanding

        # A/S 미수/누적미수: 접수일 기준
        as_unpaid_qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="UNPAID")
        for a in as_unpaid_qs:
            if not a.received_at:
                continue
            if a.received_at <= pe:
                row["unpaid_cum"] += int(a.amount or 0)
            if ps <= a.received_at <= pe:
                row["unpaid"] += int(a.amount or 0)

        row["real"] = row["total"] - row["refund"]
        row["tax_label"] = _tax_label(taxes)
        dep_rows.append(row)

        # totals
        dep_total["qty"] += row["qty"]
        dep_total["total"] += row["total"]
        dep_total["nhis_deposit"] += row["nhis_deposit"]
        dep_total["copay"] += row["copay"]
        dep_total["self_add"] += row["self_add"]
        dep_total["followup"] += row.get("followup", 0)
        dep_total["unpaid"] += row["unpaid"]
        dep_total["unpaid_cum"] += row["unpaid_cum"]
        dep_total["as_cost"] += row["as_cost"]
        dep_total["refund"] += row["refund"]

    pay_total["real"] = pay_total["total"] - pay_total["refund"]
    dep_total["real"] = dep_total["total"] - dep_total["refund"]
    dep_total["tax_label"] = _tax_label(dep_tax_flags)


    # =========================
    # 매출상세(view=detail)
    # - 결제일기준: PaymentTransaction(수납) 단위
    # - 입금일기준: 수납 + 공단입금(케이스 입금일/입금액) 이벤트 단위
    #   ※ 고객 식별(고객ID/고객명)이 가능한 형태로 출력
    # =========================
    detail_pay_rows: list[dict] = []
    detail_dep_rows: list[dict] = []

    if view_mode == "detail":
        def _period_label_for_date(d: datetime.date, g: str) -> str:
            if g == "day":
                return d.strftime("%Y-%m-%d")
            if g == "month":
                return f"{d.year:04d}.{d.month:02d}"
            # week: 일~토
            ws = d - datetime.timedelta(days=(d.weekday() + 1) % 7)
            we = ws + datetime.timedelta(days=6)
            return f"{ws.strftime('%Y-%m-%d')}~{we.strftime('%Y-%m-%d')}"

        def _case_qty_detail(c: CustomerCase) -> int:
            return 2 if getattr(c, "has_sub", False) else 1

        def _case_expected_total(c: CustomerCase) -> int:
            nhis = int(getattr(c, "nhis_amount", 0) or 0)
            copay = int(getattr(c, "copay_amount", 0) or 0)
            add = int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0
            return nhis + copay + add

        # A/S 비용(일자별, 고객별) 맵 (net = amount - refund_amount)
        as_map: dict[tuple[int, datetime.date], int] = {}
        as_qs = (
            AfterService.objects.filter(
                customer__is_deleted=False,
                paid_at__isnull=False,
                paid_at__gte=start,
                paid_at__lte=end,
                is_paid=True,
                amount__gt=0,
            )
            .exclude(status="CANCELED")
            .values_list("customer_id", "paid_at", "amount", "refund_amount")
        )
        for cust_id, paid_at, amount, refund_amount in as_qs:
            net = int(amount or 0) - int(refund_amount or 0)
            if net < 0:
                net = 0
            key = (int(cust_id), paid_at)
            as_map[key] = as_map.get(key, 0) + net

        # 후기적합(차수별 50,000) 이벤트 맵
        fu_sub_map: dict[tuple[int, datetime.date], int] = {}
        fu_dep_map: dict[tuple[int, datetime.date], int] = {}
        for c in cases:
            if not c or not getattr(c, 'customer_id', None):
                continue
            for n in (1, 2, 3, 4):
                dt = getattr(c, f'fu{n}_submitted_at', None)
                if dt:
                    try:
                        d = dt.date()
                    except Exception:
                        d = None
                    if d and start <= d <= end:
                        fu_sub_map[(c.customer_id, d)] = fu_sub_map.get((c.customer_id, d), 0) + FOLLOWUP_AMOUNT
                dd = getattr(c, f'fu{n}_deposit_date', None)
                if dd and start <= dd <= end:
                    fu_dep_map[(c.customer_id, dd)] = fu_dep_map.get((c.customer_id, dd), 0) + FOLLOWUP_AMOUNT

        # 결제일기준 상세 (수납 이력)
        pay_qs = (
            PaymentTransaction.objects.filter(
                # ✅ paid_at이 DateTimeField인 경우 end가 00:00:00으로 비교되어
                # 같은 날짜의 오후 결제가 누락될 수 있습니다.
                # 날짜 단위(일 전체)로 필터링합니다.
                paid_at__gte=start,
                paid_at__lte=end,
                case__customer__is_deleted=False,
            )
            .select_related("case", "case__customer")
            .order_by("case__customer__name", "case__customer__id", "paid_at", "id")
        )

        for t in pay_qs:
            c = t.case
            cust = c.customer
            paid_d = t.paid_at.date() if hasattr(t.paid_at, "date") else t.paid_at
            nhis = int(getattr(c, "nhis_amount", 0) or 0)
            copay = int(getattr(c, "copay_amount", 0) or 0)
            self_add = int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0
            total_case = nhis + copay + self_add
            amt = int(t.amount or 0)
            refund_amt = abs(amt) if amt < 0 else 0
            real_amt = amt if amt > 0 else 0
            detail_pay_rows.append(
                {
                    "period": _period_label_for_date(paid_d, gran),
                    "customer_id": cust.id,
                    "customer_name": cust.name,
                    "qty": _case_qty_detail(c),
                    "total": total_case,
                    "real": real_amt,
                    "tax_label": t.tax_type or "-",
                    "nhis": nhis,
                    "copay": copay,
                    "self_add": self_add,
                    "as_cost": as_map.get((cust.id, paid_d), 0),
                    "refund": refund_amt,
                    "followup": 0,
                    "_sort_date": paid_d,
                    "_kind": "pay",
                }
            )


        # 후기적합(제출완료일 기준) 상세 행 추가
        for c in cases:
            cust = getattr(c, 'customer', None)
            if not cust or getattr(cust, 'is_deleted', False):
                continue
            for n in (1, 2, 3, 4):
                dt = getattr(c, f'fu{n}_submitted_at', None)
                if not dt:
                    continue
                try:
                    d = dt.date()
                except Exception:
                    continue
                if not (start <= d <= end):
                    continue
                detail_pay_rows.append(
                    {
                        'period': _period_label_for_date(d, gran),
                        'customer_id': cust.id,
                        'customer_name': cust.name,
                        'qty': 0,
                        'total': FOLLOWUP_AMOUNT,
                        'real': FOLLOWUP_AMOUNT,
                        'tax_label': '-',
                        'nhis': 0,
                        'copay': 0,
                        'self_add': 0,
                        'as_cost': as_map.get((cust.id, d), 0),
                        'refund': 0,
                        'followup': FOLLOWUP_AMOUNT,
                        '_sort_date': d,
                        '_kind': 'fu',
                    }
                )

        # 정렬(고객명/ID/일자, 결제(수납) -> 후기적합 순)
        kind_order = {'pay': 0, 'fu': 1}
        detail_pay_rows.sort(key=lambda r: (r.get('customer_name',''), int(r.get('customer_id',0) or 0), r.get('_sort_date', start), kind_order.get(r.get('_kind','pay'), 9)))
        for r in detail_pay_rows:
            r.pop('_sort_date', None)
            r.pop('_kind', None)

        # 입금일기준 상세: 수납 + 공단입금 이벤트
        dep_cases = (
            CustomerCase.objects.filter(
                nhis_deposit_date__isnull=False,
                nhis_deposit_date__gte=start,
                nhis_deposit_date__lte=end,
                customer__is_deleted=False,
            )
            .select_related("customer")
            .order_by("customer__name", "customer__id", "nhis_deposit_date", "id")
        )

        # 케이스 캐시
        case_cache: dict[int, CustomerCase] = {c.id: c for c in cases}

        # 케이스별 수납 이벤트(수납일 기준)
        pay_amounts_by_case_date: dict[tuple[int, datetime.date], list[int]] = {}
        for t in pay_qs:
            d = t.paid_at.date() if hasattr(t.paid_at, "date") else t.paid_at
            pay_amounts_by_case_date.setdefault((t.case_id, d), []).append(int(t.amount or 0))

        dep_amount_by_case: dict[int, int] = {}
        for c in dep_cases:
            dep_amount_by_case[c.id] = int(getattr(c, "nhis_deposit_amount", 0) or 0)

        # 미수금 러닝(고객별 누적)
        unpaid_by_case: dict[int, int] = {}
        received_by_case: dict[int, int] = {}
        cases_by_customer: dict[int, list[int]] = {}
        for c in cases:
            if getattr(c.customer, "is_deleted", False):
                continue
            cases_by_customer.setdefault(c.customer_id, []).append(c.id)
            unpaid_by_case[c.id] = _case_expected_total(c)
            received_by_case[c.id] = 0

        # 이벤트 목록 생성 (고객, 일자, 케이스, 종류)
        events: list[tuple[int, datetime.date, int, str]] = []
        for (case_id, d), _amts in pay_amounts_by_case_date.items():
            c = case_cache.get(case_id)
            if not c:
                continue
            events.append((c.customer_id, d, case_id, "pay"))
        for c in dep_cases:
            events.append((c.customer_id, c.nhis_deposit_date, c.id, "nhis"))

        # 후기적합(입금일 기준) 이벤트 추가
        for c in cases:
            if getattr(c.customer, 'is_deleted', False):
                continue
            for n in (1, 2, 3, 4):
                dd = getattr(c, f'fu{n}_deposit_date', None)
                if dd and start <= dd <= end:
                    events.append((c.customer_id, dd, c.id, 'fu'))
        events.sort(key=lambda x: (x[0], x[1], x[2], x[3]))

        for cust_id, d, case_id, kind in events:
            c = case_cache.get(case_id)
            if not c:
                continue
            cust = c.customer
            if not cust or getattr(cust, "is_deleted", False):
                continue

            # ✅ 미수 누적은 고객 단위로 항상 계산 가능해야 합니다.
            # 후기적합(fu) 이벤트에서도 누적값 표시가 필요하므로 기본값을 먼저 정의합니다.
            unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))

            if kind == "pay":
                for amt in pay_amounts_by_case_date.get((case_id, d), []):
                    received_by_case[case_id] = received_by_case.get(case_id, 0) + int(amt)
                    unpaid_by_case[case_id] = max(_case_expected_total(c) - received_by_case[case_id], 0)
                    unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))

                    detail_dep_rows.append(
                        {
                            "period": _period_label_for_date(d, gran),
                            "customer_id": cust.id,
                            "customer_name": cust.name,
                            "qty": _case_qty_detail(c),
                            "total": int(amt) if int(amt) > 0 else 0,
                            "real": int(amt),
                            "nhis_deposit": 0,
                            "copay": int(getattr(c, "copay_amount", 0) or 0),
                            "self_add": int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0,
                            "unpaid": unpaid_by_case[case_id],
                            "unpaid_cum": unpaid_cum,
                            "as_cost": as_map.get((cust.id, d), 0),
                            "followup": 0,
                        }
                    )
            elif kind == 'nhis':
                dep_amt = dep_amount_by_case.get(case_id, 0)
                if dep_amt == 0:
                    continue
                received_by_case[case_id] = received_by_case.get(case_id, 0) + int(dep_amt)
                unpaid_by_case[case_id] = max(_case_expected_total(c) - received_by_case[case_id], 0)
                unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))

                detail_dep_rows.append(
                    {
                        "period": _period_label_for_date(d, gran),
                        "customer_id": cust.id,
                        "customer_name": cust.name,
                        "qty": _case_qty_detail(c),
                        "total": int(dep_amt),
                        "real": int(dep_amt),
                        "nhis_deposit": int(dep_amt),
                        "copay": int(getattr(c, "copay_amount", 0) or 0),
                        "self_add": int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0,
                        "unpaid": unpaid_by_case[case_id],
                        "unpaid_cum": unpaid_cum,
                        "as_cost": as_map.get((cust.id, d), 0),
                        "followup": 0,
                    }
                )

            else:  # kind == 'fu'
                # 후기적합은 보청기 미수/누적미수 산식에 영향을 주지 않음
                detail_dep_rows.append(
                    {
                        "period": _period_label_for_date(d, gran),
                        "customer_id": cust.id,
                        "customer_name": cust.name,
                        "qty": 0,
                        "total": FOLLOWUP_AMOUNT,
                        "real": FOLLOWUP_AMOUNT,
                        "nhis_deposit": 0,
                        "copay": 0,
                        "self_add": 0,
                        "unpaid": unpaid_by_case.get(case_id, 0),
                        "unpaid_cum": unpaid_cum,
                        "as_cost": as_map.get((cust.id, d), 0),
                        "followup": FOLLOWUP_AMOUNT,
                    }
                )

    ctx = {
        "start": start,
        "end": end,
        "gran": gran,
        "view_mode": view_mode,
        "detail_pay_rows": detail_pay_rows,
        "detail_dep_rows": detail_dep_rows,
        "start_month": start_month if gran == "month" else "",
        "end_month": end_month if gran == "month" else "",
        "searched": searched,
        "pay_rows": pay_rows,
        "dep_rows": dep_rows,
        "pay_total": pay_total,
        "dep_total": dep_total,
    }
    return render(request, "customers/sales_analysis.html", ctx)


@login_required
def _as_paid_sum_by_pay(ps: datetime.date, pe: datetime.date):
    qs = (
        AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            amount__gt=0,
            paid_at__isnull=False,
            paid_at__gte=ps,
            paid_at__lte=pe,
        )
        .exclude(status="CANCELED")
    )
    total = sum(int(a.amount or 0) for a in qs)
    taxes = set((a.tax_type or "") for a in qs)
    return total, taxes

def _as_paid_sum_by_dep(ps: datetime.date, pe: datetime.date):
    qs = (
        AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            payment_status="PAID",
            amount__gt=0,
        )
        .exclude(status="CANCELED")
    )
    total = 0
    taxes = set()
    for a in qs:
        d = a.deposited_at or a.paid_at
        if not d or d < ps or d > pe:
            continue
        total += int(a.amount or 0)
        taxes.add(a.tax_type or "")
    return total, taxes

def _as_refund_sum(ps: datetime.date, pe: datetime.date):
    qs = (
        AfterService.objects.filter(
            customer__is_deleted=False,
            is_paid=True,
            refund_amount__gt=0,
            refund_at__isnull=False,
            refund_at__gte=ps,
            refund_at__lte=pe,
        )
        .exclude(status="CANCELED")
    )
    total = sum(int(a.refund_amount or 0) for a in qs)
    taxes = set((a.tax_type or "") for a in qs)
    return total, taxes


    # 결제일 기준 rows
    pay_rows = []
    pay_total = {"qty": 0, "total": 0, "real": 0, "tax_label": "-", "nhis": 0, "copay": 0, "self_add": 0, "as_cost": 0, "refund": 0}
    for label, ps, pe in periods:
        row = {"period": label, "qty": 0, "total": 0, "real": 0, "tax_label": "-", "nhis": 0, "copay": 0, "self_add": 0, "as_cost": 0, "refund": 0}
        taxes: set[str] = set()
        # 케이스 1회 집계: 첫 결제일/구매일이 기간에 들어오면 포함
        for c in cases:
            d = case_first_date.get(c.id)
            if not d:
                continue
            if ps <= d <= pe:
                q = _case_qty(c)
                nhis, copay, add, total = _case_amounts(c)
                row["qty"] += q
                row["nhis"] += nhis
                row["copay"] += copay
                row["self_add"] += add
                row["total"] += total
                tx = case_first_tx.get(c.id)
                if tx and tx.tax_type:
                    taxes.add(tx.tax_type)

        # A/S (결제일 기준)
        as_paid, as_tax = as_paid_sum_by_pay(ps, pe)
        as_refund, as_ref_tax = as_refund_sum_by_pay(ps, pe)
        row["as_cost"] += as_paid
        row["total"] += as_paid
        row["refund"] += as_refund
        taxes |= {t for t in (as_tax | as_ref_tax) if t}

        row["refund"] += _refund_sum(ps, pe)
        row["real"] = row["total"] - row["refund"]

        # 과세구분 라벨(기간 단위): 과세/면세/혼합, 없으면 '-'
        taxes = {t for t in taxes if t}
        if not taxes:
            row["tax_label"] = "-"
        elif len(taxes) >= 2:
            row["tax_label"] = "혼합"
        else:
            row["tax_label"] = next(iter(taxes))

        pay_rows.append(row)
        for k in ["qty", "total", "nhis", "copay", "self_add", "as_cost", "refund", "real"]:
            pay_total[k] += row[k]
    # 합계행의 tax_label은 '-' 고정
    pay_total["tax_label"] = "-"

    # 입금일 기준 rows
    dep_rows = []
    dep_total = {"qty": 0, "total": 0, "real": 0, "nhis_deposit": 0, "copay": 0, "self_add": 0, "as_cost": 0, "unpaid": 0, "unpaid_cum": 0}
    # 누적미수는 '기간 종료일(pe) 기준' 케이스별 (총계 - 누적입금) 합으로 계산
    def _case_paid_sum_until(c: CustomerCase, cutoff: datetime.date) -> int:
        txs = PaymentTransaction.objects.filter(case=c, paid_at__lte=cutoff)
        return sum(int(t.amount or 0) for t in txs)

    for label, ps, pe in periods:
        row = {"period": label, "qty": 0, "total": 0, "real": 0, "nhis_deposit": 0, "copay": 0, "self_add": 0, "as_cost": 0, "unpaid": 0, "unpaid_cum": 0, "refund": 0}
        # 기간 내 입금 합(양수만 총입금액, 음수는 환불로 별도)
        txs = PaymentTransaction.objects.filter(paid_at__gte=ps, paid_at__lte=pe)
        pos = [t for t in txs if int(t.amount or 0) > 0]
        neg = [t for t in txs if int(t.amount or 0) < 0]
        row["total"] = sum(int(t.amount or 0) for t in pos)
        row["refund"] = sum(abs(int(t.amount or 0)) for t in neg)

        # A/S (입금일 기준)
        as_paid, _as_tax = _as_paid_sum_by_dep(ps, pe)
        as_refund, _as_ref_tax = _as_refund_sum(ps, pe)
        row["as_cost"] += as_paid
        row["total"] += as_paid
        row["refund"] += as_refund
        row["real"] = row["total"] - row["refund"]

        # 보청기 수량/본인부담/자부담은 '기간 내 첫 결제일이 포함되는 케이스'로 집계 (결제일 기준과 정합)
        for c in cases:
            d = case_first_date.get(c.id)
            if not d:
                continue
            if ps <= d <= pe:
                q = _case_qty(c)
                nhis, copay, add, _total_case = _case_amounts(c)
                row["qty"] += q
                row["copay"] += copay
                row["self_add"] += add

        # 미수(기간 신규): 기간 종료일 기준 누적미수 - 기간 시작일 전 기준 누적미수
        cum_end = 0
        cum_prev = 0
        for c in cases:
            nhis, copay, add, total_case = _case_amounts(c)
            paid_end = _case_paid_sum_until(c, pe)
            paid_prev = _case_paid_sum_until(c, ps - datetime.timedelta(days=1))
            cum_end += max(total_case - paid_end, 0)
            cum_prev += max(total_case - paid_prev, 0)
        row["unpaid_cum"] = cum_end
        row["unpaid"] = max(cum_end - cum_prev, 0)

        # A/S 미수 추가(접수일 기준)
        as_unpaid_qs = AfterService.objects.filter(customer__is_deleted=False, is_paid=True, payment_status="UNPAID")
        as_cum_end = sum(int(a.amount or 0) for a in as_unpaid_qs if a.received_at and a.received_at <= pe)
        as_cum_prev = sum(int(a.amount or 0) for a in as_unpaid_qs if a.received_at and a.received_at <= (ps - datetime.timedelta(days=1)))
        row["unpaid_cum"] += as_cum_end
        row["unpaid"] += max(as_cum_end - as_cum_prev, 0)

        dep_rows.append(row)
        for k in ["qty", "total", "real", "copay", "self_add", "as_cost", "unpaid"]:
            dep_total[k] += row[k]
    # 누적미수금액 합계행은 최종 값(마지막 기간 값)
    dep_total["unpaid_cum"] = dep_rows[-1]["unpaid_cum"] if dep_rows else 0

    # 사업자 정보 (설정 > 프로필의 값을 사용)
    # - BusinessProfile은 계정 1:1 구조이므로, 현재 로그인한 계정의 프로필을 우선 사용합니다.
    # - 구버전 코드에서 business_number 같은 잘못된 필드명을 참조하면 공란이 됩니다.
    biz = None
    try:
        biz = BusinessProfile.objects.filter(user=request.user).first()
    except Exception:
        biz = None
    if not biz:
        # 혹시 데이터가 단일 프로필 형태로 남아있는 경우를 위한 안전장치
        try:
            biz = BusinessProfile.objects.first()
        except Exception:
            biz = None

    biz_name = getattr(biz, "business_name", "") if biz else ""
    biz_no = getattr(biz, "business_reg_no", "") if biz else ""

    def _html_xls(title: str, headers: list[str], rows: list[list[str]]) -> bytes:
        # 간단한 엑셀용 HTML(.xls) - 열 너비 지정
        col_widths = {
            "기간": 180,
            "보청기 수량": 110,
            "총 매출": 140,
            "총 입금액": 140,
            "실 매출": 140,
            "과세구분": 100,
            "공단 금액": 140,
            "본인부담액": 140,
            "자부담금액": 140,
            "환불": 120,
            "미수금액": 140,
            "누적미수금액": 160,
        }
        cols = "".join([f'<col style="width:{col_widths.get(h, 140)}px;">' for h in headers])
        ths = "".join([f'<th style="border:1px solid #ddd; background:#f5f6f7; padding:8px; text-align:center; white-space:nowrap;">{h}</th>' for h in headers])

        def td(v: str, align: str = "right"):
            # 텍스트로 고정(#### 방지). 숫자도 콤마 포함 텍스트로 넣습니다.
            return f'<td style="border:1px solid #eee; padding:6px 8px; text-align:{align}; white-space:nowrap;">{v}</td>'

        trs = []
        for r in rows:
            tds = []
            for i, v in enumerate(r):
                if i == 0:
                    tds.append(td(v, "left"))
                elif headers[i] == "보청기 수량":
                    tds.append(td(v, "center"))
                elif headers[i] == "과세구분":
                    tds.append(td(v, "center"))
                else:
                    tds.append(td(v, "right"))
            trs.append("<tr>" + "".join(tds) + "</tr>")

        html = f"""<html>
<head>
<meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
</head>
<body>
<table style="border-collapse:collapse; font-family:Arial, 'Malgun Gothic', sans-serif; font-size:12px;">
  <tr><td colspan="{len(headers)}" style="padding:6px 2px; font-weight:bold;">사업자명 : {biz_name}</td></tr>
  <tr><td colspan="{len(headers)}" style="padding:6px 2px; font-weight:bold;">사업자등록번호 : {biz_no}</td></tr>
  <tr><td colspan="{len(headers)}" style="padding:6px 2px;"></td></tr>
</table>
<table style="border-collapse:collapse; font-family:Arial, 'Malgun Gothic', sans-serif; font-size:12px;">
  <colgroup>{cols}</colgroup>
  <thead><tr>{ths}</tr></thead>
  <tbody>
    {''.join(trs)}
  </tbody>
</table>
</body></html>"""
        return html.encode("utf-8")

    # 결제일기준 파일 rows 구성 (합계행 포함)
    pay_headers = ["기간", "보청기 수량", "총 매출", "실 매출", "과세구분", "공단 금액", "본인부담액", "자부담금액", "후기적합", "as비용", "환불"]
    pay_data = []
    for r in pay_rows:
        pay_data.append([
            r["period"],
            f'{r["qty"]}개',
            _fmt(r["total"]),
            _fmt(r["real"]),
            r["tax_label"],
            _fmt(r["nhis"]),
            _fmt(r["copay"]),
            _fmt(r["self_add"]),
            _fmt(r.get("followup", 0)),
            _fmt(r["as_cost"]),
            _fmt(r["refund"]),
        ])
    pay_data.append([
        "전체 합계",
        f'{pay_total["qty"]}개',
        _fmt(pay_total["total"]),
        _fmt(pay_total["real"]),
        "-",
        _fmt(pay_total["nhis"]),
        _fmt(pay_total["copay"]),
        _fmt(pay_total["self_add"]),
        _fmt(pay_total.get("followup", 0)),
        _fmt(pay_total["as_cost"]),
        _fmt(pay_total["refund"]),
    ])

    # 입금일기준 파일 rows 구성 (과세구분/환불 제거, 합계행 포함)
    dep_headers = ["기간", "보청기 수량", "총 입금액", "실 매출", "본인부담액", "자부담금액", "후기적합", "as비용", "미수금액", "누적미수금액"]
    dep_data = []
    for r in dep_rows:
        dep_data.append([
            r["period"],
            f'{r["qty"]}개',
            _fmt(r["total"]),
            _fmt(r["real"]),
            _fmt(r["copay"]),
            _fmt(r["self_add"]),
            _fmt(r.get("followup", 0)),
            _fmt(r["as_cost"]),
            _fmt(r["unpaid"]),
            _fmt(r["unpaid_cum"]),
        ])
    dep_data.append([
        "전체 합계",
        f'{dep_total["qty"]}개',
        _fmt(dep_total["total"]),
        _fmt(dep_total["real"]),
        _fmt(dep_total["copay"]),
        _fmt(dep_total["self_add"]),
        _fmt(dep_total.get("followup", 0)),
        _fmt(dep_total["as_cost"]),
        _fmt(dep_total["unpaid"]),
        _fmt(dep_total["unpaid_cum"]),
    ])

    # 다운로드 로그(화면 노출 없음)
    try:
        SalesDownloadLog.objects.create(user=request.user, start_date=start, end_date=end, gran=gran)
    except Exception:
        pass

    period_key = (
        f"{start.strftime('%Y-%m')}_{end.strftime('%Y-%m')}"
        if gran == "month"
        else f"{start.isoformat()}_{end.isoformat()}"
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        # ✅ 파일명은 영문(ASCII)으로 고정
        zf.writestr(f"sales_paid_basis_{period_key}.xls", _html_xls("결제일기준", pay_headers, pay_data))
        zf.writestr(f"sales_deposit_basis_{period_key}.xls", _html_xls("입금일기준", dep_headers, dep_data))

    resp = HttpResponse(buf.getvalue(), content_type="application/zip")
    # 파일명: sales_data_(기간).zip
    resp["Content-Disposition"] = f'attachment; filename="sales_data_{period_key}.zip"'
    return resp


@login_required
def sales_analysis_vat_export(request: HttpRequest) -> HttpResponse:
    """설정 > 매출분석: 매출자료 다운로드(zip)

    **기존 양식(HTML-table .xls + zip) 유지**
    - 결제일 기준 파일: 제품/결제(PaymentTransaction.paid_at) + 유상 A/S(AfterService.paid_at)
    - 입금일 기준 파일: 제품/결제(PaymentTransaction.paid_at) + 유상 A/S(AfterService.deposited_at)
    - A/S 포함 조건(절대): 유상 + 결제추가(수납완료) + 완료(COMPLETED) + 취소 제외
    """
    if not (request.user.is_staff or request.user.is_superuser):
        return HttpResponseForbidden("Forbidden")

    today = timezone.localdate()

    def _parse_date(key: str, default: datetime.date) -> datetime.date:
        s = (request.GET.get(key) or "").strip()
        if not s:
            return default
        try:
            return datetime.date.fromisoformat(s)
        except Exception:
            return default

    end = _parse_date("end", today)
    start = _parse_date("start", end - datetime.timedelta(days=30))
    if start > end:
        start, end = end, start

    from .models import PaymentTransaction, AfterService, BusinessProfile, CustomerCase

    # 사업자 정보
    biz_name = ""
    biz_reg = ""
    try:
        bp = BusinessProfile.objects.filter(user=request.user).first()
        if bp:
            biz_name = (bp.business_name or "").strip()
            biz_reg = (bp.business_registration_number or "").strip()
    except Exception:
        pass

    gran = (request.GET.get("gran") or "day").strip()
    if gran not in ("day", "week", "month"):
        gran = "day"

    def _periods(start_d: datetime.date, end_d: datetime.date, g: str):
        periods: list[tuple[str, datetime.date, datetime.date]] = []
        if g == "day":
            cur = start_d
            while cur <= end_d:
                periods.append((cur.strftime("%Y-%m-%d"), cur, cur))
                cur = cur + datetime.timedelta(days=1)
            return periods
        if g == "week":
            # 일~토: 일요일이 주 시작
            cur = start_d - datetime.timedelta(days=(start_d.weekday() + 1) % 7)
            while cur <= end_d:
                w_start = cur
                w_end = cur + datetime.timedelta(days=6)
                ps = max(w_start, start_d)
                pe = min(w_end, end_d)
                label = f"{ps.strftime('%Y-%m-%d')}~{pe.strftime('%Y-%m-%d')}"
                periods.append((label, ps, pe))
                cur = cur + datetime.timedelta(days=7)
            return [p for p in periods if p[1] <= p[2]]

        # month
        cur = datetime.date(start_d.year, start_d.month, 1)
        while cur <= end_d:
            last_day = calendar.monthrange(cur.year, cur.month)[1]
            m_start = cur
            m_end = datetime.date(cur.year, cur.month, last_day)
            label = cur.strftime("%Y.%m")
            periods.append((label, max(m_start, start_d), min(m_end, end_d)))
            if cur.month == 12:
                cur = datetime.date(cur.year + 1, 1, 1)
            else:
                cur = datetime.date(cur.year, cur.month + 1, 1)
        return [p for p in periods if p[1] <= p[2]]

    periods = _periods(start, end, gran)

    def _sum_by_day_map(day_map: dict, ps: datetime.date, pe: datetime.date) -> int:
        total = 0
        cur = ps
        while cur <= pe:
            total += int(day_map.get(cur, 0) or 0)
            cur += datetime.timedelta(days=1)
        return total

    # 제품/결제: 결제일 기준(수납일 paid_at)
    tx_qs = PaymentTransaction.objects.filter(paid_at__gte=start, paid_at__lte=end)
    tx_amt_by_day = {}
    tx_tax_by_day = {}
    tx_refund_by_day = {}
    for paid_at, amount, tax_type, origin_tx_id in tx_qs.values_list("paid_at", "amount", "tax_type", "origin_tx_id"):
        d = paid_at
        tx_amt_by_day[d] = tx_amt_by_day.get(d, 0) + int(amount or 0)
        if tax_type:
            tx_tax_by_day.setdefault(d, set()).add(tax_type)
        # 환불 거래(음수 또는 origin_tx 연결된 자식)를 환불 컬럼으로 집계(절댓값)
        if (amount or 0) < 0 or origin_tx_id:
            tx_refund_by_day[d] = tx_refund_by_day.get(d, 0) + abs(int(amount or 0))

    # 케이스(판매): 구매일 기준으로 수량/공단/본인부담/자부담(추가) 집계
    case_qs = CustomerCase.objects.filter(purchase_date__gte=start, purchase_date__lte=end)
    qty_by_day = {}
    nhis_by_day = {}
    copay_by_day = {}
    self_add_by_day = {}
    for d, nhis, copay in case_qs.values_list("purchase_date", "nhis_amount", "copay_amount"):
        if not d:
            continue
        qty_by_day[d] = qty_by_day.get(d, 0) + 1
        nhis_by_day[d] = nhis_by_day.get(d, 0) + int(nhis or 0)
        copay_by_day[d] = copay_by_day.get(d, 0) + int(copay or 0)

    # 자부담금액(추가)은 purchase_date_add 기준으로 집계
    case_add_qs = CustomerCase.objects.filter(purchase_date_add__gte=start, purchase_date_add__lte=end)
    for d, self_add in case_add_qs.values_list("purchase_date_add", "self_pay_amount_add"):
        if not d:
            continue
        self_add_by_day[d] = self_add_by_day.get(d, 0) + int(self_add or 0)

    # A/S 포함 조건:
    # - 유상(is_paid=True) + 결제완료(결제일 paid_at 존재)만 포함
    # - 취소(CANCELED) 제외
    # - "입금일 기준" 파일도 입금일을 보지 않고 결제완료 즉시 포함 => paid_at 기준으로 동일 집계
    as_base = (
        AfterService.objects.filter(
            is_paid=True,
            amount__gt=0,
            paid_at__isnull=False,
        )
        .exclude(status="CANCELED")
    )

    # 결제일 기준 A/S 비용/환불 (net = amount - refund_amount)
    as_paid_qs = as_base.filter(paid_at__gte=start, paid_at__lte=end)
    as_cost_paid = {}
    as_refund_paid = {}
    for d, amount, refund_amount in as_paid_qs.values_list("paid_at", "amount", "refund_amount"):
        if not d:
            continue
        amt = int(amount or 0)
        ref = int(refund_amount or 0)
        net = amt - ref
        if net < 0:
            net = 0
        as_cost_paid[d] = as_cost_paid.get(d, 0) + net
        if ref > 0:
            as_refund_paid[d] = as_refund_paid.get(d, 0) + ref

    # 입금일 기준 A/S 비용/환불: 요구사항대로 결제완료 즉시 포함 => paid_at 기준 동일 집계
    as_dep_qs = as_paid_qs
    as_cost_dep = dict(as_cost_paid)
    as_refund_dep = dict(as_refund_paid)

    def _tax_label(taxes: set[str]) -> str:
        taxes = {t for t in taxes if t}
        if not taxes:
            return "-"
        if len(taxes) == 1:
            return list(taxes)[0]
        return "혼합"

    def _fmt_money(n: int) -> str:
        try:
            return f"{int(n):,}"
        except Exception:
            return "0"

    def _escape(s: str) -> str:
        return html.escape(s or "", quote=True)

    def _html_xls(title: str, headers: list[str], rows: list[list[str]]) -> str:
        # 샘플 파일 형식(HTML table .xls)과 동일한 톤으로 생성
        col_widths = [180, 110, 140, 140, 100, 140, 140, 140, 140, 110]
        if len(headers) != len(col_widths):
            col_widths = [140] * len(headers)

        cols = "".join([f'<col style="width:{w}px;">' for w in col_widths])

        ths = "".join(
            [
                '<th style="border:1px solid #ddd; background:#f5f6f7; padding:8px; text-align:center; white-space:nowrap;">'
                + _escape(h)
                + "</th>"
                for h in headers
            ]
        )

        def td(v: str, align: str = "right") -> str:
            return (
                f'<td style="border:1px solid #eee; padding:6px 8px; text-align:{align}; white-space:nowrap;">{_escape(v)}</td>'
            )

        trs = []
        for r in rows:
            tds = []
            for i, v in enumerate(r):
                h = headers[i]
                if h == "기간":
                    tds.append(td(v, "center"))
                elif h == "보청기 수량":
                    tds.append(td(v, "center"))
                elif h == "과세구분":
                    tds.append(td(v, "center"))
                else:
                    tds.append(td(v, "right"))
            trs.append("<tr>" + "".join(tds) + "</tr>")

        return f"""<html>
<head>
<meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
</head>
<body>
<table style="border-collapse:collapse; font-family:Arial, 'Malgun Gothic', sans-serif; font-size:12px;">
  <tr><td colspan="{len(headers)}" style="padding:6px 2px; font-weight:bold;">사업자명 : {_escape(biz_name)}</td></tr>
  <tr><td colspan="{len(headers)}" style="padding:6px 2px; font-weight:bold;">사업자등록번호 : {_escape(biz_reg)}</td></tr>
  <tr><td colspan="{len(headers)}" style="padding:6px 2px;"></td></tr>
</table>
<table style="border-collapse:collapse; font-family:Arial, 'Malgun Gothic', sans-serif; font-size:12px;">
  <colgroup>{cols}</colgroup>
  <thead><tr>{ths}</tr></thead>
  <tbody>
    {''.join(trs)}
  </tbody>
</table>
</body>
</html>"""

    # ✅ 매출상세 다운로드(mode=detail): 결제일 기준 + 입금일 기준 2개 파일을 zip으로 제공합니다.
    if (request.GET.get("mode") or "").strip() == "detail":
        # 상세는 일(day) 기준만
        gran = "day"
        periods = _periods(start, end, gran)

        # A/S (유상 + 결제완료 + 취소 제외) 맵: (customer_id, date) -> net
        as_map: dict[tuple[int, datetime.date], int] = {}
        as_qs = (
            AfterService.objects.filter(
                is_paid=True,
                amount__gt=0,
                paid_at__isnull=False,
            )
            .exclude(status="CANCELED")
            .values_list("customer_id", "paid_at", "amount", "refund_amount")
        )
        for cust_id, paid_at, amount, refund_amount in as_qs:
            net = int(amount or 0) - int(refund_amount or 0)
            if net < 0:
                net = 0
            as_map[(int(cust_id), paid_at)] = as_map.get((int(cust_id), paid_at), 0) + net

        FOLLOWUP_AMOUNT = 50000

        def _case_qty_detail(c: CustomerCase) -> int:
            return 2 if getattr(c, "has_sub", False) else 1

        def _case_expected_total(c: CustomerCase) -> int:
            nhis = int(getattr(c, "nhis_amount", 0) or 0)
            copay = int(getattr(c, "copay_amount", 0) or 0)
            self_add = int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0
            return nhis + copay + self_add

        # 결제일 기준 상세 rows
        headers_paid_detail = [
            "기간",
            "고객id",
            "고객명",
            "보청기 수량",
            "총 매출",
            "실 매출",
            "과세구분",
            "공단 금액",
            "본인부담액",
            "자부담금액",
            "후기적합",
            "as비용",
            "환불",
        ]

        paid_detail_rows: list[list[str]] = []
        tx_qs = (
            PaymentTransaction.objects.filter(
                paid_at__gte=start,
                paid_at__lte=end,
                case__customer__is_deleted=False,
            )
            .select_related("case", "case__customer")
            .order_by("case__customer__name", "case__customer__id", "paid_at", "id")
        )
        for t in tx_qs:
            c = t.case
            cust = c.customer
            nhis = int(getattr(c, "nhis_amount", 0) or 0)
            copay = int(getattr(c, "copay_amount", 0) or 0)
            self_add = int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0
            total_case = nhis + copay + self_add
            amt = int(t.amount or 0)
            refund_amt = abs(amt) if amt < 0 else 0
            real_amt = amt if amt > 0 else 0
            paid_detail_rows.append(
                [
                    t.paid_at.strftime("%Y-%m-%d"),
                    str(cust.id),
                    cust.name,
                    f"{_case_qty_detail(c)}개",
                    _fmt_money(total_case),
                    _fmt_money(real_amt),
                    (t.tax_type or "-"),
                    _fmt_money(nhis),
                    _fmt_money(copay),
                    _fmt_money(self_add),
                    _fmt_money(0),
                    _fmt_money(as_map.get((cust.id, t.paid_at), 0)),
                    _fmt_money(refund_amt),
                ]
            )

        # 후기적합(제출완료일 기준) 상세 행
        cases = CustomerCase.objects.filter(customer__is_deleted=False).select_related("customer")
        for c in cases:
            cust = getattr(c, "customer", None)
            if not cust or getattr(cust, "is_deleted", False):
                continue
            for n in (1, 2, 3, 4):
                dt = getattr(c, f"fu{n}_submitted_at", None)
                if not dt:
                    continue
                try:
                    d = dt.date()
                except Exception:
                    continue
                if not (start <= d <= end):
                    continue
                paid_detail_rows.append(
                    [
                        d.strftime("%Y-%m-%d"),
                        str(cust.id),
                        cust.name,
                        "0개",
                        _fmt_money(FOLLOWUP_AMOUNT),
                        _fmt_money(FOLLOWUP_AMOUNT),
                        "-",
                        _fmt_money(0),
                        _fmt_money(0),
                        _fmt_money(0),
                        _fmt_money(FOLLOWUP_AMOUNT),
                        _fmt_money(as_map.get((cust.id, d), 0)),
                        _fmt_money(0),
                    ]
                )

        # 정렬(고객명/ID/일자)
        paid_detail_rows.sort(key=lambda r: (r[2], int(r[1]), r[0]))

        # 입금일 기준 상세 rows
        headers_dep_detail = [
            "기간",
            "고객id",
            "고객명",
            "보청기 수량",
            "총 입금액",
            "실 매출",
            "공단 입금액",
            "본인부담액",
            "자부담금액",
            "후기적합",
            "미수금액",
            "누적미수금액",
            "as비용",
        ]

        # 케이스/수납/공단입금 기반 이벤트 생성(화면 매출상세와 동일 산식)
        cases = (
            CustomerCase.objects.filter(customer__is_deleted=False)
            .select_related("customer")
            .order_by("customer__name", "customer__id", "id")
        )
        case_cache: dict[int, CustomerCase] = {c.id: c for c in cases}

        dep_cases = (
            CustomerCase.objects.filter(
                nhis_deposit_date__isnull=False,
                nhis_deposit_date__gte=start,
                nhis_deposit_date__lte=end,
                customer__is_deleted=False,
            )
            .select_related("customer")
            .order_by("customer__name", "customer__id", "nhis_deposit_date", "id")
        )

        pay_amounts_by_case_date: dict[tuple[int, datetime.date], list[int]] = {}
        for t in tx_qs:
            pay_amounts_by_case_date.setdefault((t.case_id, t.paid_at), []).append(int(t.amount or 0))

        dep_amount_by_case: dict[int, int] = {}
        for c in dep_cases:
            dep_amount_by_case[c.id] = int(getattr(c, "nhis_deposit_amount", 0) or 0)

        unpaid_by_case: dict[int, int] = {}
        received_by_case: dict[int, int] = {}
        cases_by_customer: dict[int, list[int]] = {}
        for c in cases:
            if getattr(c.customer, "is_deleted", False):
                continue
            cases_by_customer.setdefault(c.customer_id, []).append(c.id)
            unpaid_by_case[c.id] = _case_expected_total(c)
            received_by_case[c.id] = 0

        events: list[tuple[int, datetime.date, int, str]] = []
        for (case_id, d), _amts in pay_amounts_by_case_date.items():
            c = case_cache.get(case_id)
            if not c:
                continue
            events.append((c.customer_id, d, case_id, "pay"))
        for c in dep_cases:
            events.append((c.customer_id, c.nhis_deposit_date, c.id, "nhis"))
        for c in cases:
            if getattr(c.customer, 'is_deleted', False):
                continue
            for n in (1, 2, 3, 4):
                dd = getattr(c, f'fu{n}_deposit_date', None)
                if dd and start <= dd <= end:
                    events.append((c.customer_id, dd, c.id, 'fu'))
        events.sort(key=lambda x: (x[0], x[1], x[2], x[3]))

        dep_detail_rows: list[list[str]] = []
        for cust_id, d, case_id, kind in events:
            c = case_cache.get(case_id)
            if not c:
                continue
            cust = c.customer
            if not cust or getattr(cust, "is_deleted", False):
                continue

            unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))

            if kind == "pay":
                for amt in pay_amounts_by_case_date.get((case_id, d), []):
                    received_by_case[case_id] = received_by_case.get(case_id, 0) + int(amt)
                    unpaid_by_case[case_id] = max(_case_expected_total(c) - received_by_case[case_id], 0)
                    unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))
                    dep_detail_rows.append(
                        [
                            d.strftime("%Y-%m-%d"),
                            str(cust.id),
                            cust.name,
                            f"{_case_qty_detail(c)}개",
                            _fmt_money(int(amt) if int(amt) > 0 else 0),
                            _fmt_money(int(amt) if int(amt) > 0 else 0),
                            _fmt_money(0),
                            _fmt_money(int(getattr(c, "copay_amount", 0) or 0)),
                            _fmt_money(int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0),
                            _fmt_money(0),
                            _fmt_money(unpaid_by_case[case_id]),
                            _fmt_money(unpaid_cum),
                            _fmt_money(as_map.get((cust.id, d), 0)),
                        ]
                    )
            elif kind == "nhis":
                dep_amt = dep_amount_by_case.get(case_id, 0)
                if dep_amt == 0:
                    continue
                received_by_case[case_id] = received_by_case.get(case_id, 0) + int(dep_amt)
                unpaid_by_case[case_id] = max(_case_expected_total(c) - received_by_case[case_id], 0)
                unpaid_cum = sum(unpaid_by_case.get(x, 0) for x in cases_by_customer.get(cust_id, []))
                dep_detail_rows.append(
                    [
                        d.strftime("%Y-%m-%d"),
                        str(cust.id),
                        cust.name,
                        f"{_case_qty_detail(c)}개",
                        _fmt_money(int(dep_amt)),
                        _fmt_money(int(dep_amt)),
                        _fmt_money(int(dep_amt)),
                        _fmt_money(int(getattr(c, "copay_amount", 0) or 0)),
                        _fmt_money(int(getattr(c, "self_pay_amount_add", 0) or 0) if getattr(c, "has_sub", False) else 0),
                        _fmt_money(0),
                        _fmt_money(unpaid_by_case[case_id]),
                        _fmt_money(unpaid_cum),
                        _fmt_money(as_map.get((cust.id, d), 0)),
                    ]
                )
            else:  # fu
                dep_detail_rows.append(
                    [
                        d.strftime("%Y-%m-%d"),
                        str(cust.id),
                        cust.name,
                        "0개",
                        _fmt_money(FOLLOWUP_AMOUNT),
                        _fmt_money(FOLLOWUP_AMOUNT),
                        _fmt_money(0),
                        _fmt_money(0),
                        _fmt_money(0),
                        _fmt_money(FOLLOWUP_AMOUNT),
                        _fmt_money(unpaid_by_case.get(case_id, 0)),
                        _fmt_money(unpaid_cum),
                        _fmt_money(as_map.get((cust.id, d), 0)),
                    ]
                )

        dep_detail_rows.sort(key=lambda r: (r[2], int(r[1]), r[0]))

        period_key = f"{start.isoformat()}_{end.isoformat()}"
        buf = io.BytesIO()
        import zipfile as _zipfile
        with _zipfile.ZipFile(buf, "w", compression=_zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"sales_detail_paid_basis_{period_key}.xls", _html_xls("결제일기준", headers_paid_detail, paid_detail_rows))
            zf.writestr(f"sales_detail_deposit_basis_{period_key}.xls", _html_xls("입금일기준", headers_dep_detail, dep_detail_rows))

        resp = HttpResponse(buf.getvalue(), content_type="application/zip")
        resp["Content-Disposition"] = f'attachment; filename="sales_detail_{period_key}.zip"'
        return resp

    headers = ["기간", "보청기 수량", "총 매출", "실 매출", "과세구분", "공단 금액", "본인부담액", "자부담금액", "as비용", "환불"]

    def _build_rows(kind: str) -> list[list[str]]:
        rows = []
        for label, ps, pe in periods:
            qty = _sum_by_day_map(qty_by_day, ps, pe)
            nhis = _sum_by_day_map(nhis_by_day, ps, pe)
            copay = _sum_by_day_map(copay_by_day, ps, pe)
            self_add = _sum_by_day_map(self_add_by_day, ps, pe)

            # 총매출(판매 기준): 공단+본인부담(+자부담)
            total_sales = (nhis or 0) + (copay or 0) + (self_add or 0)

            # 실매출(수납 기준): 제품 수납(순액)만(환불은 환불 컬럼으로 별도 표기)
            real_sales = _sum_by_day_map(tx_amt_by_day, ps, pe)

            # 과세구분: 기간 내 과세구분 합쳐서 표시
            tax_set: set[str] = set()
            cur = ps
            while cur <= pe:
                tax_set |= set(tx_tax_by_day.get(cur, set()) or set())
                cur += datetime.timedelta(days=1)
            tax_label = _tax_label(tax_set)

            # A/S 비용/환불
            as_cost_map = as_cost_paid if kind == "paid" else as_cost_dep
            as_refund_map = as_refund_paid if kind == "paid" else as_refund_dep
            as_cost = _sum_by_day_map(as_cost_map, ps, pe)
            refunds = _sum_by_day_map(tx_refund_by_day, ps, pe) + _sum_by_day_map(as_refund_map, ps, pe)

            rows.append(
                [
                    label,
                    str(qty),
                    _fmt_money(total_sales),
                    _fmt_money(real_sales),
                    tax_label,
                    _fmt_money(nhis),
                    _fmt_money(copay),
                    _fmt_money(self_add),
                    _fmt_money(as_cost),
                    _fmt_money(refunds),
                ]
            )
        return rows

    pay_rows = _build_rows("paid")
    dep_rows = _build_rows("dep")

    period_key = (
        f"{start.strftime('%Y-%m')}_{end.strftime('%Y-%m')}"
        if gran == "month"
        else f"{start.isoformat()}_{end.isoformat()}"
    )

    buf = io.BytesIO()
    import zipfile as _zipfile
    with _zipfile.ZipFile(buf, "w", compression=_zipfile.ZIP_DEFLATED) as zf:
        # ✅ 파일명은 영문(ASCII)으로 고정
        zf.writestr(f"sales_paid_basis_{period_key}.xls", _html_xls("결제일기준", headers, pay_rows))
        zf.writestr(f"sales_deposit_basis_{period_key}.xls", _html_xls("입금일기준", headers, dep_rows))

    resp = HttpResponse(buf.getvalue(), content_type="application/zip")
    resp["Content-Disposition"] = f'attachment; filename="sales_data_{period_key}.zip"'
    return resp
def settings_profile(request: HttpRequest) -> HttpResponse:
    """설정 > 프로필: 로그인한 센터(계정) 사업자 정보."""

    profile, _ = BusinessProfile.objects.get_or_create(user=request.user)

    CONSENT_TEXT = (
        "입력하신 사업자 정보는 공단 청구 및 후기적합 신청 관련 서류의 자동 작성·처리를 위한 목적으로만 이용됩니다.\n"
        "해당 정보는 위 목적 외에 마케팅 활용, 제3자 제공 등 다른 용도로는 수집·이용되지 않습니다."
    )

    def _client_ip(req: HttpRequest) -> str:
        xf = req.META.get("HTTP_X_FORWARDED_FOR")
        if xf:
            return xf.split(",")[0].strip()
        return (req.META.get("REMOTE_ADDR") or "").strip()

    if request.method == "POST":
        form = BusinessProfileForm(request.POST, instance=profile, user=request.user)
        if form.is_valid():
            obj = form.save(commit=False)

            # 주민등록번호: 입력값이 있으면 그대로 저장(폼에서 13자리 검증)
            rrn_direct = (form.cleaned_data.get("rep_rrn_full") or "").strip()
            if rrn_direct:
                obj.rep_rrn_full = rrn_direct

            # 동의 체크(최초 1회 필수)
            if not profile.consent_agreed:
                if request.POST.get("consent_agree") != "on":
                    form.add_error(None, "사업자 정보 이용 동의가 필요합니다.")
                    ctx = {
                        "form": form,
                        "profile": profile,
                        "saved": False,
                        "rep_rrn_masked": profile.rep_rrn_masked,
                        "can_view_full_rrn": bool(request.user.is_staff or request.user.is_superuser),
                        "consent_required": True,
                        "consent_text": CONSENT_TEXT,
                    }
                    return render(request, "customers/profile.html", ctx)
                obj.consent_agreed = True
                obj.consent_agreed_at = timezone.now()
                obj.consent_text = CONSENT_TEXT

            obj.save()

            # 동의 로그(최초 동의 시)
            if obj.consent_agreed and not profile.consent_agreed:
                from .models import BusinessProfileConsentLog
                BusinessProfileConsentLog.objects.create(
                    profile=obj,
                    user=request.user,
                    consent_text=CONSENT_TEXT,
                    ip=_client_ip(request),
                    user_agent=(request.META.get("HTTP_USER_AGENT") or "")[:255],
                )
            # 저장 직후 새로고침 없이 즉시 마스킹 값이 화면에 보이도록 즉시 렌더링합니다.
            profile = obj
            form = BusinessProfileForm(instance=profile, user=request.user)
            ctx = {
                "form": form,
                "profile": profile,
                "saved": True,
                "rep_rrn_masked": profile.rep_rrn_masked,
                "can_view_full_rrn": bool(request.user.is_staff or request.user.is_superuser),
                "consent_required": (not profile.consent_agreed),
                "consent_text": CONSENT_TEXT,
            }
            return render(request, "customers/profile.html", ctx)
    else:
        form = BusinessProfileForm(instance=profile, user=request.user)

    ctx = {
        "form": form,
        "profile": profile,
        "saved": request.GET.get("saved") == "1",
        "rep_rrn_masked": profile.rep_rrn_masked,
        "can_view_full_rrn": bool(request.user.is_staff or request.user.is_superuser),
        "consent_required": (not profile.consent_agreed),
        "consent_text": CONSENT_TEXT,
    }
    return render(request, "customers/profile.html", ctx)


@login_required
def settings_profile_reveal_rrn(request: HttpRequest) -> JsonResponse:
    """대표자 주민등록번호 전체 보기(JSON). 관리자만 가능 + 접근 로그 기록."""
    if not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    profile, _ = BusinessProfile.objects.get_or_create(user=request.user)
    if not profile.consent_agreed:
        return JsonResponse({"ok": False, "error": "consent_required"}, status=400)

    def _client_ip(req: HttpRequest) -> str:
        xf = req.META.get("HTTP_X_FORWARDED_FOR")
        if xf:
            return xf.split(",")[0].strip()
        return (req.META.get("REMOTE_ADDR") or "").strip()

    from .models import BusinessProfileAccessLog
    BusinessProfileAccessLog.objects.create(
        profile=profile,
        user=request.user,
        action="RRN_REVEAL",
        ip=_client_ip(request),
        user_agent=(request.META.get("HTTP_USER_AGENT") or "")[:255],
    )

    return JsonResponse({"ok": True, "rrn": profile.rep_rrn_full})


@login_required
def settings_access_logs(request: HttpRequest) -> HttpResponse:
    """설정 > 접근 로그(관리자용)."""
    if not request.user.is_superuser:
        return HttpResponseForbidden("Forbidden")

    from .models import BusinessProfileAccessLog, BusinessProfileConsentLog

    access_qs = BusinessProfileAccessLog.objects.select_related("user", "profile").order_by("-created_at")[:500]
    consent_qs = BusinessProfileConsentLog.objects.select_related("user", "profile").order_by("-created_at")[:500]

    return render(
        request,
        "customers/access_logs.html",
        {"access_logs": access_qs, "consent_logs": consent_qs},
    )


def _purge_expired_trash() -> None:
    cutoff = timezone.now() - datetime.timedelta(days=30)
    expired = Customer.objects.filter(is_deleted=True, deleted_at__isnull=False, deleted_at__lt=cutoff)
    for c in expired:
        from .models import CustomerTrashLog
        CustomerTrashLog.objects.create(customer=c, user=None, action="purge")
        c.delete()


@login_required
def trash_list(request: HttpRequest) -> HttpResponse:
    _purge_expired_trash()
    per_page = request.GET.get("per_page") or "30"
    if per_page not in ("30", "50", "100", "200"):
        per_page = "30"
    per_page_n = int(per_page)

    qs = Customer.objects.filter(is_deleted=True).order_by("-deleted_at", "-id")
    paginator = Paginator(qs, per_page_n)
    page_number = request.GET.get("page") or "1"
    page_obj = paginator.get_page(page_number)

    return render(
        request,
        "customers/trash_list.html",
        {"customers": page_obj.object_list, "page_obj": page_obj, "paginator": paginator, "total": qs.count(), "per_page": per_page},
    )




@login_required
def trash_bulk_restore(request: HttpRequest) -> HttpResponse:
    if request.method != "POST":
        return redirect("/customers/settings/trash/")

    ids = request.POST.getlist("ids")
    if not ids:
        return redirect("/customers/settings/trash/")

    qs = Customer.objects.filter(id__in=ids, is_deleted=True)
    for c in qs:
        c.is_deleted = False
        c.deleted_at = None
        c.deleted_by = None
        c.save(update_fields=["is_deleted", "deleted_at", "deleted_by"])
        from .models import CustomerTrashLog
        CustomerTrashLog.objects.create(customer=c, user=request.user, action="restore")
    return redirect("/customers/settings/trash/")


@login_required
def trash_bulk_purge(request: HttpRequest) -> HttpResponse:
    if request.method != "POST":
        return redirect("/customers/settings/trash/")

    ids = request.POST.getlist("ids")
    if not ids:
        return redirect("/customers/settings/trash/")

    qs = Customer.objects.filter(id__in=ids, is_deleted=True)
    for c in qs:
        from .models import CustomerTrashLog
        CustomerTrashLog.objects.create(customer=c, user=request.user, action="purge")
        c.delete()
    return redirect("/customers/settings/trash/")

@login_required
def trash_restore(request: HttpRequest, pk: int) -> HttpResponse:
    if request.method != "POST":
        return redirect("/customers/settings/trash/")

    c = get_object_or_404(Customer, pk=pk, is_deleted=True)
    c.is_deleted = False
    c.deleted_at = None
    c.deleted_by = None
    c.save(update_fields=["is_deleted", "deleted_at", "deleted_by"])
    from .models import CustomerTrashLog
    CustomerTrashLog.objects.create(customer=c, user=request.user, action="restore")
    return redirect("/customers/settings/trash/")


@login_required
def trash_purge(request: HttpRequest, pk: int) -> HttpResponse:
    if request.method != "POST":
        return redirect("/customers/settings/trash/")

    c = get_object_or_404(Customer, pk=pk, is_deleted=True)
    from .models import CustomerTrashLog
    CustomerTrashLog.objects.create(customer=c, user=request.user, action="purge")
    c.delete()
    return redirect("/customers/settings/trash/")


@login_required
def clear_product_payment(request: HttpRequest, pk: int) -> HttpResponse:
    # 안전장치: customer_detail에서도 처리하지만, 별도 엔드포인트도 유지
    customer = get_object_or_404(Customer, pk=pk, is_deleted=False)
    if request.method != "POST":
        return redirect(_redirect_url("제품/결제"))

    case_id = request.POST.get("case_id")
    case = None
    if case_id and str(case_id).isdigit():
        case = CustomerCase.objects.filter(id=int(case_id), customer=customer).first()
    if case is None:
        case = _get_latest_case(customer)
    if case is None:
        return redirect(_redirect_url("제품/결제"))

    case.manufacturer = ""
    case.model_name = ""
    case.serial_number = ""
    case.standard_code = ""
    case.manufacture_date = None
    case.receiver = ""
    case.pre_fit_date = None
    case.purchase_date = None
    case.side = ""
    case.earmold_made_date = None
    case.nhis_amount = None
    case.copay_amount = None

    case.manufacturer_add = ""
    case.model_name_add = ""
    case.serial_number_add = ""
    case.standard_code_add = ""
    case.manufacture_date_add = None
    case.receiver_add = ""
    case.pre_fit_date_add = None
    case.purchase_date_add = None
    case.side_add = ""
    case.self_pay_amount_add = None

    case.has_sub = False
    case.save()
    update_customer_stage(customer)
    return redirect(_redirect_url("제품/결제", case))

# ----------------------------
# Calendar (상담 탭 방문예약 연동)
# ----------------------------
@login_required
def calendar_home(request: HttpRequest) -> HttpResponse:
    """캘린더(월/주/일). 상담 탭의 방문예약(Consultation.visit_reservation_at)과 연결."""
    view = (request.GET.get("view") or "month").strip()
    if view not in ("month", "week", "day"):
        view = "week"

    # 기준 날짜
    date_raw = (request.GET.get("date") or "").strip()
    try:
        base_date = datetime.date.fromisoformat(date_raw) if date_raw else timezone.localdate()
    except Exception:
        base_date = timezone.localdate()

    # 알림/이동 하이라이트용(선택 날짜 강조)
    highlight_date = (request.GET.get("hl_date") or "").strip()

    # range 계산
    # 주 시작 요일: 일요일(일~토)로 고정
    # - Python weekday(): 월=0 ... 일=6
    # - 일요일 시작 offset: (weekday + 1) % 7
    if view == "month":
        first = base_date.replace(day=1)
        last_day = calendar.monthrange(first.year, first.month)[1]
        last = base_date.replace(day=last_day)
        start = first - datetime.timedelta(days=((first.weekday() + 1) % 7))  # 일요일 시작
        end = last + datetime.timedelta(days=((5 - last.weekday()) % 7))      # 토요일 끝
    elif view == "day":
        start = base_date
        end = base_date
    else:
        start = base_date - datetime.timedelta(days=((base_date.weekday() + 1) % 7))
        end = start + datetime.timedelta(days=6)

    day_list = []
    cur = start
    while cur <= end:
        day_list.append(cur)
        cur += datetime.timedelta(days=1)

    start_dt = timezone.make_aware(datetime.datetime.combine(start, datetime.time.min))
    end_dt = timezone.make_aware(datetime.datetime.combine(end, datetime.time.max))

    is_super = bool(request.user.is_superuser)
    # 관리자 전용: 취소 일정 보기 모드
    cancel_mode = is_super and ((request.GET.get("cancel") or "").strip() == "1")
    center_param = (request.GET.get("center") or "1").strip()
    center_on_param = (center_param != "0")
    extra_qs = f"&center={'1' if center_on_param else '0'}" + ("&cancel=1" if cancel_mode else "")

    # 사용자 선호(센터일정 표시 ON/OFF) - 기본값 ON
    # - 일반 모드: 세션 값을 프론트 초기값으로 사용
    # - 취소 모드: URL 파라미터(center=0/1)가 우선
    center_on_pref = bool(request.session.get("calendar_center_on", True))

    qs = (
        Consultation.objects
        .select_related("customer")
        .filter(customer__is_deleted=False, visit_reservation_at__isnull=False)
        .filter(visit_reservation_at__gte=start_dt, visit_reservation_at__lte=end_dt)
        .order_by("visit_reservation_at", "id")
    )
    # 표시 모드:
    # - 기본: 취소 일정 숨김
    # - 관리자 + cancel=1:
    #   - center=1: 센터 취소 일정만 표시(고객 일정은 숨김)
    #   - center=0: 고객 취소 일정만 표시(센터 일정은 숨김)
    if cancel_mode:
        if center_on_param:
            qs = Consultation.objects.none()
        else:
            qs = qs.filter(outcome="취소")
    else:
        qs = qs.exclude(outcome="취소")

    events_by_day: dict[str, list[dict]] = {d.isoformat(): [] for d in day_list}
    for c in qs:
        dt = timezone.localtime(c.visit_reservation_at)
        key = dt.date().isoformat()
        if key not in events_by_day:
            continue
        is_canceled = (c.outcome == "취소")
        events_by_day[key].append({
            "consultation_id": c.id,
            "customer_id": c.customer_id,
            "customer_name": c.customer.name,
            "time": dt.strftime("%H:%M"),
            "title": f"{c.customer.name}",
            "url": f"/customers/{c.customer_id}/?tab=상담&hl={c.id}",
            "dt_value": dt.strftime("%Y-%m-%dT%H:%M"),
            "status": ("CANCELED" if is_canceled else "ACTIVE"),
        })


    # 센터 일정 (휴가/외근/회의)
    center_events_by_day: dict[str, list[dict]] = {d.isoformat(): [] for d in day_list}
    center_qs = (
        CenterEvent.objects
        .select_related("created_by")
        .filter(start_at__lte=end_dt, end_at__gte=start_dt)
        .order_by("start_at", "id")
    )
    if cancel_mode:
        if center_on_param:
            center_qs = center_qs.filter(status="CANCELED")
        else:
            center_qs = CenterEvent.objects.none()
    else:
        center_qs = center_qs.filter(status="ACTIVE")
    for ev in center_qs:
        start_local = timezone.localtime(ev.start_at)
        end_local = timezone.localtime(ev.end_at)
        cur_day = start_local.date()
        last_day = end_local.date()
        while cur_day <= last_day:
            key = cur_day.isoformat()
            if key in center_events_by_day:
                display_title = ev.title.strip() if (ev.event_type == "기타" and ev.title) else ev.event_type
                center_events_by_day[key].append({
                    "id": ev.id,
                    "type": ev.event_type,
                    "title": display_title,
                    "time": start_local.strftime("%H:%M") if cur_day == start_local.date() else "",
                    "start": start_local.strftime("%Y-%m-%dT%H:%M"),
                    "end": end_local.strftime("%Y-%m-%dT%H:%M"),
                    "memo": ev.memo,
                    "created_by": (ev.created_by.get_username() if ev.created_by else ""),
                    "status": ev.status,
                })
            cur_day += datetime.timedelta(days=1)

    # 오늘 미리보기 (week/day에서만 사용)
    today = timezone.localdate()
    today_key = today.isoformat()
    today_events = events_by_day.get(today_key, [])

    # 네비게이션용 prev/next
    if view == "month":
        prev_date = (base_date.replace(day=1) - datetime.timedelta(days=1)).replace(day=1)
        next_date = (base_date.replace(day=calendar.monthrange(base_date.year, base_date.month)[1]) + datetime.timedelta(days=1)).replace(day=1)
        title = f"{base_date.year}년 {base_date.month}월"
    elif view == "day":
        prev_date = base_date - datetime.timedelta(days=1)
        next_date = base_date + datetime.timedelta(days=1)
        title = base_date.strftime("%Y-%m-%d")
    else:
        prev_date = start - datetime.timedelta(days=7)
        next_date = start + datetime.timedelta(days=7)
        title = f"{start.strftime('%Y-%m-%d')} ~ {end.strftime('%Y-%m-%d')}"

    ctx = {
        "view": view,
        "base_date": base_date,
        "title": title,
        "days": day_list,
        "events_by_day": events_by_day,
        "center_events_by_day": center_events_by_day,
        "is_superuser": is_super,
        "cancel_mode": bool(cancel_mode),
        "center_on_param": bool(center_on_param),
        "center_on_pref": bool(center_on_pref),
        "extra_qs": extra_qs,
        "today": today,
        "today_events": today_events,
        "prev_date": prev_date,
        "next_date": next_date,
        "highlight_date": highlight_date,
    }
    return render(request, "customers/calendar.html", ctx)


@login_required
def calendar_prefs(request: HttpRequest) -> JsonResponse:
    """캘린더 화면 토글(센터일정/취소일정) 선호 저장 (AJAX)."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method_not_allowed"}, status=405)

    try:
        import json
        payload = json.loads((request.body or b"{}").decode("utf-8"))
    except Exception:
        payload = {}

    if "center_on" in payload:
        v = payload.get("center_on")
        on = True if v in (True, 1, "1", "true", "True", "on", "ON") else False
        request.session["calendar_center_on"] = on

    if "cancel_on" in payload:
        v = payload.get("cancel_on")
        on = True if v in (True, 1, "1", "true", "True", "on", "ON") else False
        request.session["calendar_cancel_on"] = on

    return JsonResponse({"ok": True})


@login_required
def calendar_update_reservation(request: HttpRequest, consultation_id: int) -> HttpResponse:
    """캘린더/상담 탭에서 방문예약 변경(사유 필수) → 원본 데이터 수정 + 변경 이력 1줄 기록."""

    def _redirect_fallback() -> HttpResponse:
        next_url = (request.POST.get("next") or "").strip()
        if next_url and url_has_allowed_host_and_scheme(
            next_url,
            allowed_hosts={request.get_host()},
            require_https=request.is_secure(),
        ):
            return redirect(next_url)
        return redirect("customers:calendar_home")

    if request.method != "POST":
        return redirect("customers:calendar_home")

    consultation = get_object_or_404(Consultation, pk=consultation_id, customer__is_deleted=False)
    new_dt_raw = (request.POST.get("visit_reservation_at") or "").strip()
    reason = (request.POST.get("reason") or "").strip()

    if not new_dt_raw or not reason:
        return _redirect_fallback()

    try:
        naive = datetime.datetime.strptime(new_dt_raw, "%Y-%m-%dT%H:%M")
        # 10분 단위로 스냅
        snapped_minute = (naive.minute // 10) * 10
        naive = naive.replace(minute=snapped_minute, second=0, microsecond=0)
        new_dt = timezone.make_aware(naive)
    except Exception:
        return _redirect_fallback()

    # 과거 날짜 예약 수정 제한: 관리자(슈퍼유저)만 허용
    today = timezone.localdate()
    if (consultation.visit_reservation_at and timezone.localtime(consultation.visit_reservation_at).date() < today) or (timezone.localtime(new_dt).date() < today):
        if not request.user.is_superuser:
            return _redirect_fallback()

    old_dt = consultation.visit_reservation_at
    consultation.visit_reservation_at = new_dt
    consultation.updated_by = request.user
    consultation.save(update_fields=["visit_reservation_at", "updated_by", "updated_at"])

    try:
        ConsultationReservationChangeLog.objects.create(
            consultation=consultation,
            customer=consultation.customer,
            changed_by=request.user,
            old_reservation_at=old_dt,
            new_reservation_at=new_dt,
            reason=reason[:120],
        )
    except Exception:
        pass

    return _redirect_fallback()

@login_required
def api_reservation_conflict(request: HttpRequest) -> JsonResponse:
    """같은 시간대(동일 datetime)의 방문예약 중복 여부를 확인합니다."""
    at_raw = (request.GET.get("at") or "").strip()
    exclude_id_raw = (request.GET.get("exclude") or "").strip()
    if not at_raw:
        return JsonResponse({"count": 0, "items": []})
    try:
        naive = datetime.datetime.strptime(at_raw, "%Y-%m-%dT%H:%M")
        snapped_minute = (naive.minute // 10) * 10
        naive = naive.replace(minute=snapped_minute, second=0, microsecond=0)
        target_dt = timezone.make_aware(naive)
    except Exception:
        return JsonResponse({"count": 0, "items": []})

    qs = (
        Consultation.objects
        .select_related("customer")
        .filter(customer__is_deleted=False, visit_reservation_at=target_dt)
    )
    if exclude_id_raw.isdigit():
        qs = qs.exclude(id=int(exclude_id_raw))

    items = []
    for c in qs[:20]:
        dt = timezone.localtime(c.visit_reservation_at) if c.visit_reservation_at else None
        items.append({
            "consultation_id": c.id,
            "customer_id": c.customer_id,
            "customer_name": c.customer.name,
            "time": dt.strftime("%H:%M") if dt else "",
            "url": f"/customers/{c.customer_id}/?tab=상담&hl={c.id}",
        })
    return JsonResponse({"count": len(items), "items": items})


@login_required
def center_event_create(request: HttpRequest) -> JsonResponse:
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method"})
    event_type = (request.POST.get("event_type") or "").strip()
    title = (request.POST.get("title") or "").strip()
    memo = (request.POST.get("memo") or "").strip()
    start_raw = (request.POST.get("start_at") or "").strip()
    end_raw = (request.POST.get("end_at") or "").strip()

    valid_types = ("휴가", "외근", "회의", "교육", "기타")
    if event_type not in valid_types or not start_raw or not end_raw:
        return JsonResponse({"ok": False, "error": "invalid"})
    # '기타'인 경우만 제목(=기타 제목) 필수
    if event_type == "기타" and not title:
        return JsonResponse({"ok": False, "error": "need_title"})

    try:
        start_naive = datetime.datetime.strptime(start_raw, "%Y-%m-%dT%H:%M")
        end_naive = datetime.datetime.strptime(end_raw, "%Y-%m-%dT%H:%M")
        start_at = timezone.make_aware(start_naive)
        end_at = timezone.make_aware(end_naive)
    except Exception:
        return JsonResponse({"ok": False, "error": "datetime"})

    if end_at < start_at:
        return JsonResponse({"ok": False, "error": "range"})

    # 유형이 기타가 아닌 경우 title은 비워둡니다(화면에서는 유형만 노출).
    save_title = (title[:80] if event_type == "기타" else "")
    ev = CenterEvent.objects.create(
        event_type=event_type,
        title=save_title,
        memo=memo[:120],
        start_at=start_at,
        end_at=end_at,
        created_by=request.user,
    )
    try:
        CenterEventLog.objects.create(
            event=ev,
            action="CREATE",
            actor=request.user,
            before_json="",
            after_json=json.dumps({
                "event_type": ev.event_type,
                "title": ev.title,
                "start_at": ev.start_at.isoformat(),
                "end_at": ev.end_at.isoformat(),
                "memo": ev.memo,
                "status": ev.status,
            }, ensure_ascii=False),
        )
    except Exception:
        pass
    return JsonResponse({"ok": True, "id": ev.id})


@login_required
def center_event_update(request: HttpRequest, event_id: int) -> JsonResponse:
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method"})
    ev = get_object_or_404(CenterEvent, pk=event_id)
    # 권한: 관리자 또는 본인 일정만 수정 가능
    if not request.user.is_superuser:
        if (not ev.created_by) or (ev.created_by_id != request.user.id):
            return JsonResponse({"ok": False, "error": "forbidden"}, status=403)
    # 취소된 일정은 복구/수정 불가
    if ev.status == "CANCELED":
        return JsonResponse({"ok": False, "error": "canceled"}, status=400)
    event_type = (request.POST.get("event_type") or ev.event_type).strip()
    title = (request.POST.get("title") or ev.title).strip()
    memo = (request.POST.get("memo") or ev.memo).strip()
    start_raw = (request.POST.get("start_at") or "").strip()
    end_raw = (request.POST.get("end_at") or "").strip()

    valid_types = ("휴가", "외근", "회의", "교육", "기타")
    if event_type not in valid_types:
        return JsonResponse({"ok": False, "error": "invalid"})
    if event_type == "기타" and not title:
        return JsonResponse({"ok": False, "error": "need_title"})

    try:
        if start_raw:
            ev.start_at = timezone.make_aware(datetime.datetime.strptime(start_raw, "%Y-%m-%dT%H:%M"))
        if end_raw:
            ev.end_at = timezone.make_aware(datetime.datetime.strptime(end_raw, "%Y-%m-%dT%H:%M"))
    except Exception:
        return JsonResponse({"ok": False, "error": "datetime"})

    if ev.end_at < ev.start_at:
        return JsonResponse({"ok": False, "error": "range"})

    before = {
        "event_type": ev.event_type,
        "title": ev.title,
        "start_at": ev.start_at.isoformat() if ev.start_at else "",
        "end_at": ev.end_at.isoformat() if ev.end_at else "",
        "memo": ev.memo,
        "status": ev.status,
    }
    ev.event_type = event_type
    ev.title = (title[:80] if event_type == "기타" else "")
    ev.memo = memo[:120]
    ev.save()
    try:
        CenterEventLog.objects.create(
            event=ev,
            action="UPDATE",
            actor=request.user,
            before_json=json.dumps(before, ensure_ascii=False),
            after_json=json.dumps({
                "event_type": ev.event_type,
                "title": ev.title,
                "start_at": ev.start_at.isoformat(),
                "end_at": ev.end_at.isoformat(),
                "memo": ev.memo,
                "status": ev.status,
            }, ensure_ascii=False),
        )
    except Exception:
        pass
    return JsonResponse({"ok": True})


@login_required
def center_event_cancel(request: HttpRequest, event_id: int) -> JsonResponse:
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method"})
    ev = get_object_or_404(CenterEvent, pk=event_id)
    # 권한: 관리자 또는 본인 일정만 취소 가능
    if not request.user.is_superuser:
        if (not ev.created_by) or (ev.created_by_id != request.user.id):
            return JsonResponse({"ok": False, "error": "forbidden"}, status=403)
    if ev.status == "CANCELED":
        return JsonResponse({"ok": True})

    before = {
        "event_type": ev.event_type,
        "title": ev.title,
        "start_at": ev.start_at.isoformat() if ev.start_at else "",
        "end_at": ev.end_at.isoformat() if ev.end_at else "",
        "memo": ev.memo,
        "status": ev.status,
    }
    ev.status = "CANCELED"
    ev.canceled_at = timezone.now()
    ev.canceled_by = request.user
    ev.save(update_fields=["status", "canceled_at", "canceled_by", "updated_at"])
    try:
        CenterEventLog.objects.create(
            event=ev,
            action="CANCEL",
            actor=request.user,
            before_json=json.dumps(before, ensure_ascii=False),
            after_json=json.dumps({
                "event_type": ev.event_type,
                "title": ev.title,
                "start_at": ev.start_at.isoformat(),
                "end_at": ev.end_at.isoformat(),
                "memo": ev.memo,
                "status": ev.status,
            }, ensure_ascii=False),
        )
    except Exception:
        pass
    return JsonResponse({"ok": True})


# ==============================
# 재고관리(Inventory) - CSV 업/다운로드
# ==============================


def _inv_get_or_create_mfr(name: str):
    from .models import InventoryManufacturer

    name = (name or "").strip()
    obj, _ = InventoryManufacturer.objects.get_or_create(name=name)
    return obj


def _inv_get_or_create_model(
    mfr,
    model_name: str,
    item_type: str = "SERIAL",
    product_kind: str = "HEARING_AID",
    importer_name: str | None = None,
):
    from .models import InventoryProductModel

    model_name = (model_name or "").strip()
    item_type = (item_type or "SERIAL").strip().upper()
    product_kind = (product_kind or "HEARING_AID").strip().upper()
    if product_kind not in ("HEARING_AID", "DOME", "RECEIVER"):
        product_kind = "HEARING_AID"
    if item_type not in ("SERIAL", "QTY"):
        item_type = "SERIAL"

    defaults = {"item_type": item_type, "product_kind": product_kind}
    if importer_name is not None:
        defaults["importer_name"] = (importer_name or "").strip()

    obj, created = InventoryProductModel.objects.get_or_create(
        manufacturer=mfr, model_name=model_name, defaults=defaults
    )

    # ✅ importer_name 동기화(값이 있으면 최신값으로)
    try:
        if importer_name is not None:
            imp = (importer_name or "").strip()
            if imp and (getattr(obj, "importer_name", "") or "") != imp:
                obj.importer_name = imp
                obj.save(update_fields=["importer_name"])
    except Exception:
        pass

    # ✅ 소프트 삭제된 모델은 재입고 시 자동 복구
    try:
        if (not created) and getattr(obj, "is_deleted", False):
            obj.is_deleted = False
            obj.deleted_at = None
            obj.deleted_reason = ""
            obj.save(update_fields=["is_deleted", "deleted_at", "deleted_reason"])
    except Exception:
        pass
    # 기존 모델이 있는데 타입/품목구분이 다른 경우: 기존 값 유지(운영 안전)
    if (not created) and (obj.item_type != item_type):
        pass
    try:
        if (not created) and getattr(obj, "product_kind", "") and (obj.product_kind != product_kind):
            pass
    except Exception:
        pass
    return obj


def _inv_unit_snapshot(u) -> dict:
    try:
        m = u.product_model
        return {
            "item_type": "SERIAL",
            "manufacturer": m.manufacturer.name,
            "model_name": m.model_name,
            "serial_no": u.serial_no,
            "standard_code": u.standard_code,
            "mfg_date": (u.mfg_date.isoformat() if u.mfg_date else ""),
            "status": u.status,
        }
    except Exception:
        return {
            "item_type": "SERIAL",
            "serial_no": getattr(u, "serial_no", ""),
            "standard_code": getattr(u, "standard_code", ""),
            "mfg_date": "",
            "status": getattr(u, "status", ""),
        }


def _inv_qty_snapshot(pm) -> dict:
    try:
        return {
            "item_type": "QTY",
            "manufacturer": pm.manufacturer.name,
            "model_name": pm.model_name,
            "standard_code": "",
            "mfg_date": "",
            "qty_current": int(pm.qty_current or 0),
        }
    except Exception:
        return {"item_type": "QTY", "qty_current": 0}


def _inv_check_and_notify(pm):
    """재고 상태 변화(음수/안전재고) 알림 생성/중복 방지/상태회복 시 플래그 초기화.

    - 알림은 공통(사용자 분리 X)
    - 30일 경과 알림은 자동 삭제
    """

    try:
        import urllib.parse
        from .models import Notification
        from django.utils import timezone
        from datetime import timedelta

        # 30일 자동삭제
        cutoff = timezone.now() - timedelta(days=30)
        try:
            Notification.objects.filter(created_at__lt=cutoff).delete()
        except Exception:
            pass

        qty = int(pm.qty_current or 0)
        mfr = pm.manufacturer.name if getattr(pm, "manufacturer", None) else ""
        model = pm.model_name
        link = f"/customers/inventory/?q={urllib.parse.quote((mfr + ' ' + model).strip())}"

        # 음수 재고
        if qty < 0 and (not getattr(pm, "negative_alerted", False)):
            Notification.objects.create(
                kind="INVENTORY",
                title="재고 음수",
                message=f"{mfr} {model} 현재 수량이 {qty} 입니다.",
                link=link,
                is_read=False,
            )
            pm.negative_alerted = True
            pm.save(update_fields=["negative_alerted"])
        elif qty >= 0 and getattr(pm, "negative_alerted", False):
            pm.negative_alerted = False
            pm.save(update_fields=["negative_alerted"])

        # 안전재고(임계치)
        threshold = int(getattr(pm, "alert_threshold", 0) or 0)
        if threshold > 0 and qty <= threshold and (not getattr(pm, "threshold_alerted", False)):
            Notification.objects.create(
                kind="INVENTORY",
                title="재고 부족",
                message=f"{mfr} {model} 현재 수량 {qty} (안전재고 {threshold})",
                link=link,
                is_read=False,
            )
            pm.threshold_alerted = True
            pm.save(update_fields=["threshold_alerted"])
        elif threshold > 0 and qty > threshold and getattr(pm, "threshold_alerted", False):
            pm.threshold_alerted = False
            pm.save(update_fields=["threshold_alerted"])

    except Exception:
        return


def _inv_read_csv_file(f) -> list:
    import csv

    raw = f.read()
    # excel utf-8-sig 대응
    try:
        text = raw.decode("utf-8-sig")
    except Exception:
        text = raw.decode("utf-8", errors="replace")

    lines = text.splitlines()
    reader = csv.DictReader(lines)
    rows = []
    for r in reader:
        rows.append({(k or "").strip(): (v or "").strip() for k, v in (r or {}).items()})
    return rows


def _inv_make_fail_csv_bytes(fail_rows: list, fieldnames: list) -> bytes:
    import csv
    import io

    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=fieldnames)
    writer.writeheader()
    for r in fail_rows:
        writer.writerow(r)
    return buf.getvalue().encode("utf-8-sig")


@login_required
def inventory_home(request: HttpRequest):
    """재고관리: 입고(업로드+양식), 출고/조정(업로드/다운로드/양식) - CSV(UTF-8 BOM)"""

    from django.db.models import Q, Count
    from django.http import HttpResponse
    import urllib.parse
    import json
    import datetime
    import base64
    import csv
    import io

    from .models import InventoryProductModel, InventoryUnit, InventoryStockEvent

    tab = (request.GET.get("tab") or "입고").strip()
    if tab not in ("입고", "조정", "출고"):
        tab = "입고"

    q = (request.GET.get("q") or "").strip()

    # pagination
    try:
        per_page = int((request.GET.get("per_page") or "50").strip())
    except Exception:
        per_page = 50
    if per_page not in (30, 50, 100):
        per_page = 50
    try:
        page = int((request.GET.get("page") or "1").strip())
    except Exception:
        page = 1
    if page < 1:
        page = 1

    # 실패 CSV 다운로드 (세션 저장)
    if (request.GET.get("download_fail") or "") == "1":
        b64 = request.session.pop("inv_fail_csv_b64", "")
        if not b64:
            return HttpResponse("No fail file", status=404)
        data = base64.b64decode(b64.encode("ascii"))
        resp = HttpResponse(data, content_type="text/csv; charset=utf-8")
        _inv_set_attachment(resp, "fail_rows.csv")
        return resp

    # 다운로드(양식/내역)
    download = (request.GET.get("download") or "").strip()
    if download:
        # 파일명용 사업자명(프로필)
        from .models import BusinessProfile

        def _sanitize_filename(s: str) -> str:
            """ASCII-only filename slug (avoids browsers saving as a generic 'download')."""
            from django.utils.text import slugify

            s = (s or "").strip()
            # ASCII only; if name is non-ASCII (e.g., Korean), slugify may return empty -> fallback.
            s = slugify(s, allow_unicode=False)
            return s or "center"

        biz_name = ""
        try:
            bp = BusinessProfile.objects.filter(user=request.user).first()
            if bp:
                biz_name = (bp.business_name or "").strip()
        except Exception:
            biz_name = ""
        biz_name_safe = _sanitize_filename(biz_name)
        ymd = timezone.localdate().strftime('%Y%m%d')

        buf = io.StringIO()
        w = csv.writer(buf)

        # 헤더(한글 고정) - 업로드된 템플릿(1.csv/2.csv/3.csv) 기준
        receive_header = ["유형", "제조사", "제조수입업소명", "모델명", "제조번호", "표준코드", "제조일", "현재 수량"]
        ship_header = ["유형", "제조사", "제조수입업소명", "모델명", "제조번호", "표준코드", "제조일", "현재 수량"]
        adjust_header = [
            "유형",
            "구분",
            "제조사",
            "제조수입업소명",
            "모델명",
            "제조번호",
            "표준코드",
            "제조일",
            "조정수량",
            "조정사유",
            "담당자",
            "조정일",
        ]

        def _kind_kor(pk: str) -> str:
            pk = (pk or "").strip().upper()
            if pk == "DOME":
                return "돔"
            if pk == "RECEIVER":
                return "리시버"
            return "보청기"

        if download == "receive_template":
            w.writerow(receive_header)
            # 업로드 템플릿(1.csv) 예시와 동일하게 제공
            w.writerow(["보청기", "벨톤", "GN", "모델A", "SN001", "STD001", "2026-01-01", ""])
            w.writerow(["돔", "벨톤", "GN", "파워돔", "", "", "", "10"])
            w.writerow(["리시버", "벨톤", "GN", "HP", "", "", "", "5"])
            filename = f"inventory_receive_template_{ymd}_{biz_name_safe}.csv"

        elif download == "ship_template":
            w.writerow(ship_header)
            # 업로드 템플릿(3.csv) 예시와 동일하게 제공
            w.writerow(["보청기", "벨톤", "GN", "모델A", "SN001", "STD001", "2026-01-01", "1"])
            w.writerow(["돔", "벨톤", "GN", "파워돔", "", "", "", "100"])
            w.writerow(["리시버", "벨톤", "GN", "HP", "", "", "", "100"])
            filename = f"inventory_release_template_{ymd}_{biz_name_safe}.csv"

        elif download == "ship_download":
            w.writerow(ship_header)
            qs = (
                InventoryStockEvent.objects.filter(event_type="SHIP")
                .select_related("unit__product_model__manufacturer", "product_model__manufacturer", "created_by")
                .order_by("-created_at", "-id")[:5000]
            )
            for ev in qs:
                if ev.unit_id:
                    u = ev.unit
                    pm = u.product_model
                    mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                    imp = (getattr(pm, "importer_name", "") or "") if pm else ""
                    model = pm.model_name if pm else ""
                    kind_k = _kind_kor(getattr(pm, "product_kind", "HEARING_AID"))
                    w.writerow([
                        kind_k,
                        mfr,
                        imp,
                        model,
                        u.serial_no,
                        u.standard_code,
                        u.mfg_date.isoformat() if u.mfg_date else "",
                        1,
                    ])
                else:
                    pm = ev.product_model
                    mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                    imp = (getattr(pm, "importer_name", "") or "") if pm else ""
                    model = pm.model_name if pm else ""
                    qty = abs(int(ev.qty_delta or 0))
                    kind_k = _kind_kor(getattr(pm, "product_kind", "HEARING_AID"))
                    w.writerow([
                        kind_k,
                        mfr,
                        imp,
                        model,
                        "",
                        "",
                        "",
                        qty,
                    ])
            filename = f"inventory_release_{ymd}_{biz_name_safe}.csv"

        elif download == "adjust_template":
            w.writerow(adjust_header)
            # 업로드 템플릿(2.csv) 예시와 동일하게 제공
            w.writerow(["보청기", "데이터수정", "벨톤", "GN", "모델A", "SN001", "STD001", "2026-01-01", "", "오타 수정", "", ""])
            w.writerow(["보청기", "분실", "벨톤", "GN", "모델B", "SN002", "STD002", "2026-01-01", "", "분실 처리", "", ""])
            w.writerow(["돔", "데이터수정", "벨톤", "GN", "파워돔", "", "", "", "-1", "수량 조정", "", ""])
            w.writerow(["리시버", "데이터수정", "벨톤", "GN", "HP", "", "", "", "-1", "수량 조정", "", ""])
            filename = f"inventory_adjust_template_{ymd}_{biz_name_safe}.csv"

        elif download == "adjust_download":
            w.writerow(adjust_header)
            qs = (
                InventoryStockEvent.objects.filter(event_type="ADJUST")
                .select_related("unit__product_model__manufacturer", "product_model__manufacturer", "created_by")
                .order_by("-created_at", "-id")[:5000]
            )
            for ev in qs:
                kind_display = ev.get_adjust_kind_display() if hasattr(ev, "get_adjust_kind_display") else (ev.adjust_kind or "")
                if ev.unit_id:
                    u = ev.unit
                    pm = u.product_model
                    mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                    imp = (getattr(pm, "importer_name", "") or "") if pm else ""
                    model = pm.model_name if pm else ""
                    kind_k = _kind_kor(getattr(pm, "product_kind", "HEARING_AID"))
                    w.writerow([
                        kind_k,
                        kind_display,
                        mfr,
                        imp,
                        model,
                        u.serial_no,
                        u.standard_code,
                        u.mfg_date.isoformat() if u.mfg_date else "",
                        "",
                        ev.reason or "",
                        ev.created_by.get_username() if ev.created_by else "",
                        ev.created_at.date().isoformat(),
                    ])
                else:
                    pm = ev.product_model
                    mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                    imp = (getattr(pm, "importer_name", "") or "") if pm else ""
                    model = pm.model_name if pm else ""
                    kind_k = _kind_kor(getattr(pm, "product_kind", "HEARING_AID"))
                    w.writerow([
                        kind_k,
                        kind_display,
                        mfr,
                        imp,
                        model,
                        "",
                        "",
                        "",
                        int(ev.qty_delta or 0),
                        ev.reason or "",
                        ev.created_by.get_username() if ev.created_by else "",
                        ev.created_at.date().isoformat(),
                    ])
            filename = f"inventory_adjust_{ymd}_{biz_name_safe}.csv"

        else:
            return HttpResponse("Invalid download", status=400)

        data = buf.getvalue().encode("utf-8-sig")  # Excel 한글 깨짐 방지
        resp = HttpResponse(data, content_type="text/csv; charset=utf-8")
        quoted = urllib.parse.quote(filename)
        resp["Content-Disposition"] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quoted}'
        return resp

    # 업로드 처리
    message = None
    error = None
    fail_available = False

    if request.method == "POST":
        action = (request.POST.get("action") or "").strip()

        # 관리자 전용: 오등록 삭제(완전 삭제)
        if action == "delete_item":
            import json

            if not (request.user.is_staff or request.user.is_superuser):
                request.session["inv_err"] = "삭제는 관리자만 가능합니다."
                return redirect(
                    "/customers/inventory/?"
                    + urllib.parse.urlencode({"tab": tab, "q": q, "per_page": per_page, "page": page})
                )

            item_type = (request.POST.get("item_type") or "").strip().upper()
            unit_id = (request.POST.get("unit_id") or "").strip()
            pm_id = (request.POST.get("product_model_id") or "").strip()
            reason = (request.POST.get("reason") or "").strip() or "오등록으로 인한 삭제"

            try:
                with transaction.atomic():
                    if item_type == "SERIAL":
                        if not unit_id:
                            raise ValueError("대상 제조번호가 없습니다")
                        u = InventoryUnit.objects.select_related("product_model__manufacturer").get(id=int(unit_id))
                        before = _inv_unit_snapshot(u)
                        InventoryStockEvent.objects.create(
                            unit=u,
                            product_model=None,
                            qty_delta=0,
                            event_type="ADJUST",
                            progress_status="DONE",
                            adjust_kind="DELETE",
                            reason=reason,
                            before_json=json.dumps(before, ensure_ascii=False),
                            after_json=json.dumps(_inv_qty_snapshot(pm), ensure_ascii=False),
                            created_by=request.user,
                        )
                        # 관련 이벤트도 함께 삭제(완전 삭제 요구)
                        InventoryStockEvent.objects.filter(unit=u).delete()
                        u.delete()
                    else:
                        if not pm_id:
                            raise ValueError("대상 품목이 없습니다")
                        pm = InventoryProductModel.objects.select_related("manufacturer").get(id=int(pm_id))
                        before = _inv_qty_snapshot(pm)
                        InventoryStockEvent.objects.create(
                            unit=None,
                            product_model=pm,
                            qty_delta=0,
                            event_type="ADJUST",
                            progress_status="DONE",
                            adjust_kind="DELETE",
                            reason=reason,
                            before_json=json.dumps(before, ensure_ascii=False),
                            after_json=json.dumps(_inv_qty_snapshot(pm), ensure_ascii=False),
                            created_by=request.user,
                        )
                        InventoryStockEvent.objects.filter(product_model=pm).delete()
                        pm.delete()
                request.session["inv_msg"] = "삭제가 완료되었습니다."
            except Exception as e:
                request.session["inv_err"] = f"삭제 실패: {e}"

            qs = {"tab": tab}
            if q:
                qs["q"] = q
            qs["per_page"] = per_page
            qs["page"] = page
            return redirect("/customers/inventory/?" + urllib.parse.urlencode(qs))

        # 화면에서 일괄 수량 조정(수량형만)
        if action == "adjust_bulk_qty":
            import json

            ids = request.POST.getlist("product_model_ids")
            qty_delta_str = (request.POST.get("qty_delta") or "").strip()
            reason = (request.POST.get("reason") or "").strip()

            if not ids:
                request.session["inv_err"] = "일괄 적용할 항목을 선택해주세요."
            elif not reason:
                request.session["inv_err"] = "조정사유는 필수입니다."
            else:
                # ✅ 오등록 삭제(소프트 삭제)
                if reason == "오등록으로 인한 삭제":
                    ok = 0
                    fail = 0
                    fail_msgs = []
                    try:
                        with transaction.atomic():
                            pms = list(
                                InventoryProductModel.objects.select_related("manufacturer").filter(
                                    id__in=[int(x) for x in ids if str(x).strip().isdigit()], item_type="QTY", is_deleted=False
                                )
                            )
                            for pm in pms:
                                try:
                                    before = _inv_qty_snapshot(pm)
                                    # ✅ 정책 변경: 출고(SHIP)는 무시하고, '현재 남아있는 재고'만 대상으로 삭제/정리
                                    # - QTY 품목을 is_deleted 처리하면 전액환불 재고원복(+1) 대상 탐색이 막힐 수 있어
                                    #   현재 수량을 0으로 정리만 하고 모델은 유지합니다.
                                    pm.qty_current = 0
                                    pm.deleted_at = timezone.now()
                                    pm.deleted_reason = reason
                                    pm.save(update_fields=["qty_current", "deleted_at", "deleted_reason"])

                                    InventoryStockEvent.objects.create(
                                        unit=None,
                                        product_model=pm,
                                        qty_delta=0,
                                        event_type="ADJUST",
                                        progress_status="DONE",
                                        adjust_kind="DELETE",
                                        reason=reason,
                                        before_json=json.dumps(before, ensure_ascii=False),
                                        after_json=json.dumps(_inv_qty_snapshot(pm), ensure_ascii=False),
                                        created_by=request.user,
                                    )
                                    ok += 1
                                except Exception as e:
                                    fail += 1
                                    fail_msgs.append(f"{pm.manufacturer.name} {pm.model_name}: {e}")

                        if ok:
                            request.session["inv_msg"] = f"삭제 완료: 성공 {ok} / 실패 {fail}"
                        if fail:
                            request.session["inv_err"] = "일부 실패: " + " | ".join(fail_msgs[:5])
                    except Exception as e:
                        request.session["inv_err"] = f"삭제 실패: {e}"

                # ✅ 일반 수량 조정
                else:
                    try:
                        qty_delta = int(float(qty_delta_str))
                    except Exception:
                        qty_delta = 0
                    if qty_delta == 0:
                        request.session["inv_err"] = "조정수량은 0이 될 수 없습니다."
                    else:
                        ok = 0
                        fail = 0
                        fail_msgs = []
                        try:
                            with transaction.atomic():
                                pms = list(
                                    InventoryProductModel.objects.select_related("manufacturer").filter(
                                        id__in=[int(x) for x in ids if str(x).strip().isdigit()], item_type="QTY", is_deleted=False
                                    )
                                )
                                for pm in pms:
                                    try:
                                        before = _inv_qty_snapshot(pm)
                                        new_qty = int(pm.qty_current or 0) + qty_delta
                                        pm.qty_current = new_qty
                                        pm.save(update_fields=["qty_current"])
                                        _inv_check_and_notify(pm)
                                        after = _inv_qty_snapshot(pm)
                                        InventoryStockEvent.objects.create(
                                            unit=None,
                                            product_model=pm,
                                            qty_delta=qty_delta,
                                            event_type="ADJUST",
                                            progress_status="DONE",
                                            adjust_kind="DATA_CORRECTION",
                                            reason=reason,
                                            before_json=json.dumps(before, ensure_ascii=False),
                                            after_json=json.dumps(after, ensure_ascii=False),
                                            created_by=request.user,
                                        )
                                        ok += 1
                                    except Exception as e:
                                        fail += 1
                                        fail_msgs.append(f"{pm.manufacturer.name} {pm.model_name}: {e}")
                            if ok:
                                request.session["inv_msg"] = f"일괄 조정 완료: 성공 {ok} / 실패 {fail}"
                            if fail:
                                request.session["inv_err"] = "일부 실패: " + " | ".join(fail_msgs[:5])
                        except Exception as e:
                            request.session["inv_err"] = f"일괄 조정 실패: {e}"

            qs = {"tab": tab}
            if q:
                qs["q"] = q
            qs["per_page"] = per_page
            qs["page"] = page
            return redirect("/customers/inventory/?" + urllib.parse.urlencode(qs))

        # 화면에서 일괄 입고(수량형만) - 보청기는 CSV 업로드 안내
        if action == "receive_bulk_qty":
            import json

            ids = request.POST.getlist("product_model_ids")
            qty_str = (request.POST.get("qty") or "").strip()
            if not ids:
                request.session["inv_err"] = "일괄 입고할 항목을 선택해주세요."
            else:
                try:
                    qty = int(float(qty_str))
                except Exception:
                    qty = 0
                if qty <= 0:
                    request.session["inv_err"] = "입고 수량은 1 이상이어야 합니다."
                else:
                    ok = 0
                    try:
                        with transaction.atomic():
                            pms = list(
                                InventoryProductModel.objects.select_related("manufacturer").filter(
                                    id__in=[int(x) for x in ids if str(x).strip().isdigit()], item_type="QTY"
                                )
                            )
                            for pm in pms:
                                before = _inv_qty_snapshot(pm)
                                pm.item_type = "QTY"
                                pm.qty_current = int(pm.qty_current or 0) + qty
                                pm.save(update_fields=["item_type", "qty_current"])
                                _inv_check_and_notify(pm)
                                after = _inv_qty_snapshot(pm)
                                InventoryStockEvent.objects.create(
                                    unit=None,
                                    product_model=pm,
                                    qty_delta=qty,
                                    event_type="RECEIVE",
                                    progress_status="DONE",
                                    adjust_kind="",
                                    reason="일괄 입고",
                                    before_json=json.dumps(before, ensure_ascii=False),
                                    after_json=json.dumps(after, ensure_ascii=False),
                                    created_by=request.user,
                                )
                                ok += 1
                        request.session["inv_msg"] = f"일괄 입고 완료: {ok}건"
                    except Exception as e:
                        request.session["inv_err"] = f"일괄 입고 실패: {e}"

            qs = {"tab": tab}
            if q:
                qs["q"] = q
            qs["per_page"] = per_page
            qs["page"] = page
            return redirect("/customers/inventory/?" + urllib.parse.urlencode(qs))

        # 화면에서 일괄 출고(수량형만) - 보청기는 CSV 업로드 안내
        if action == "ship_bulk_qty":
            import json

            ids = request.POST.getlist("product_model_ids")
            qty_str = (request.POST.get("qty") or "").strip()
            reason = (request.POST.get("reason") or "").strip()
            if not ids:
                request.session["inv_err"] = "일괄 출고할 항목을 선택해주세요."
            elif not reason:
                request.session["inv_err"] = "사유는 필수입니다."
            else:
                try:
                    qty = int(float(qty_str))
                except Exception:
                    qty = 0
                if qty <= 0:
                    request.session["inv_err"] = "출고 수량은 1 이상이어야 합니다."
                else:
                    ok = 0
                    try:
                        with transaction.atomic():
                            pms = list(
                                InventoryProductModel.objects.select_related("manufacturer").filter(
                                    id__in=[int(x) for x in ids if str(x).strip().isdigit()], item_type="QTY"
                                )
                            )
                            for pm in pms:
                                before = _inv_qty_snapshot(pm)
                                pm.item_type = "QTY"
                                pm.qty_current = int(pm.qty_current or 0) - qty
                                pm.save(update_fields=["item_type", "qty_current"])
                                _inv_check_and_notify(pm)
                                after = _inv_qty_snapshot(pm)
                                InventoryStockEvent.objects.create(
                                    unit=None,
                                    product_model=pm,
                                    qty_delta=-qty,
                                    event_type="SHIP",
                                    progress_status="DONE",
                                    adjust_kind="",
                                    reason=reason,
                                    before_json=json.dumps(before, ensure_ascii=False),
                                    after_json=json.dumps(after, ensure_ascii=False),
                                    created_by=request.user,
                                )
                                ok += 1
                        request.session["inv_msg"] = f"일괄 출고 완료: {ok}건"
                    except Exception as e:
                        request.session["inv_err"] = f"일괄 출고 실패: {e}"

            qs = {"tab": tab}
            if q:
                qs["q"] = q
            qs["per_page"] = per_page
            qs["page"] = page
            return redirect("/customers/inventory/?" + urllib.parse.urlencode(qs))
        # 화면에서 즉시 조정(A안)
        if action == "adjust_single":
            from .models import InventoryManufacturer

            def _parse_date(val: str):
                val = (val or "").strip()
                if not val:
                    return None
                try:
                    return datetime.date.fromisoformat(val)
                except Exception:
                    return None

            item_type = (request.POST.get("item_type") or "").strip().upper()
            unit_id = (request.POST.get("unit_id") or "").strip()
            pm_id = (request.POST.get("product_model_id") or "").strip()
            progress_in = (request.POST.get("progress_status") or "완료").strip()
            kind_in = (request.POST.get("adjust_kind") or "데이터수정").strip()
            reason = (request.POST.get("reason") or "").strip()

            if not reason:
                error = "조정사유는 필수입니다."
            else:
                prog_map = {"진행중": "IN_PROGRESS", "완료": "DONE", "취소": "CANCELED"}
                progress_status = prog_map.get(progress_in, "DONE")
                kind_map = {"데이터수정": "DATA_CORRECTION", "분실": "LOST", "폐기": "DISCARD"}
                adj = kind_map.get(kind_in, "DATA_CORRECTION")

                try:
                    with transaction.atomic():
                        if item_type == "SERIAL":
                            if not unit_id:
                                raise ValueError("대상 제조번호가 없습니다")
                            u = InventoryUnit.objects.select_related("product_model__manufacturer").get(id=int(unit_id))

                            if adj in ("LOST", "DISCARD"):
                                # 분실/폐기: 출고 상태로 전환 + 로그
                                before = _inv_unit_snapshot(u)
                                if (u.status or "").strip() != "SHIPPED":
                                    u.status = "SHIPPED"
                                    u.save(update_fields=["status"])
                                after = _inv_unit_snapshot(u)
                                InventoryStockEvent.objects.create(
                                    unit=u,
                                    product_model=None,
                                    qty_delta=0,
                                    event_type="ADJUST",
                                    progress_status=progress_status,
                                    adjust_kind=adj,
                                    reason=reason,
                                    before_json=json.dumps(before, ensure_ascii=False),
                                    after_json=json.dumps(after, ensure_ascii=False),
                                    created_by=request.user,
                                )
                                ship_reason = "분실" if adj == "LOST" else "폐기"
                                InventoryStockEvent.objects.create(
                                    unit=u,
                                    product_model=None,
                                    qty_delta=0,
                                    event_type="SHIP",
                                    progress_status="DONE",
                                    reason=ship_reason,
                                    before_json="",
                                    after_json=json.dumps(after, ensure_ascii=False),
                                    created_by=request.user,
                                )
                            else:
                                # 데이터수정: 제조사/모델/제조번호/표준코드/제조일 수정
                                mfr = (request.POST.get("manufacturer") or "").strip()
                                model = (request.POST.get("model_name") or "").strip()
                                serial_no = (request.POST.get("serial_no") or "").strip()
                                standard_code = (request.POST.get("standard_code") or "").strip()
                                mfg_date = _parse_date(request.POST.get("mfg_date") or "")
                                if not (mfr and model and serial_no):
                                    raise ValueError("제조사/모델명/제조번호는 필수입니다")
                                if InventoryUnit.objects.filter(serial_no=serial_no).exclude(id=u.id).exists():
                                    raise ValueError("이미 등록된 제조번호입니다")
                                before = _inv_unit_snapshot(u)

                                mfr_obj = _inv_get_or_create_mfr(mfr)
                                pk = "HEARING_AID"
                                try:
                                    pk = getattr(u.product_model, "product_kind", "HEARING_AID") or "HEARING_AID"
                                except Exception:
                                    pk = "HEARING_AID"
                                pm = _inv_get_or_create_model(mfr_obj, model, item_type="SERIAL", product_kind=pk)

                                u.product_model = pm
                                u.serial_no = serial_no
                                u.standard_code = standard_code
                                u.mfg_date = mfg_date
                                u.save()

                                after = _inv_unit_snapshot(u)
                                InventoryStockEvent.objects.create(
                                    unit=u,
                                    product_model=None,
                                    qty_delta=0,
                                    event_type="ADJUST",
                                    progress_status=progress_status,
                                    adjust_kind="DATA_CORRECTION",
                                    reason=reason,
                                    before_json=json.dumps(before, ensure_ascii=False),
                                    after_json=json.dumps(after, ensure_ascii=False),
                                    created_by=request.user,
                                )

                        else:
                            # 수량형(QTY) 조정
                            if not pm_id:
                                raise ValueError("대상 품목이 없습니다")
                            pm = InventoryProductModel.objects.select_related("manufacturer").get(id=int(pm_id))
                            qty_delta_str = (request.POST.get("qty_delta") or "").strip()
                            try:
                                qty_delta = int(float(qty_delta_str))
                            except Exception:
                                qty_delta = 0
                            if qty_delta == 0:
                                raise ValueError("조정수량은 0이 될 수 없습니다")

                            before = _inv_qty_snapshot(pm)
                            new_qty = int(pm.qty_current or 0) + qty_delta
                            pm.item_type = "QTY"
                            pm.qty_current = new_qty
                            pm.save(update_fields=["item_type", "qty_current"])
                            _inv_check_and_notify(pm)
                            after = _inv_qty_snapshot(pm)
                            InventoryStockEvent.objects.create(
                                unit=None,
                                product_model=pm,
                                qty_delta=qty_delta,
                                event_type="ADJUST",
                                progress_status=progress_status,
                                adjust_kind="DATA_CORRECTION",
                                reason=reason,
                                before_json=json.dumps(before, ensure_ascii=False),
                                after_json=json.dumps(after, ensure_ascii=False),
                                created_by=request.user,
                            )

                    message = "조정이 반영되었습니다."
                except Exception as e:
                    error = f"조정 실패: {e}"

            # 결과를 GET으로 리다이렉트
            qs = {"tab": tab}
            if q:
                qs["q"] = q
            url = "/customers/inventory/?" + urllib.parse.urlencode(qs)
            if message:
                request.session["inv_msg"] = message
            if error:
                request.session["inv_err"] = error
            return redirect(url)

        if action in ("receive_csv", "ship_csv", "adjust_csv"):
            f = request.FILES.get("csv_file")
            if not f:
                error = "CSV 파일을 선택해주세요."
            else:
                rows = _inv_read_csv_file(f)
                if not rows:
                    error = "CSV 내용이 비어있습니다."
                else:
                    ok = 0
                    fail = 0
                    fail_rows = []
                    errors = []  # [{row:int, reason:str}]

                    # 원본 헤더 유지 + 오류사유 추가
                    header_keys = list(rows[0].keys())
                    if "오류사유" not in header_keys:
                        header_keys.append("오류사유")

                    def _row_fail(r, reason, row_no: int | None = None):
                        nonlocal fail
                        fail += 1
                        rr = dict(r)
                        rr["오류사유"] = reason
                        fail_rows.append(rr)
                        if row_no is not None:
                            errors.append({"row": int(row_no), "reason": str(reason)})

                    def _gv(r, *keys):
                        for k in keys:
                            if k in r and (r.get(k) is not None):
                                return (r.get(k) or "").strip()
                        return ""

                    def _parse_date(val: str, field_name: str):
                        val = (val or "").strip()
                        if not val:
                            return None
                        try:
                            return datetime.date.fromisoformat(val)
                        except Exception:
                            raise ValueError(f"{field_name} 포맷은 YYYY-MM-DD 입니다")

                    def _parse_int(val: str, default=1):
                        val = (val or "").strip()
                        if val == "":
                            return default
                        try:
                            return int(float(val))
                        except Exception:
                            return None

                    if action == "receive_csv":
                        # 입고: 템플릿(1.csv) 기준 / 하나라도 누락되면 전체 업로드 금지
                        required_cols = ["유형", "제조사", "제조수입업소명", "모델명", "제조번호", "표준코드", "제조일", "현재 수량"]
                        missing_cols = [c for c in required_cols if c not in header_keys]
                        if missing_cols:
                            for i, r in enumerate(rows, start=2):
                                _row_fail(r, "헤더 누락: " + ", ".join(missing_cols), i)
                        else:
                            normalized = []
                            kind_map = {
                                "보청기": "HEARING_AID",
                                "hearing_aid": "HEARING_AID",
                                "hearingaid": "HEARING_AID",
                                "돔": "DOME",
                                "dome": "DOME",
                                "리시버": "RECEIVER",
                                "receiver": "RECEIVER",
                            }

                            # 1) 전체 검증
                            for i, r in enumerate(rows, start=2):
                                kind_raw = (_gv(r, "유형", "type", "품목", "product_kind") or "").strip()
                                mfr = (_gv(r, "제조사", "manufacturer") or "").strip()
                                importer = (_gv(r, "제조수입업소명", "importer") or "").strip()
                                model = (_gv(r, "모델명", "model_name") or "").strip()
                                serial = (_gv(r, "제조번호", "serial_no") or "").strip()
                                std = (_gv(r, "표준코드", "standard_code") or "").strip()
                                mfg = (_gv(r, "제조일", "mfg_date") or "").strip()
                                qty_str = (_gv(r, "현재 수량", "현재수량", "qty") or "").strip()

                                # 필수값 체크
                                # - 공통 필수: 유형/제조사/제조수입업소명/모델명
                                # - 보청기: 제조번호 필수, 수량은 공란이면 1(항상 1)
                                # - 보청기 제외: 제조번호/표준코드/제조일 공란 허용
                                if not kind_raw:
                                    _row_fail(r, "유형은 필수입니다", i)
                                    continue
                                if not mfr:
                                    _row_fail(r, "제조사는 필수입니다", i)
                                    continue
                                # 제조수입업소명은 업로드에서 필수로 강제하지 않습니다(기존 동작 유지).
                                # 비어있으면 빈값으로 저장되며, 공단 서류/프로필 연동은 별도 흐름에서 처리합니다.
                                if not model:
                                    _row_fail(r, "모델명은 필수입니다", i)
                                    continue

                                kind_key = kind_raw.lower()
                                product_kind = kind_map.get(kind_key, kind_map.get(kind_raw, "HEARING_AID"))

                                # 보청기는 제조번호 필수 + 수량은 항상 1(공란이면 1)
                                if product_kind == "HEARING_AID":
                                    if not serial:
                                        _row_fail(r, "보청기는 제조번호가 필수입니다", i)
                                        continue
                                    if not std:
                                        _row_fail(r, "보청기는 표준코드가 필수입니다", i)
                                        continue
                                    if not mfg:
                                        _row_fail(r, "보청기는 제조일이 필수입니다", i)
                                        continue
                                    qty = 1
                                else:
                                    # 보청기 제외는 제조번호/표준코드/제조일 공란 허용
                                    # 수량형(QTY)으로 처리되는 경우에는 수량 필수
                                    if not serial:
                                        if not qty_str:
                                            _row_fail(r, "현재 수량은 필수입니다", i)
                                            continue
                                    qty = _parse_int(qty_str, default=1)
                                    if qty is None or qty <= 0:
                                        _row_fail(r, "현재 수량은 1 이상이어야 합니다", i)
                                        continue

                                # 제조일(공란 허용)
                                if mfg:
                                    try:
                                        mfg_date = _parse_date(mfg, "제조일")
                                    except Exception as e:
                                        _row_fail(r, str(e), i)
                                        continue
                                else:
                                    mfg_date = None

                                # 제조번호가 있는 품목은 수량 1로만 처리(공란이면 1)
                                if serial:
                                    qty = 1

                                normalized.append(
                                    {
                                        "row_no": i,
                                        "product_kind": product_kind,
                                        "mfr": mfr,
                                        "importer": importer,
                                        "model": model,
                                        "serial": serial,
                                        "std": std,
                                        "mfg_date": mfg_date,
                                        "qty": qty,
                                    }
                                )

                            # 2) 저장(전체 통과 시에만)
                            if not fail_rows:
                                try:
                                    with transaction.atomic():
                                        for it in normalized:
                                            mfr_obj = _inv_get_or_create_mfr(it["mfr"])
                                            if it["serial"]:
                                                if InventoryUnit.objects.filter(serial_no=it["serial"]).exists():
                                                    raise ValueError(f"이미 등록된 제조번호입니다: {it['serial']}")
                                                pm = _inv_get_or_create_model(
                                                    mfr_obj,
                                                    it["model"],
                                                    item_type="SERIAL",
                                                    product_kind=it["product_kind"],
                                                    importer_name=it["importer"],
                                                )
                                                u = InventoryUnit.objects.create(
                                                    product_model=pm,
                                                    serial_no=it["serial"],
                                                    standard_code=it["std"],
                                                    mfg_date=it["mfg_date"],
                                                    status="IN_STOCK",
                                                )
                                                InventoryStockEvent.objects.create(
                                                    unit=u,
                                                    product_model=None,
                                                    qty_delta=0,
                                                    event_type="RECEIVE",
                                                    progress_status="DONE",
                                                    reason="",
                                                    before_json="",
                                                    after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                                    created_by=request.user,
                                                )
                                            else:
                                                pm = _inv_get_or_create_model(
                                                    mfr_obj,
                                                    it["model"],
                                                    item_type="QTY",
                                                    product_kind=it["product_kind"],
                                                    importer_name=it["importer"],
                                                )
                                                pm.item_type = "QTY"
                                                before = _inv_qty_snapshot(pm)
                                                pm.qty_current = int(pm.qty_current or 0) + int(it["qty"])
                                                pm.save(update_fields=["item_type", "qty_current"])
                                                _inv_check_and_notify(pm)
                                                after = _inv_qty_snapshot(pm)
                                                InventoryStockEvent.objects.create(
                                                    unit=None,
                                                    product_model=pm,
                                                    qty_delta=int(it["qty"]),
                                                    event_type="RECEIVE",
                                                    progress_status="DONE",
                                                    reason="",
                                                    before_json=json.dumps(before, ensure_ascii=False),
                                                    after_json=json.dumps(after, ensure_ascii=False),
                                                    created_by=request.user,
                                                )

                                    ok = len(normalized)
                                except Exception as e:
                                    # 전체 실패 처리
                                    for i, r in enumerate(rows, start=2):
                                        _row_fail(r, f"저장 실패: {e}", i)

                    elif action == "ship_csv":
                        # 출고: 템플릿(3.csv) 기준도 허용
                        for r in rows:
                            kind_raw = (_gv(r, "유형", "type", "품목", "product_kind") or "").strip()
                            mfr = (_gv(r, "제조사", "manufacturer") or "").strip()
                            importer = (_gv(r, "제조수입업소명", "importer") or "").strip()
                            model = (_gv(r, "모델명", "model_name") or "").strip()
                            serial = (_gv(r, "제조번호", "serial_no") or "").strip()
                            qty_str = (_gv(r, "현재 수량", "수량", "qty") or "").strip()
                            reason = (_gv(r, "사유", "reason") or "").strip()  # 템플릿에 없으면 빈값 허용

                            kind_map = {
                                "보청기": "HEARING_AID",
                                "hearing_aid": "HEARING_AID",
                                "hearingaid": "HEARING_AID",
                                "돔": "DOME",
                                "dome": "DOME",
                                "리시버": "RECEIVER",
                                "receiver": "RECEIVER",
                            }
                            kind_key = kind_raw.lower()
                            product_kind = kind_map.get(kind_key, kind_map.get(kind_raw, "HEARING_AID"))

                            qty = _parse_int(qty_str, default=1)
                            if qty is None or qty <= 0:
                                _row_fail(r, "현재 수량/수량은 1 이상이어야 합니다")
                                continue

                            if serial:
                                # 시리얼형
                                if qty != 1:
                                    _row_fail(r, "제조번호가 있는 품목은 수량이 1이어야 합니다")
                                    continue
                                u = InventoryUnit.objects.filter(serial_no=serial).select_related("product_model__manufacturer").first()
                                if u is None:
                                    _row_fail(r, "재고에 없는 제조번호입니다")
                                    continue
                                if (u.status or "").strip() == "SHIPPED":
                                    _row_fail(r, "이미 판매 된 보청기 입니다")
                                    continue
                                try:
                                    with transaction.atomic():
                                        u.status = "SHIPPED"
                                        u.save(update_fields=["status"])
                                        InventoryStockEvent.objects.create(
                                            unit=u,
                                            product_model=None,
                                            qty_delta=0,
                                            event_type="SHIP",
                                            progress_status="DONE",
                                            reason=reason,
                                            before_json="",
                                            after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                            created_by=request.user,
                                        )
                                    ok += 1
                                except Exception as e:
                                    _row_fail(r, f"저장 실패: {e}")
                            else:
                                # 수량형
                                if not (mfr and model):
                                    _row_fail(r, "제조번호가 없는 품목은 제조사/모델명이 필수입니다")
                                    continue
                                try:
                                    with transaction.atomic():
                                        mfr_obj = _inv_get_or_create_mfr(mfr)
                                        pm = _inv_get_or_create_model(
                                            mfr_obj,
                                            model,
                                            item_type="QTY",
                                            product_kind=product_kind,
                                            importer_name=(importer or None),
                                        )
                                        pm.item_type = "QTY"
                                        before = _inv_qty_snapshot(pm)
                                        pm.qty_current = int(pm.qty_current or 0) - qty
                                        pm.save(update_fields=["item_type", "qty_current"])
                                        _inv_check_and_notify(pm)
                                        after = _inv_qty_snapshot(pm)
                                        InventoryStockEvent.objects.create(
                                            unit=None,
                                            product_model=pm,
                                            qty_delta=-qty,
                                            event_type="SHIP",
                                            progress_status="DONE",
                                            reason=reason,
                                            before_json=json.dumps(before, ensure_ascii=False),
                                            after_json=json.dumps(after, ensure_ascii=False),
                                            created_by=request.user,
                                        )
                                    ok += 1
                                except Exception as e:
                                    _row_fail(r, f"저장 실패: {e}")

                    else:
                        # adjust_csv
                        for r in rows:
                            kind_in = _gv(r, "구분", "adjust_type")
                            kind_raw = _gv(r, "유형", "type", "품목", "product_kind")
                            mfr = _gv(r, "제조사", "manufacturer")
                            importer = _gv(r, "제조수입업소명", "importer")
                            model = _gv(r, "모델명", "model_name")
                            serial_cell = _gv(r, "제조번호", "serial_no")
                            std = _gv(r, "표준코드", "standard_code")
                            mfg = _gv(r, "제조일", "mfg_date")
                            qty_delta_str = _gv(r, "조정수량", "qty_delta")
                            reason = _gv(r, "조정사유", "사유", "reason")

                            if not kind_in:
                                _row_fail(r, "구분은 필수입니다")
                                continue
                            if not reason:
                                _row_fail(r, "조정사유는 필수입니다")
                                continue

                            if not (importer or "").strip():
                                _row_fail(r, "제조수입업소명은 필수입니다")
                                continue

                            # 진행상태는 템플릿에서 제거됨: DONE 고정
                            progress_status = "DONE"

                            kind_map = {
                                "보청기": "HEARING_AID",
                                "hearing_aid": "HEARING_AID",
                                "hearingaid": "HEARING_AID",
                                "돔": "DOME",
                                "dome": "DOME",
                                "리시버": "RECEIVER",
                                "receiver": "RECEIVER",
                            }
                            kind_key = (kind_raw or "").strip().lower()
                            product_kind = kind_map.get(kind_key, kind_map.get((kind_raw or "").strip(), "HEARING_AID"))

                            kind_map = {"데이터수정": "DATA_CORRECTION", "분실": "LOST", "폐기": "DISCARD",
                                        "DATA_CORRECTION": "DATA_CORRECTION", "LOST": "LOST", "DISCARD": "DISCARD"}
                            adj = kind_map.get(kind_in.strip(), "")
                            if adj not in ("DATA_CORRECTION", "LOST", "DISCARD"):
                                _row_fail(r, "구분은 데이터수정/분실/폐기 중 하나입니다")
                                continue

                            # 날짜
                            try:
                                mfg_date = _parse_date(mfg, "제조일") if mfg else None
                            except Exception as e:
                                _row_fail(r, str(e))
                                continue

                            # 제조번호 기반: 있으면 시리얼 조정, 없으면 수량형 조정
                            if serial_cell:
                                # 시리얼 데이터수정에서 제조번호 변경 지원: "OLD->NEW" 또는 "OLD→NEW"
                                target_serial = serial_cell
                                after_serial = serial_cell
                                for sep in ("->", "→"):
                                    if sep in serial_cell:
                                        parts = [p.strip() for p in serial_cell.split(sep, 1)]
                                        target_serial = parts[0] or ""
                                        after_serial = parts[1] or ""
                                        break
                                if not target_serial:
                                    _row_fail(r, "제조번호가 올바르지 않습니다")
                                    continue
                                u = InventoryUnit.objects.filter(serial_no=target_serial).select_related("product_model__manufacturer").first()
                                if u is None:
                                    _row_fail(r, "재고에 없는 제조번호입니다")
                                    continue

                                if adj in ("LOST", "DISCARD"):
                                    try:
                                        with transaction.atomic():
                                            if (u.status or "").strip() != "SHIPPED":
                                                u.status = "SHIPPED"
                                                u.save(update_fields=["status"])

                                            InventoryStockEvent.objects.create(
                                                unit=u,
                                                product_model=None,
                                                qty_delta=0,
                                                event_type="ADJUST",
                                                progress_status=progress_status,
                                                adjust_kind=adj,
                                                reason=reason,
                                                before_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                                after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                                created_by=request.user,
                                            )
                                            ship_reason = "분실" if adj == "LOST" else "폐기"
                                            InventoryStockEvent.objects.create(
                                                unit=u,
                                                product_model=None,
                                                qty_delta=0,
                                                event_type="SHIP",
                                                progress_status="DONE",
                                                reason=ship_reason,
                                                before_json="",
                                                after_json=json.dumps(_inv_unit_snapshot(u), ensure_ascii=False),
                                                created_by=request.user,
                                            )
                                        ok += 1
                                    except Exception as e:
                                        _row_fail(r, f"저장 실패: {e}")
                                else:
                                    # DATA_CORRECTION: after 값은 현재 행 컬럼 기준
                                    if not (mfr and model):
                                        _row_fail(r, "데이터수정은 제조사/모델명이 필수입니다")
                                        continue
                                    if not after_serial:
                                        _row_fail(r, "데이터수정은 변경 제조번호가 필요합니다(예: SN001->SN001A)")
                                        continue
                                    if after_serial != u.serial_no and InventoryUnit.objects.filter(serial_no=after_serial).exclude(id=u.id).exists():
                                        _row_fail(r, "이미 등록된 제조번호입니다")
                                        continue
                                    try:
                                        with transaction.atomic():
                                            before = _inv_unit_snapshot(u)
                                            after_dict = {
                                                "item_type": "SERIAL",
                                                "manufacturer": mfr,
                                                "model_name": model,
                                                "serial_no": after_serial,
                                                "standard_code": std,
                                                "mfg_date": mfg_date.isoformat() if mfg_date else "",
                                                "status": u.status,
                                            }
                                            InventoryStockEvent.objects.create(
                                                unit=u,
                                                product_model=None,
                                                qty_delta=0,
                                                event_type="ADJUST",
                                                progress_status=progress_status,
                                                adjust_kind="DATA_CORRECTION",
                                                reason=reason,
                                                before_json=json.dumps(before, ensure_ascii=False),
                                                after_json=json.dumps(after_dict, ensure_ascii=False),
                                                created_by=request.user,
                                            )
                                            mfr_obj = _inv_get_or_create_mfr(mfr)
                                            pm = _inv_get_or_create_model(
                                                mfr_obj,
                                                model,
                                                item_type="SERIAL",
                                                product_kind=product_kind,
                                                importer_name=(importer or None),
                                            )
                                            u.product_model = pm
                                            u.serial_no = after_serial
                                            u.standard_code = std
                                            u.mfg_date = mfg_date
                                            u.save()
                                        ok += 1
                                    except Exception as e:
                                        _row_fail(r, f"저장 실패: {e}")
                            else:
                                # 수량형 조정
                                if not (mfr and model):
                                    _row_fail(r, "제조번호가 없는 품목은 제조사/모델명이 필수입니다")
                                    continue
                                qty_delta = _parse_int(qty_delta_str, default=0)
                                if qty_delta is None or qty_delta == 0:
                                    _row_fail(r, "조정수량은 0이 될 수 없습니다")
                                    continue
                                try:
                                    with transaction.atomic():
                                        mfr_obj = _inv_get_or_create_mfr(mfr)
                                        pm = _inv_get_or_create_model(
                                            mfr_obj,
                                            model,
                                            item_type="QTY",
                                            product_kind=product_kind,
                                            importer_name=(importer or None),
                                        )
                                        pm.item_type = "QTY"
                                        before = _inv_qty_snapshot(pm)
                                        new_qty = int(pm.qty_current or 0) + qty_delta
                                        if new_qty < 0:
                                            _row_fail(r, "조정 후 수량이 0 미만이 될 수 없습니다")
                                            continue
                                        pm.qty_current = new_qty
                                        pm.save(update_fields=["item_type", "qty_current"])
                                        after = _inv_qty_snapshot(pm)
                                        InventoryStockEvent.objects.create(
                                            unit=None,
                                            product_model=pm,
                                            qty_delta=qty_delta,
                                            event_type="ADJUST",
                                            progress_status=progress_status,
                                            adjust_kind="DATA_CORRECTION",
                                            reason=reason,
                                            before_json=json.dumps(before, ensure_ascii=False),
                                            after_json=json.dumps(after, ensure_ascii=False),
                                            created_by=request.user,
                                        )
                                    ok += 1
                                except Exception as e:
                                    _row_fail(r, f"저장 실패: {e}")

                    if fail_rows:
                        data = _inv_make_fail_csv_bytes(fail_rows, header_keys)
                        request.session["inv_fail_csv_b64"] = base64.b64encode(data).decode("ascii")
                        fail_available = True

                    message = f"성공 {ok} / 실패 {fail}"

                    # 입고 업로드 결과(모달 표시용) - A안: 세션에 저장 후 GET에서 꺼내 표시
                    if action == "receive_csv":
                        # 에러 행 범위 그룹화
                        def _group_error_ranges(err_list):
                            if not err_list:
                                return []
                            # reason -> sorted rows
                            buckets = {}
                            for it in err_list:
                                try:
                                    rn = int(it.get("row") or 0)
                                except Exception:
                                    rn = 0
                                rsn = (it.get("reason") or "").strip()
                                if rn <= 0 or not rsn:
                                    continue
                                buckets.setdefault(rsn, []).append(rn)
                            lines = []
                            for rsn in sorted(buckets.keys()):
                                rows_sorted = sorted(set(buckets[rsn]))
                                if not rows_sorted:
                                    continue
                                start = prev = rows_sorted[0]
                                for rn in rows_sorted[1:]:
                                    if rn == prev + 1:
                                        prev = rn
                                        continue
                                    if start == prev:
                                        lines.append(f"{start}행 : {rsn}")
                                    else:
                                        lines.append(f"{start}행~{prev}행 : {rsn}")
                                    start = prev = rn
                                if start == prev:
                                    lines.append(f"{start}행 : {rsn}")
                                else:
                                    lines.append(f"{start}행~{prev}행 : {rsn}")
                            return lines

                        total = len(rows)
                        result = {
                            "filename": getattr(f, "name", "") or "",
                            "total": int(total),
                            "success": int(ok),
                            "fail": int(fail),
                            "summary_lines": _group_error_ranges(errors),
                        }
                        request.session["inv_receive_upload_result"] = result

        # 결과를 GET으로 리다이렉트
        qs = {"tab": tab}
        if q:
            qs["q"] = q
        url = "/customers/inventory/?" + urllib.parse.urlencode(qs)
        if message:
            request.session["inv_msg"] = message
        if error:
            request.session["inv_err"] = error
        # 입고 업로드: 파일 단위 오류도 모달로 표시(1회성)
        if (request.POST.get("action") or "").strip() == "receive_csv":
            if error and not request.session.get("inv_receive_upload_result"):
                request.session["inv_receive_upload_result"] = {
                    "filename": "",
                    "total": 0,
                    "success": 0,
                    "fail": 0,
                    "summary_lines": [str(error)],
                }
        if fail_available:
            request.session["inv_fail_available"] = True
        return redirect(url)

    # messages
    if request.session.get("inv_msg"):
        message = request.session.pop("inv_msg")
    if request.session.get("inv_err"):
        error = request.session.pop("inv_err")
    # 입고 업로드 결과 모달(1회성)
    upload_result = request.session.pop("inv_receive_upload_result", None)
    fail_available = bool(request.session.pop("inv_fail_available", False)) or bool(request.session.get("inv_fail_csv_b64"))

    # list rows
    units_qs = InventoryUnit.objects.select_related("product_model__manufacturer")
    qty_qs = InventoryProductModel.objects.select_related("manufacturer").filter(item_type="QTY", is_deleted=False, qty_current__gt=0)

    if q:
        units_qs = units_qs.filter(
            Q(serial_no__icontains=q)
            | Q(standard_code__icontains=q)
            | Q(product_model__model_name__icontains=q)
            | Q(product_model__manufacturer__name__icontains=q)
        )
        qty_qs = qty_qs.filter(Q(model_name__icontains=q) | Q(manufacturer__name__icontains=q))

    from django.core.paginator import Paginator

    receive_units = []
    receive_qty_models = []
    adjust_units = []
    adjust_qty_models = []
    ship_events = []
    adjust_events = []
    ship_current_units = []
    ship_current_qty_models = []

    pager = None
    page_obj = None

    if tab == "입고":
        receive_units = list(units_qs.filter(status="IN_STOCK").order_by("-created_at", "-id")[:5000])
        receive_qty_models = list(qty_qs.order_by("manufacturer__name", "model_name")[:5000])
    elif tab == "출고":
        # 출고 탭: 현재 출고 가능 재고(입고/조정 반영) + 출고 내역
        ship_current_units = list(units_qs.filter(status="IN_STOCK").order_by("-created_at", "-id")[:5000])
        ship_current_qty_models = list(qty_qs.order_by("manufacturer__name", "model_name")[:5000])

        ship_events_qs = (
            InventoryStockEvent.objects.filter(event_type="SHIP")
            .select_related("unit__product_model__manufacturer", "product_model__manufacturer", "created_by")
            .order_by("-created_at", "-id")
        )
        pager = Paginator(ship_events_qs, per_page)
        page_obj = pager.get_page(page)
        ship_events = list(page_obj.object_list)
    else:
        # 조정 탭: 현재 재고(조정 대상) + 조정 내역
        adjust_units = list(units_qs.filter(status="IN_STOCK").order_by("-created_at", "-id")[:5000])
        adjust_qty_models_all = qty_qs.order_by("manufacturer__name", "model_name")
        pager = Paginator(adjust_qty_models_all, per_page)
        page_obj = pager.get_page(page)
        adjust_qty_models = list(page_obj.object_list)

        adjust_events_qs = (
            InventoryStockEvent.objects.filter(event_type="ADJUST")
            .select_related("unit__product_model__manufacturer", "product_model__manufacturer", "created_by")
            .order_by("-created_at", "-id")
        )
        adjust_events = list(adjust_events_qs[:200])

    # SERIAL qty_map (모델별 보유 수량)
    qty_map = {}
    try:
        for row in (
            InventoryUnit.objects.filter(status="IN_STOCK")
            .values("product_model_id")
            .annotate(cnt=Count("id"))
        ):
            qty_map[int(row["product_model_id"])] = int(row["cnt"])
    except Exception:
        qty_map = {}

    # view-friendly rows (수량은 항상 양수로)
    ship_rows = []
    for ev in ship_events:
        if ev.unit_id:
            u = ev.unit
            pm = u.product_model
            ship_rows.append({
                "제조사": pm.manufacturer.name if pm and pm.manufacturer else "",
                "모델명": pm.model_name if pm else "",
                "제조번호": u.serial_no,
                "표준코드": u.standard_code,
                "제조일": (u.mfg_date.isoformat() if u.mfg_date else ""),
                "수량": 1,
                "사유": ev.reason or "",
                "담당자": ev.created_by.get_username() if ev.created_by else "",
                "출고일": ev.created_at.date().isoformat(),
            })
        else:
            pm = ev.product_model
            ship_rows.append({
                "제조사": pm.manufacturer.name if pm and pm.manufacturer else "",
                "모델명": pm.model_name if pm else "",
                "제조번호": "",
                "표준코드": "",
                "제조일": "",
                "수량": abs(int(ev.qty_delta or 0)),
                "사유": ev.reason or "",
                "담당자": ev.created_by.get_username() if ev.created_by else "",
                "출고일": ev.created_at.date().isoformat(),
            })

    adjust_rows = []
    kind_display_map = {"DATA_CORRECTION": "데이터수정", "LOST": "분실", "DISCARD": "폐기"}
    for ev in adjust_events:
        progress_disp = ev.get_progress_status_display() if hasattr(ev, "get_progress_status_display") else (ev.progress_status or "")
        kind_disp = kind_display_map.get((ev.adjust_kind or "").strip(), ev.get_adjust_kind_display() if hasattr(ev, "get_adjust_kind_display") else (ev.adjust_kind or ""))
        if ev.unit_id:
            u = ev.unit
            pm = u.product_model
            adjust_rows.append({
                "진행상태": progress_disp,
                "구분": kind_disp,
                "제조사": pm.manufacturer.name if pm and pm.manufacturer else "",
                "모델명": pm.model_name if pm else "",
                "제조번호": u.serial_no,
                "표준코드": u.standard_code,
                "제조일": (u.mfg_date.isoformat() if u.mfg_date else ""),
                "조정수량": "",
                "조정사유": ev.reason or "",
                "담당자": ev.created_by.get_username() if ev.created_by else "",
                "조정일": ev.created_at.date().isoformat(),
            })
        else:
            pm = ev.product_model
            adjust_rows.append({
                "진행상태": progress_disp,
                "구분": kind_disp,
                "제조사": pm.manufacturer.name if pm and pm.manufacturer else "",
                "모델명": pm.model_name if pm else "",
                "제조번호": "",
                "표준코드": "",
                "제조일": "",
                "조정수량": int(ev.qty_delta or 0),
                "조정사유": ev.reason or "",
                "담당자": ev.created_by.get_username() if ev.created_by else "",
                "조정일": ev.created_at.date().isoformat(),
            })

    ctx = {
        "tab": tab,
        "q": q,
        "per_page": per_page,
        "page": page,
        "page_obj": page_obj,
        "message": message,
        "error": error,
        "upload_result_json": json.dumps(upload_result, ensure_ascii=False) if upload_result else "",
        "receive_rows": receive_units,
        "receive_qty_models": receive_qty_models,
        "adjust_current_rows": adjust_units,
        "adjust_current_qty_models": adjust_qty_models,
        "ship_current_rows": ship_current_units,
        "ship_current_qty_models": ship_current_qty_models,
        "ship_rows": ship_rows,
        "adjust_rows": adjust_rows,
        "qty_map": qty_map,
        "fail_available": fail_available,
        "is_admin": bool(request.user.is_staff or request.user.is_superuser),
    }
    return render(request, "customers/inventory.html", ctx)




# ==============================
# Inventory APIs (제품/결제 모달/제조번호 자동조회)
# ==============================


from django.http import JsonResponse


@login_required
def api_inventory_manufacturers(request: HttpRequest):
    """제조사 목록"""
    from .models import InventoryManufacturer

    kind = (request.GET.get("kind") or "").strip().upper()

    qs = InventoryManufacturer.objects.all()
    if kind in ("HEARING_AID", "DOME", "RECEIVER"):
        qs = qs.filter(product_models__product_kind=kind, product_models__is_deleted=False).distinct()

    items = list(qs.order_by("name", "id").values_list("name", flat=True))
    return JsonResponse({"items": [{"name": n} for n in items]})


@login_required
def api_inventory_units(request: HttpRequest):
    """제조사별 재고 목록

    - HEARING_AID: 제조번호(시리얼) 기반(InventoryUnit)
    - DOME/RECEIVER: 수량형 기반(InventoryProductModel)

    제품/결제 모달에서 공통 사용.
    """
    from django.db.models import Q
    from .models import InventoryUnit, InventoryProductModel

    manufacturer = (request.GET.get("manufacturer") or "").strip()
    kind = (request.GET.get("kind") or "HEARING_AID").strip().upper()
    q = (request.GET.get("q") or "").strip()
    in_stock = (request.GET.get("in_stock") or "1").strip()
    exclude_serials = (request.GET.get("exclude_serials") or "").strip()
    exclude_list = [s.strip() for s in exclude_serials.split(",") if s.strip()]

    # 수량형(돔/리시버)
    if kind in ("DOME", "RECEIVER"):
        qs = InventoryProductModel.objects.select_related("manufacturer").filter(item_type="QTY")
        qs = qs.filter(product_kind=kind)
        # soft delete 제외
        if hasattr(InventoryProductModel, "is_deleted"):
            qs = qs.filter(is_deleted=False)
        if manufacturer:
            qs = qs.filter(manufacturer__name=manufacturer)
        if q:
            qs = qs.filter(Q(model_name__icontains=q))
        qs = qs.order_by("model_name", "id")[:200]

        items = []
        for pm in qs:
            mfr = getattr(getattr(pm, "manufacturer", None), "name", "")
            items.append(
                {
                    "manufacturer": mfr,
                    "model_name": pm.model_name,
                    "qty_current": int(pm.qty_current or 0),
                    "status": "IN_STOCK" if int(pm.qty_current or 0) > 0 else "OUT_OF_STOCK",
                }
            )
        return JsonResponse({"items": items})

    # 시리얼형(보청기)
    qs = InventoryUnit.objects.select_related("product_model__manufacturer")
    if kind in ("HEARING_AID",):
        qs = qs.filter(product_model__product_kind=kind)
    if manufacturer:
        qs = qs.filter(product_model__manufacturer__name=manufacturer)
    if in_stock == "1":
        qs = qs.filter(status="IN_STOCK")
    if exclude_list:
        qs = qs.exclude(serial_no__in=exclude_list)
    if q:
        qs = qs.filter(Q(product_model__model_name__icontains=q) | Q(serial_no__icontains=q))

    # 제품/결제 모달: 제조번호 없는 항목은 제외
    qs = qs.exclude(serial_no="")

    qs = qs.order_by("product_model__model_name", "serial_no", "id")[:200]

    items = []
    for u in qs:
        pm = getattr(u, "product_model", None)
        mfr = getattr(getattr(pm, "manufacturer", None), "name", "") if pm else ""
        items.append(
            {
                "manufacturer": mfr,
                "model_name": pm.model_name if pm else "",
                "serial_no": u.serial_no,
                "standard_code": u.standard_code or "",
                "mfg_date": u.mfg_date.isoformat() if u.mfg_date else "",
                "status": u.status,
            }
        )

    return JsonResponse({"items": items})


@login_required
def api_inventory_unit_by_serial(request: HttpRequest):
    """제조번호 단건 조회 - 제품/결제 탭 자동기입용"""
    from .models import InventoryUnit

    serial = (request.GET.get("serial") or "").strip()
    kind = (request.GET.get("kind") or "").strip().upper()
    if not serial:
        return JsonResponse({"found": False, "reason": "NO_SERIAL"})

    qs = InventoryUnit.objects.select_related("product_model__manufacturer")

    # 1) 완전 일치(대소문자 무시)
    u = qs.filter(serial_no__iexact=serial).first()

    # 2) 흔한 입력 변형(공백/하이픈 제거 등) 재시도
    if not u:
        variants = [
            serial.replace(" ", ""),
            serial.replace("-", ""),
            serial.replace(" ", "").replace("-", ""),
        ]
        seen = set([serial])
        for v in variants:
            v = (v or "").strip()
            if not v or v in seen:
                continue
            seen.add(v)
            u = qs.filter(serial_no__iexact=v).first()
            if u:
                break

    # 3) 정규화 비교(공백/특수문자 차이 흡수) - 후보 50건 내 매칭
    if not u:
        norm = re.sub(r"[^0-9A-Za-z]+", "", serial).upper()
        if norm:
            chunk = norm[-6:] if len(norm) >= 6 else norm
            cands = list(qs.filter(serial_no__icontains=chunk)[:50])
            matches = []
            for cu in cands:
                cnorm = re.sub(r"[^0-9A-Za-z]+", "", (cu.serial_no or "")).upper()
                if cnorm == norm:
                    matches.append(cu)
            if len(matches) == 1:
                u = matches[0]
            elif len(matches) > 1:
                return JsonResponse({
                    "found": False,
                    "reason": "MULTIPLE_MATCHES",
                    "items": [
                        {
                            "manufacturer": getattr(getattr(getattr(mu, "product_model", None), "manufacturer", None), "name", ""),
                            "model_name": getattr(getattr(mu, "product_model", None), "model_name", ""),
                            "serial_no": mu.serial_no,
                            "mfg_date": mu.mfg_date.isoformat() if mu.mfg_date else "",
                        }
                        for mu in matches[:10]
                    ],
                })

    if not u:
        return JsonResponse({"found": False, "reason": "NOT_FOUND"})

    pm = getattr(u, "product_model", None)
    try:
        if kind in ("HEARING_AID", "DOME", "RECEIVER") and pm and getattr(pm, "product_kind", "") and pm.product_kind != kind:
            return JsonResponse({"found": False, "reason": "KIND_MISMATCH"})
    except Exception:
        pass
    mfr = getattr(getattr(pm, "manufacturer", None), "name", "") if pm else ""

    return JsonResponse(
        {
            "found": True,
            "manufacturer": mfr,
            "model_name": pm.model_name if pm else "",
            "serial_no": u.serial_no,
            "standard_code": u.standard_code or "",
            "mfg_date": u.mfg_date.isoformat() if u.mfg_date else "",
            "status": u.status,
        }
    )


@login_required
@require_POST
def api_notifications_mark_read(request: HttpRequest):
    """재고 알림 읽음 처리 (방문 알림은 DB 저장 안함)"""
    import json

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        payload = {}

    kind = (payload.get("kind") or "INVENTORY").strip().upper()
    if kind != "INVENTORY":
        kind = "INVENTORY"

    try:
        from .models import Notification
        from datetime import timedelta
        from django.utils import timezone

        cutoff = timezone.now() - timedelta(days=30)
        Notification.objects.filter(kind=kind, is_read=False, created_at__gte=cutoff).update(is_read=True)
        return JsonResponse({"ok": True})
    except Exception:
        return JsonResponse({"ok": True})


@login_required
def settings_inventory_history(request: HttpRequest) -> HttpResponse:
    """설정 > 재고 히스토리 (통합 로그)"""
    import csv
    import io
    import urllib.parse
    from django.http import HttpResponse

    from .models import InventoryStockEvent

    q = (request.GET.get("q") or "").strip()
    download = (request.GET.get("download") or "").strip()

    qs = (
        InventoryStockEvent.objects
        .select_related("unit__product_model__manufacturer", "product_model__manufacturer", "created_by")
        .order_by("-created_at", "-id")
    )
    if q:
        from django.db.models import Q
        qs = qs.filter(
            Q(unit__serial_no__icontains=q)
            | Q(unit__standard_code__icontains=q)
            | Q(unit__product_model__model_name__icontains=q)
            | Q(unit__product_model__manufacturer__name__icontains=q)
            | Q(product_model__model_name__icontains=q)
            | Q(product_model__manufacturer__name__icontains=q)
            | Q(reason__icontains=q)
        )

    if download == "1":
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(["일시", "구분", "유형", "제조사", "모델명", "제조번호", "표준코드", "제조일", "수량변동", "사유", "담당자"])
        for ev in qs[:5000]:
            dt = timezone.localtime(ev.created_at)
            ev_type = ev.event_type
            kind = "보청기"
            mfr = ""
            model = ""
            serial = ""
            std = ""
            mfg = ""
            qty_delta = ev.qty_delta or 0
            if ev.unit_id:
                u = ev.unit
                pm = u.product_model
                kind = "보청기" if getattr(pm, "product_kind", "HEARING_AID") == "HEARING_AID" else ("돔" if pm.product_kind == "DOME" else "리시버")
                mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                model = pm.model_name if pm else ""
                serial = u.serial_no
                std = u.standard_code
                mfg = u.mfg_date.isoformat() if u.mfg_date else ""
                if ev_type == "RECEIVE":
                    qty_delta = 1
                elif ev_type in ("SHIP", "DISCARD"):
                    qty_delta = -1
                else:
                    qty_delta = 0
            else:
                pm = ev.product_model
                kind = "보청기" if getattr(pm, "product_kind", "HEARING_AID") == "HEARING_AID" else ("돔" if pm.product_kind == "DOME" else "리시버")
                mfr = pm.manufacturer.name if pm and pm.manufacturer else ""
                model = pm.model_name if pm else ""
            w.writerow([
                dt.strftime("%Y-%m-%d %H:%M"),
                ev_type,
                kind,
                mfr,
                model,
                serial,
                std,
                mfg,
                int(qty_delta),
                ev.reason or "",
                ev.created_by.get_username() if ev.created_by else "",
            ])

        ymd = timezone.localdate().strftime("%Y%m%d")
        filename = f"inventory_history_{ymd}.csv"
        data = buf.getvalue().encode("utf-8-sig")
        resp = HttpResponse(data, content_type="text/csv; charset=utf-8")
        quoted = urllib.parse.quote(filename)
        resp["Content-Disposition"] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quoted}'
        return resp

    from django.core.paginator import Paginator
    try:
        per_page = int((request.GET.get("per_page") or "50").strip())
    except Exception:
        per_page = 50
    if per_page not in (30, 50, 100):
        per_page = 50
    try:
        page = int((request.GET.get("page") or "1").strip())
    except Exception:
        page = 1
    if page < 1:
        page = 1
    pager = Paginator(qs, per_page)
    page_obj = pager.get_page(page)

    ctx = {"q": q, "per_page": per_page, "page": page, "page_obj": page_obj}
    return render(request, "customers/inventory_history.html", ctx)



# ==========================
# 서류 출력(일반: 공단/후기)
# ==========================
import zipfile
from pathlib import Path


def _doc_template_dir() -> Path:
    # customers/doc_templates/general/
    return Path(__file__).resolve().parent / "doc_templates" / "general"


def _safe_filename(name: str) -> str:
    # 윈도우 파일명 금지문자 제거
    bad = r'\\/:*?"<>|'
    for ch in bad:
        name = name.replace(ch, "_")
    name = name.strip()
    return name or "documents"


def _format_date(d) -> str:
    try:
        if not d:
            return ""
        return d.strftime("%Y-%m-%d")
    except Exception:
        return ""


def _format_year_only(d) -> str:
    """날짜에서 연도(YYYY)만 반환합니다. (월/일은 서류에서 직접 기입)"""
    try:
        if not d:
            return ""
        return str(d.year)
    except Exception:
        return ""


def _fmt_int(v) -> str:
    """금액 표기: 콤마 포함."""
    try:
        return f"{int(v):,}"
    except Exception:
        return ""


def _parse_birth_from_rrn(rrn: str) -> str:
    """주민등록번호(YYMMDD-XXXXXXX)에서 생년월일(YYYY-MM-DD)을 뽑습니다."""
    if not rrn:
        return ""
    s = re.sub(r"[^0-9]", "", rrn)
    if len(s) < 7:
        return ""
    yymmdd = s[:6]
    gender = s[6]
    try:
        yy = int(yymmdd[:2])
        mm = int(yymmdd[2:4])
        dd = int(yymmdd[4:6])
    except Exception:
        return ""

    # 1,2: 1900 / 3,4: 2000 (외국인 등은 범위가 더 넓지만 일반 케이스 우선)
    century = 1900
    if gender in ("3", "4", "7", "8"):
        century = 2000
    year = century + yy

    try:
        datetime.date(year, mm, dd)
    except Exception:
        return ""
    return f"{year:04d}-{mm:02d}-{dd:02d}"




def _resolve_importer_name_for_case(case: CustomerCase) -> str:
    """제조(수입)업소명: 재고 모델의 importer_name을 우선 사용합니다.

    - 제품/결제에서 제조사가 '벨톤'인 경우 요구사항에 따라 제품별 제조수입업소명을 사용합니다.
    - 해당 모델에 importer_name이 있으면 이를 사용합니다.
    """
    try:
        mfr = (case.manufacturer or "").strip()
        model = (case.model_name or "").strip()
        if not model:
            return ""
        qs = InventoryProductModel.objects.filter(is_deleted=False, model_name=model)
        if mfr:
            qs = qs.filter(manufacturer__name__iexact=mfr)
        pm = qs.first()
        if pm and (pm.importer_name or "").strip():
            return (pm.importer_name or "").strip()
        return ""
    except Exception:
        return ""

def _build_general_context(request: HttpRequest, case: CustomerCase) -> dict[str, str]:
    customer = case.customer
    bp = BusinessProfile.objects.filter(user=request.user).first()

    nhis_amount = int(case.nhis_amount or 0)
    copay_amount = int(case.copay_amount or 0)
    main_sum = nhis_amount + copay_amount
    grand_total = int(getattr(case, "grand_total", 0) or 0)

    level = (customer.exam_disability_level or "").strip()
    severe = level == "심각한 장애"
    not_severe = level == "심각하지 않은 장애"

    ctx: dict[str, str] = {
        # 날짜(연도만)
        "청구_연도": str(timezone.localdate().year),
        "청구연도": str(timezone.localdate().year),

        # 고객
        "고객명": customer.name or "",
        "연락처": customer.phone or "",
        "주소": customer.address_summary or "",
        "주민등록번호": customer.rrn_full or "",
        "생년월일": _parse_birth_from_rrn(customer.rrn_full or ""),

        # 장애 체크(표준 + 기존호환)
        "심한장애_체크": "√" if severe else "",
        "심하지않은장애_체크": "√" if not_severe else "",
        "일반_심한장애_체크": "√" if severe else "",
        "일반_심하지않은장애_체크": "√" if not_severe else "",

        # 사업자(센터)
        "사업자명": (bp.business_name if bp else "") or "",
        "대표자명": (bp.representative_name if bp else "") or "",
        "사업자등록번호": (bp.business_reg_no if bp else "") or "",
        "사업자_연락처": (bp.business_phone if bp else "") or "",
        "사업자_주소": (bp.business_address if bp else "") or "",
        "사업자_업태": (bp.business_type if bp else "") or "",
        "사업자_종목": (bp.business_item if bp else "") or "",
        "은행명": (bp.bank_name if bp else "") or "",
        "계좌번호": (bp.bank_account if bp else "") or "",

        # 급여지급청구서 하단(전체 표기)
        "사업자_주민등록번호": (bp.rep_rrn_full if bp else "") or "",

        # 안내 문구
        "날짜_안내": "※ 날짜는 YYYY-MM-DD 형식으로 자동 반영됩니다.",

        # 제품(메인만 공단 청구용)
        "모델명_대표": case.model_name or "",
        "모델명": case.model_name or "",
        "제조번호": case.serial_number or "",
        "표준코드": case.standard_code or "",
        "제조일": _format_date(case.manufacture_date),
        "구입일": _format_date(case.purchase_date),
        "착용일_구매일": _format_date(case.purchase_date),
        "리시버": case.receiver or "",
        "제조수입업소명": case.manufacturer or "",
        "제조_수입_업소명": "",

        # 표준계약서(좌/우 라인)
        "좌_모델명": "",
        "우_모델명": "",
        "좌_제조사": "",
        "우_제조사": "",
        "좌_단가": "",
        "우_단가": "",

        # 합계(계)
        "계": _fmt_int(grand_total),

        # 금액
        "공단인정금액": _fmt_int(nhis_amount),
        "공단_인정금액": _fmt_int(nhis_amount),
        "본인부담액": _fmt_int(copay_amount),
        "메인_합계": _fmt_int(main_sum),
        "실구입금액": _fmt_int(main_sum),
        "고시금액": _fmt_int(nhis_amount),
        "청구금액": _fmt_int(nhis_amount),
        "기준액": "1110000",
    }

    # 표준계약서: 메인 좌/우 선택에 따라 기입 라인 결정
    side = (case.side or "").strip()
    if side not in ("좌", "우"):
        side = "좌"
    unit_price = _fmt_int(main_sum)
    if side == "좌":
        ctx["좌_모델명"] = case.model_name or ""
        ctx["좌_제조사"] = case.manufacturer or ""
        ctx["좌_단가"] = unit_price
    else:
        ctx["우_모델명"] = case.model_name or ""
        ctx["우_제조사"] = case.manufacturer or ""
        ctx["우_단가"] = unit_price

    # 템플릿/기존 호환 키
    ctx.setdefault("제조사", case.manufacturer or "")
    ctx.setdefault("단가", unit_price)
    ctx.setdefault("총결제금액", ctx.get("계", ""))

    # 기존 템플릿 호환
    ctx.setdefault("사업자번호", ctx.get("사업자등록번호", ""))
    ctx.setdefault("사업자_유선", ctx.get("사업자_연락처", ""))
    ctx.setdefault("사업자주소", ctx.get("사업자_주소", ""))
    ctx.setdefault("사업자_주소", ctx.get("사업자_주소", ""))

    # 제조(수입)업소명: 제품별 importer_name 우선
    importer_name = _resolve_importer_name_for_case(case)
    if importer_name:
        ctx["제조_수입_업소명"] = importer_name
    else:
        ctx["제조_수입_업소명"] = (case.manufacturer or "")

    return ctx


def _fill_docx_bytes(template_path: Path, ctx: dict[str, str]) -> bytes:
    """docx 텍스트 치환.

    지원 범위:
    - {{키}} 형태의 텍스트 치환
    - Word 메일머지 MERGEFIELD(필드) 결과값 치환

    주의:
    - 서식 보존보다 "값이 제대로 들어가는 것"을 우선합니다.
    """
    from docx import Document
    from xml.sax.saxutils import escape as _xml_escape

    repl = {**{f"{{{{{k}}}}}": (v or "") for k, v in ctx.items()}, **{f"{{{k}}}": (v or "") for k, v in ctx.items()}}

    def _replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            if not p.runs:
                continue
            full = "".join(r.text for r in p.runs)
            new = full
            for k, v in repl.items():
                if k in new:
                    new = new.replace(k, v)
            if new != full:
                # 포맷은 단순화되지만, 변수 치환 정확도를 우선
                p.text = new

    def _replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)
                for t in cell.tables:
                    _replace_in_table(t)

    with open(template_path, "rb") as f:
        doc = Document(io.BytesIO(f.read()))

    _replace_in_paragraphs(doc.paragraphs)
    for t in doc.tables:
        _replace_in_table(t)

    # 헤더/푸터도 치환
    for sec in doc.sections:
        _replace_in_paragraphs(sec.header.paragraphs)
        _replace_in_paragraphs(sec.footer.paragraphs)
        for t in sec.header.tables:
            _replace_in_table(t)
        for t in sec.footer.tables:
            _replace_in_table(t)

    out = io.BytesIO()
    doc.save(out)
    filled = out.getvalue()

    # ✅ Word 메일머지(MERGEFIELD) 결과값 치환
    # - aftercare 템플릿은 MERGEFIELD를 사용하므로, docx(zip)의 word/*.xml을 직접 수정
    try:
        repl_field = {str(k): (v or "") for k, v in ctx.items()}

        def _replace_mergefields_in_xml(xml: str) -> str:
            # fldSimple / complex field 모두 지원 (표시 결과 텍스트(w:t)만 교체)
            for key, val in repl_field.items():
                safe_val = _xml_escape(val)

                # (1) <w:fldSimple ... w:instr="...MERGEFIELD key ..."> ... <w:t>...</w:t>
                p_simple = re.compile(
                    rf'(<w:fldSimple[^>]*w:instr="[^"]*MERGEFIELD\s+{re.escape(key)}[^\"]*"[^>]*>.*?<w:t[^>]*>)(.*?)(</w:t>)',
                    re.DOTALL,
                )
                xml = p_simple.sub(rf'\1{safe_val}\3', xml)

                # (2) Complex field: instrText ... MERGEFIELD key ... separate ... <w:t>...</w:t>
                p_complex = re.compile(
                    rf'(<w:instrText[^>]*>[^<]*MERGEFIELD\s+{re.escape(key)}[^<]*</w:instrText>.*?<w:fldChar[^>]*w:fldCharType="separate"[^>]*/>.*?<w:t[^>]*>)(.*?)(</w:t>)',
                    re.DOTALL,
                )
                xml = p_complex.sub(rf'\1{safe_val}\3', xml)

            return xml

        in_bio = io.BytesIO(filled)
        out_bio = io.BytesIO()
        with zipfile.ZipFile(in_bio, "r") as zin, zipfile.ZipFile(out_bio, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        txt = data.decode("utf-8")
                    except Exception:
                        zout.writestr(item, data)
                        continue

                    if "MERGEFIELD" in txt:
                        txt = _replace_mergefields_in_xml(txt)
                        data = txt.encode("utf-8")
                zout.writestr(item, data)

        filled = out_bio.getvalue()
    except Exception:
        # 메일머지 치환 실패 시에도 {{}} 치환 결과는 유지
        pass

    return filled




def _collect_missing_items_for_documents(user, case: CustomerCase) -> list[str]:
    """다운로드는 허용하되, 누락 항목을 모달로 안내하기 위한 리스트를 반환합니다."""
    missing: list[str] = []
    bp = BusinessProfile.objects.filter(user=user).first()

    def _chk(label: str, val: str | None):
        if not (val or "").strip():
            missing.append(label)

    if not bp:
        missing.extend([
            "사업자명",
            "대표자명",
            "사업자등록번호",
            "사업자 연락처",
            "사업자 주소",
            "은행",
            "계좌번호",
            "대표자 주민등록번호(전체)",
        ])
    else:
        _chk("사업자명", bp.business_name)
        _chk("대표자명", bp.representative_name)
        _chk("사업자등록번호", bp.business_reg_no)
        _chk("사업자 연락처", bp.business_phone)
        _chk("사업자 주소", bp.business_address)
        _chk("은행", bp.bank_name)
        _chk("계좌번호", bp.bank_account)
        _chk("대표자 주민등록번호(전체)", bp.rep_rrn_full)

    if (case.manufacturer or "").strip() == "벨톤":
        importer_name = _resolve_importer_name_for_case(case)
        if not importer_name:
            missing.append("제조(수입)업소명(벨톤 제품: 재고 모델에 등록 필요)")

    return missing


def _download_wrapper_response(request: HttpRequest, direct_url: str, missing_items: list[str]) -> HttpResponse:
    """다운로드 wrapper (iframe 전용).

    단일 모달 원칙:
    - 서버는 UI(토스트/모달/페이지)를 렌더링하지 않습니다.
    - iframe 내부에서 다운로드를 트리거하고, 누락 항목은 postMessage로 부모에게 전달합니다.
    """
    missing_json = json.dumps(missing_items, ensure_ascii=False)
    # direct_url은 JS 문자열로 안전하게 넣기 위해 json.dumps 사용
    direct_url_js = json.dumps(direct_url)

    body = f"""<!doctype html>
<html lang=\"ko\">
<head>
  <meta charset=\"utf-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
  <title>download</title>
</head>
<body style=\"margin:0;padding:0;background:transparent;overflow:hidden;\">
  <script>
    (function(){{
      const missingItems = {missing_json};
      try {{
        if (window.parent && window.parent !== window) {{
          window.parent.postMessage({{ type: 'DOC_DOWNLOAD', ok: true, missingItems: missingItems }}, '*');
        }}
      }} catch (e) {{}}

      // iframe 내부에서 파일 다운로드 트리거
      try {{ window.location.replace({direct_url_js}); }} catch (e) {{ window.location.href = {direct_url_js}; }}
    }})();
  </script>
</body>
</html>"""

    return HttpResponse(body, content_type="text/html; charset=utf-8")

def _zip_response(files: list[tuple[str, bytes]], zip_name: str) -> HttpResponse:
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for fname, b in files:
            z.writestr(fname, b)
    resp = HttpResponse(bio.getvalue(), content_type="application/zip")
    resp["Content-Disposition"] = f'attachment; filename="{_safe_filename(zip_name)}.zip"'
    return resp


@login_required
def download_general_public_documents(request: HttpRequest, case_id: int) -> HttpResponse:
    case = get_object_or_404(CustomerCase, id=case_id, customer__is_deleted=False)
    track = (case.customer.track or "").strip()
    if track not in ("일반", "의료", "차상위"):
        return HttpResponseForbidden("잘못된 구분값입니다.")

    # 다운로드는 항상 진행하되, 프로필 누락/매핑 누락은 모달로 안내
    if request.GET.get('direct') != '1':
        missing_items = _collect_missing_items_for_documents(request.user, case)
        q = request.GET.copy()
        q['direct'] = '1'
        direct_url = request.path + ('?' + q.urlencode() if q else '')
        return _download_wrapper_response(request, direct_url, missing_items)

    ctx = _build_general_context(request, case)

    # 장애정도(후기 의료 급여청구서: 체크가 아닌 텍스트 표기)
    level_txt = (case.customer.exam_disability_level or '').strip()
    ctx.update({
        '심각한 장애': level_txt if level_txt == '심각한 장애' else '',
        '심하지 않은 장애': level_txt if level_txt == '심각하지 않은 장애' else '',
    })
    # 공통: 일부 서식 호환 키 매핑
    ctx["사업자번호"] = ctx.get("사업자등록번호", "")

    # ✅ 의료/차상위: 없는 값은 대체하지 않고 공란 유지(요구사항)
    if track in ("의료", "차상위"):
        importer_name = _resolve_importer_name_for_case(case)
        ctx["제조_수입_업소명"] = importer_name or ""

    log = DocumentDownloadLog.objects.create(
        case=case,
        user=request.user,
        document_type="GENERAL_PUBLIC",
        round_no=None,
        status="START",
    )

    try:
        tdir = _doc_template_dir()
        files: list[tuple[str, bytes]] = []
        files.append(("01_거래명세서.docx", _fill_docx_bytes(tdir / "invoice_public.docx", ctx)))
        files.append(("02_표준계약서.docx", _fill_docx_bytes(tdir / "contract_public.docx", ctx)))
        files.append(("03_위임장.docx", _fill_docx_bytes(tdir / "power_public.docx", ctx)))

        # ✅ 구분(track)별 추가 서류 템플릿 선택
        if track == "일반":
            extra_template = "claim_public.docx"
            extra_filename = "04_공단_보조기기_급여_지급청구서.docx"
        elif track == "의료":
            extra_template = "claim_medical.docx"
            extra_filename = "04_의료_보조기기_급여비_지급청구서.docx"
        else:  # 차상위
            extra_template = "claim_lowincome.docx"
            extra_filename = "04_차상위_보조기기_급여_지급청구서.docx"

        files.append((extra_filename, _fill_docx_bytes(tdir / extra_template, ctx)))

        log.status = "SUCCESS"
        log.save(update_fields=["status"])
        # ✅ ZIP 파일명은 영문(ASCII)로 고정
        today = timezone.localdate().strftime('%Y%m%d')
        zip_name = f"nhis_documents_case{case.id}_{today}"
        return _zip_response(files, zip_name)
    except Exception as e:
        log.status = "FAIL"
        log.error_message = str(e)
        log.save(update_fields=["status", "error_message"])
        return HttpResponseBadRequest("서류 생성 중 오류가 발생했습니다.")


@login_required
def download_general_aftercare_documents(request: HttpRequest, case_id: int, round_no: int) -> HttpResponse:
    case = get_object_or_404(CustomerCase, id=case_id, customer__is_deleted=False)
    track = (case.customer.track or "").strip()
    if track not in ("일반", "의료"):
        return HttpResponseForbidden("지원하지 않는 구분값입니다.")

    round_no = int(round_no)
    if round_no not in (1, 2, 3, 4):
        return HttpResponseBadRequest("잘못된 차수입니다.")

    # 다운로드는 항상 진행하되, 프로필 누락/매핑 누락은 모달로 안내
    if request.GET.get("direct") != "1":
        missing_items = _collect_missing_items_for_documents(request.user, case)
        q = request.GET.copy()
        q["direct"] = "1"
        direct_url = request.path + ("?" + q.urlencode() if q else "")
        return _download_wrapper_response(request, direct_url, missing_items)

    ctx = _build_general_context(request, case)

    # 장애정도(후기 의료 급여청구서: 텍스트 표기)
    level_txt = (case.customer.exam_disability_level or '').strip()
    ctx['장애정도'] = level_txt

    # 공통: 후기 거래명세서 고정 금액
    ctx.update({
        "후기_총액": "50,000",
        "후기_본인부담액": "5,000",
        "후기_공급가액": "40,909",
        "후기_세액": "4,091",
        # 거래명세서(후기) 호환 키
        "메인_합계": "50,000",
        "실구입금액": "50,000",
        "본인부담액": "5,000",
        "공급가액": "40,909",
        "세액": "4,091",
    })

    today_date = timezone.localdate()
    side = (case.side or "").strip()
    if side not in ("좌", "우"):
        side = "좌"

    purchase_date = getattr(case, "purchase_date", None)
    if purchase_date:
        purchase_date_kor = f"{purchase_date.year}년{purchase_date.month}월{purchase_date.day}일"
    else:
        purchase_date_kor = ""

    fu_manager = getattr(case, f"fu{round_no}_manager", "") or ""
    fu_progress_date = getattr(case, f"fu{round_no}_progress_date", None)
    if fu_progress_date:
        fu_progress_date_kor = f"{fu_progress_date.year}년{fu_progress_date.month}월{fu_progress_date.day}일"
    else:
        fu_progress_date_kor = ""

    ctx.update({
        "청구_월": str(today_date.month),
        "청구_일": str(today_date.day),

        # 체크박스(차수)
        "회차1": "√" if round_no == 1 else "",
        "회차2": "√" if round_no == 2 else "",
        "회차3": "√" if round_no == 3 else "",
        "회차4": "√" if round_no == 4 else "",

        # 착용(메인 좌/우 기준)
        "착용좌": "√" if side == "좌" else "",
        "착용우": "√" if side == "우" else "",
        "우": "√" if side == "우" else "",
        "좌": "√" if side == "좌" else "",

        # 적합관리일자/제공자
        "적합관리일자": fu_progress_date_kor,
        "후기담당자": fu_manager,

        # 템플릿 호환 키(기존/변형)
        "후기적합 진행일자": fu_progress_date_kor,
        "후기적합 담당자": fu_manager,
        "착용일/구매일": purchase_date_kor,
        "착용일_구매일": purchase_date_kor,

        # 의료 확인서: N차 출력 시 해당 차수 1행만 채우고 나머지 공란
        "후기적합담당자": fu_manager,
        "후기적합담당자1": fu_manager if round_no == 1 else "",
        "후기적합담당자2": fu_manager if round_no == 2 else "",
        "후기적합담당자3": fu_manager if round_no == 3 else "",
        "후기적합담당자4": fu_manager if round_no == 4 else "",
    })

    log = DocumentDownloadLog.objects.create(
        case=case,
        user=request.user,
        document_type="GENERAL_AFTERCARE",
        round_no=round_no,
        status="START",
    )

    try:
        tdir = _doc_template_dir()
        files = []

        # 공통(거래명세서) + TRACK별 추가 서류
        files.append(("01_후기_거래명세서.docx", _fill_docx_bytes(tdir / "invoice_aftercare.docx", ctx)))

        if track == "일반":
            files.append(("02_후기_일반_급여청구서.docx", _fill_docx_bytes(tdir / "aftercare_claim.docx", ctx)))
            files.append(("03_후기_일반_위임장.docx", _fill_docx_bytes(tdir / "aftercare_power.docx", ctx)))
        else:  # 의료
            files.append(("02_후기_의료_급여청구서.docx", _fill_docx_bytes(tdir / "aftercare_claim_medical.docx", ctx)))
            files.append(("03_후기_의료_확인서.docx", _fill_docx_bytes(tdir / "aftercare_confirm_medical.docx", ctx)))

        log.status = "SUCCESS"
        log.save(update_fields=["status"])

        today = timezone.localdate().strftime("%Y%m%d")
        zip_name = f"aftercare_documents_{track}_round{round_no}_case{case.id}_{today}"
        return _zip_response(files, zip_name)
    except Exception as e:
        log.status = "FAIL"
        log.error_message = str(e)
        log.save(update_fields=["status", "error_message"])
        return HttpResponseBadRequest("서류 생성 중 오류가 발생했습니다.")
